"""
Script tạo phần 6.5: Hướng dẫn sử dụng tính năng Trợ lý AI (AI Assistant) cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG 6.5 =====================

def write_section_6_5(doc):
    """6.5. Hướng dẫn sử dụng tính năng Trợ lý AI (AI Assistant)"""

    add_heading1(doc, "6.5. Hướng dẫn sử dụng các tính năng Trợ lý AI (Bizen AI)")

    add_para(doc,
        "Điểm khác biệt vượt trội của hệ thống quản trị Bizen KPI/OKR so với các phần mềm truyền thống "
        "là việc tích hợp Trí tuệ nhân tạo qua model gateway có kiểm soát quyền và nguồn. Trợ lý AI không hoạt động độc lập "
        "mà đóng vai trò là một tác nhân hỗ trợ xuyên suốt (Co-pilot), tự động liên kết dữ liệu ngữ cảnh "
        "để tư vấn, sinh chỉ tiêu và cảnh báo rủi ro tiến độ theo thời gian thực."
    )

    # 6.5.1. Vai trò của Trợ lý AI trong hệ thống
    add_heading2(doc, "6.5.1. Vai trò của Trợ lý AI trong hệ thống")
    add_para(doc, "Trợ lý AI hỗ trợ người dùng ở các góc độ vận hành sau:")
    add_bullet(doc, "Cung cấp khung hội thoại tự nhiên trượt mở ở mọi màn hình, giúp nhân sự hỏi đáp giải quyết khó khăn ngay tại chỗ.", bold_prefix="Tư vấn ngữ cảnh thời gian thực")
    add_bullet(doc, "Giúp Trưởng phòng sinh nhanh các chỉ tiêu KPI định lượng phù hợp từ mục tiêu phòng ban, tránh tình trạng thiết lập KPI mơ hồ.", bold_prefix="Tự động hóa thiết lập mục tiêu")
    add_bullet(doc, "Giúp Giám đốc đọc hiểu lượng dữ liệu khổng lồ của doanh nghiệp, phân tích nguyên nhân chậm tiến độ và đưa ra đề xuất cải tiến.", bold_prefix="Hỗ trợ ra quyết định chiến lược")

    # 6.5.2. Hướng dẫn sử dụng các tính năng AI chính
    add_heading2(doc, "6.5.2. Hướng dẫn vận hành các tính năng AI chính")

    # Tính năng 1
    add_heading3(doc, "a) Widget Chat hỗ trợ chuyên môn (Bizen AI Chatbot)")
    add_para(doc,
        "Widget Chat hỗ trợ giải quyết khó khăn nghiệp vụ dựa trên ngữ cảnh thực tế của tài khoản đăng nhập:\n"
        "1. Click biểu tượng Chatbot màu hồng ở góc dưới bên phải màn hình để trượt mở panel chat.\n"
        "2. Nhập câu hỏi nghiệp vụ hoặc click các phím tắt nhanh được thiết kế sẵn (Ví dụ: 'Phân tích KPI của tôi', 'Tìm rủi ro tiến độ').\n"
        "3. Chat Advisor đóng gói snapshot KPI/OKR và check-in đã duyệt trong đúng phạm vi; RAG chỉ lấy tài liệu qua tenant/ACL do server sinh.\n"
        "4. Model phải trả câu trả lời có source ID hợp lệ hoặc chủ động báo thiếu dữ liệu. Panel escape nội dung, hiển thị citation và chỉ tư vấn; "
        "không tự thay đổi dữ liệu nghiệp vụ.",
        italic=False
    )

    # Tính năng 2
    add_heading3(doc, "b) Công cụ sinh chỉ tiêu KPI thông minh (AI KPI Generator)")
    add_para(doc,
        "Giúp Trưởng phòng sinh nhanh KPI chuẩn hóa khi giao chỉ tiêu cho nhân viên:\n"
        "1. Tại biểu mẫu 'Giao KPI mới' (URL: `/KPIs/Create`), chọn OKR phòng ban cần liên kết.\n"
        "2. Click nút 'AI gợi ý KPI', chọn kỳ đang mở và phạm vi được phép. Server chỉ gửi snapshot tối thiểu, không gửi tên, mã, email hoặc số điện thoại nhân viên.\n"
        "3. Advisor trả 3-5 bản nháp có nguồn gồm tên KPI, target, đơn vị, chiều và ngưỡng đạt/trượt đã được server kiểm tra. "
        "Quản lý click 'Áp dụng bản nháp', kiểm tra lại phạm vi/trọng số rồi tự gửi biểu mẫu; AI không tự tạo KPI chính thức.",
        italic=False
    )

    # Tính năng 3
    add_heading3(doc, "c) Phân tích hiệu suất doanh nghiệp bằng AI (AI Performance Analysis)")
    add_para(doc,
        "Giúp ban Giám đốc đánh giá nhanh sức khỏe doanh nghiệp cuối tháng hoặc cuối kỳ:\n"
        "1. Tại trang Dashboard chính của Giám đốc (URL: `/Dashboard`), click nút 'AI Performance Analysis'.\n"
        "2. AIDataService chỉ tổng hợp check-in đã duyệt thuộc tenant và phạm vi người dùng được phép xem; "
        "nếu chưa có tiến độ đo lường thì hệ thống dừng và thông báo thiếu bằng chứng.\n"
        "3. Model gateway trả strict JSON gồm tổng quan, điểm mạnh, rủi ro, hành động và citation. "
        "Kết quả chỉ mang tính tham khảo, không tự đổi điểm, xếp loại, trạng thái duyệt hoặc thưởng.",
        italic=False
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 6.5. Vai trò AI Assistant")
    print("      ✓ 6.5.1. Vai trò của AI trong hệ thống")
    print("      ✓ 6.5.2. Thao tác các tính năng AI (Widget, Generator, Analysis)")
    write_section_6_5(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 6.5 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
