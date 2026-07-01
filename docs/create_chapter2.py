"""
Script tạo CHƯƠNG 2: PHÂN TÍCH HỆ THỐNG - Phần 2.1 Yêu cầu người dùng / Danh sách chức năng
Dựa trên phân tích hệ thống thực tế và cấu trúc mẫu tốt nghiệp FPT
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    run = p.add_run(text.upper())
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Căn lề: cột 0 (STT), cột 1 (Mã UC), cột 3 (Tác nhân) căn giữa. Cột 2 (Tên UC), cột 4 (Mô tả) căn trái.
            if c_idx in [0, 1, 3]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    run = p.add_run(caption)
    set_font(run, size=12, italic=True)


# ===================== NỘI DUNG CHƯƠNG 2 =====================

def write_chapter2(doc):
    """CHƯƠNG 2: PHÂN TÍCH HỆ THỐNG - 2.1 Yêu cầu người dùng / Danh sách chức năng"""

    add_chapter_title(doc, "CHƯƠNG 2: PHÂN TÍCH HỆ THỐNG")

    # ============================================================
    # 2.1. Yêu cầu người dùng / Danh sách chức năng
    # ============================================================
    add_heading1(doc, "2.1. Yêu cầu người dùng / Danh sách chức năng")

    add_para(doc,
        "Phân tích hệ thống là giai đoạn chuyển hóa các yêu cầu khảo sát thực tế của doanh nghiệp "
        "thành các đặc tả chức năng phần mềm cụ thể. Dựa trên cơ cấu tổ chức và nhu cầu vận hành đa cấp "
        "tại các doanh nghiệp vừa và nhỏ, hệ thống xác định 5 nhóm tác nhân (Actors) chính tương tác với "
        "hệ thống, bao gồm: Admin (Quản trị hệ thống SaaS), Director (Giám đốc công ty), HR (Nhân sự), "
        "Manager (Quản lý phòng ban) và Employee (Nhân viên)."
    )

    add_para(doc,
        "Dưới đây là sơ đồ phân cấp Use Case tổng quan và bảng danh sách chi tiết các Use Case chức năng "
        "được bóc tách trực tiếp từ kiến trúc 24 controller và 45 bảng dữ liệu của hệ thống hiện tại:"
    )

    # Đường dẫn ảnh thực tế
    script_dir = os.path.dirname(os.path.abspath(__file__))
    img_path = os.path.join(script_dir, "use_case_diagram.png")

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        p_img.add_run().add_picture(img_path, width=Cm(15.5))
    else:
        run_img = p_img.add_run("[SƠ ĐỒ USE CASE TỔNG QUAN HỆ THỐNG - HÌNH ẢNH LỖI]")
        set_font(run_img, bold=True, color=RGBColor(128, 128, 128))
    
    p_img_sub = doc.add_paragraph()
    p_img_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf_img_sub = p_img_sub.paragraph_format
    pf_img_sub.space_before = Pt(6)
    pf_img_sub.space_after = Pt(12)
    run_img_sub = p_img_sub.add_run("Hình 1: Sơ đồ Use Case tổng quát hệ thống vận hành thông minh")
    set_font(run_img_sub, size=12, italic=True)

    # Bảng 2.1: Danh sách Use Case
    headers_uc = ["STT", "Mã UC", "Tên Use Case", "Tác nhân chính", "Mô tả tóm tắt chức năng"]
    rows_uc = [
        # Nhóm I: Admin (Hệ thống SaaS)
        ["1", "UC_AD_01", "Quản lý gói SaaS", "Admin hệ thống", "CRUD các gói dịch vụ (SaaSPackage), giá, giới hạn user và quyền AI."],
        ["2", "UC_AD_02", "Duyệt đăng ký mua", "Admin hệ thống", "Theo dõi và duyệt đăng ký mua gói của các công ty (PurchaseRegistration)."],
        ["3", "UC_AD_03", "Cấu hình tham số", "Admin hệ thống", "Cấu hình tham số hệ thống chung (tần suất check-in, số lượng mục tiêu tối đa)."],
        ["4", "UC_AD_04", "Xem nhật ký thao tác", "Admin hệ thống", "Tra cứu lịch sử hoạt động toàn bộ hệ thống qua Audit Logs."],
        
        # Nhóm II: Director (Ban Giám Đốc)
        ["5", "UC_DI_01", "Quản lý mục tiêu năm", "Director / Giám đốc", "Thiết lập sứ mệnh, tầm nhìn và mục tiêu tài chính hàng năm (MissionVisions)."],
        ["6", "UC_DI_02", "Thiết lập OKR công ty", "Director / Giám đốc", "Tạo OKR cấp công ty và gán các kết quả then chốt (OKRKeyResults)."],
        ["7", "UC_DI_03", "Phân bổ OKR phòng ban", "Director / Giám đốc", "Phân bổ chỉ tiêu OKR xuống các phòng ban, thiết lập mối liên kết đa cấp."],
        ["8", "UC_DI_04", "Phê duyệt kết quả kỳ", "Director / Giám đốc", "Duyệt bảng đánh giá hiệu suất cuối kỳ của toàn công ty (EvaluationResults)."],
        ["9", "UC_DI_05", "Dự báo ngân sách thưởng", "Director / Giám đốc", "Xem dự toán quỹ thưởng realtime dựa trên điểm số KPI thực tế cuối kỳ."],
        ["10", "UC_DI_06", "Smart Alerts AI", "Director / Giám đốc", "Nhận các cảnh báo rủi ro về tiến độ và vận hành do AI phân tích."],
        
        # Nhóm III: HR (Nhân sự)
        ["11", "UC_HR_01", "Quản lý cơ cấu phòng ban", "HR / Nhân sự", "CRUD sơ đồ cây phòng ban phân cấp (Departments) và gán quản lý."],
        ["12", "UC_HR_02", "Quản lý chức danh", "HR / Nhân sự", "CRUD danh mục chức vụ (Positions) và RankLevel phục vụ tính lương thưởng."],
        ["13", "UC_HR_03", "Quản lý hồ sơ nhân viên", "HR / Nhân sự", "CRUD thông tin nhân sự, import từ file Excel, auto-gen mã nhân viên."],
        ["14", "UC_HR_04", "Quản lý kỳ đánh giá", "HR / Nhân sự", "Tạo mới và đóng/mở các kỳ đánh giá hiệu suất (EvaluationPeriods) theo Quý/Năm."],
        ["15", "UC_HR_05", "Thiết lập công thức thưởng", "HR / Nhân sự", "Cấu hình GradingRanks (7 bậc S->D) và quy tắc tính thưởng tương ứng."],
        ["16", "UC_HR_06", "Xuất báo cáo tổng hợp", "HR / Nhân sự", "Tổng hợp và export bảng đánh giá, bảng lương thưởng ra Excel (EPPlus)."],
        
        # Nhóm IV: Manager (Trưởng phòng)
        ["17", "UC_MA_01", "Thiết lập OKR phòng ban", "Manager / Trưởng phòng", "Nhận OKR công ty phân rã và thiết lập OKR cấp phòng ban."],
        ["18", "UC_MA_02", "Giao KPI nhân viên", "Manager / Trưởng phòng", "Tạo KPI, gán trọng số, target và giao cho nhân viên thuộc phòng ban."],
        ["19", "UC_MA_03", "Phê duyệt Check-in", "Manager / Trưởng phòng", "Duyệt/Từ chối báo cáo tiến độ check-in của nhân viên trong Review Queue."],
        ["20", "UC_MA_04", "Quản lý dự án & Kanban", "Manager / Trưởng phòng", "Tạo dự án phòng ban, giao công việc (WorkItems) và theo dõi trên Kanban."],
        ["21", "UC_MA_05", "Lên lịch họp 1-on-1", "Manager / Trưởng phòng", "Lên lịch và lưu nội dung cuộc họp riêng định kỳ với nhân viên dưới quyền."],
        ["22", "UC_MA_06", "Đánh giá nhân viên kỳ", "Manager / Trưởng phòng", "Thực hiện chấm điểm và viết nhận xét hiệu suất nhân viên cuối kỳ."],
        
        # Nhóm V: Employee (Nhân viên)
        ["23", "UC_EM_01", "Xem mục tiêu & KPI", "Employee / Nhân viên", "Tra cứu OKR cá nhân/phòng ban và danh sách KPI được giao trong kỳ."],
        ["24", "UC_EM_02", "Check-in tiến độ", "Employee / Nhân viên", "Báo cáo giá trị thực hiện KPI theo lịch check-in, gửi ghi chú giải trình."],
        ["25", "UC_EM_03", "Quản lý Kanban cá nhân", "Employee / Nhân viên", "Xem công việc được phân công, cập nhật trạng thái thẻ kéo thả Kanban."],
        ["26", "UC_EM_04", "Tương tác 1-on-1", "Employee / Nhân viên", "Xác nhận lịch họp, gửi phản hồi và trao đổi ý kiến với quản lý trực tiếp."],
        ["27", "UC_EM_05", "Xem dự toán thưởng", "Employee / Nhân viên", "Theo dõi điểm hiệu suất cá nhân và mức thưởng dự toán tương ứng thực tế."],
        ["28", "UC_EM_06", "Trợ lý AI Gemini", "Employee / Nhân viên", "Chat với AI tư vấn lập kế hoạch, nhận gợi ý KPI thông minh cho công việc."],
    ]

    create_table(doc, headers_uc, rows_uc, col_widths=[1.0, 2.0, 3.5, 3.5, 6.8])
    add_table_caption(doc, "Bảng 9: Danh sách Use Case chức năng chi tiết của hệ thống")

    # Phân tích chi tiết yêu cầu người dùng
    add_heading2(doc, "2.1.2. Phân tích chi tiết yêu cầu người dùng (Yêu cầu phi chức năng)")
    
    add_para(doc,
        "Bên cạnh các yêu cầu chức năng kể trên, hệ thống được thiết kế để đáp ứng các tiêu chuẩn phi chức "
        "năng nghiêm ngặt của một hệ thống quản trị doanh nghiệp hiện đại:"
    )

    add_bullet(doc, "Hệ thống áp dụng cơ chế phân quyền Claims-based Auth kết hợp RBAC, đảm bảo tính bảo mật "
                   "dữ liệu giữa các phòng ban. Dữ liệu chỉ được truy cập theo phạm vi Access Scope (ví dụ: Manager chỉ "
                   "xem dữ liệu phòng ban của mình, Employee chỉ xem dữ liệu cá nhân). Các thông tin nhạy cảm như "
                   "mật khẩu được mã hóa SHA-256 trước khi lưu.", bold_prefix="Yêu cầu bảo mật")

    add_bullet(doc, "Giao diện được xây dựng responsive hoàn toàn bằng Bootstrap 5, tối ưu hóa hiển thị từ màn hình "
                   "desktop của quản trị viên đến màn hình điện thoại di động của nhân viên khi thực hiện check-in nhanh.", bold_prefix="Yêu cầu tương thích")

    add_bullet(doc, "Tích hợp cơ chế rate limit (15 req/phút) khi gọi tới Gemini API để tránh tràn băng thông và quá "
                   "tải chi phí. Tối ưu hóa truy vấn SQL Server bằng AsNoTracking đối với các truy vấn đọc, giúp hệ thống phản "
                   "hồi nhanh dưới 500ms đối với các tác vụ thông thường.", bold_prefix="Yêu cầu hiệu năng")

    add_bullet(doc, "Các tác vụ dọn dẹp dữ liệu lịch sử AI cũ được chạy ngầm tự động thông qua AIHistoryCleanupService "
                   "nhằm đảm bảo dung lượng lưu trữ tối ưu cho database dài hạn.", bold_prefix="Yêu cầu vận hành")

    # Kết luận
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Bằng việc xác định rõ 5 nhóm tác nhân và bóc tách thành 28 Use Case chi tiết, hệ thống đảm bảo "
        "bao phủ toàn bộ quy trình vận hành thực tế của doanh nghiệp vừa và nhỏ, tạo tiền đề vững chắc "
        "cho việc xây dựng các biểu đồ luồng hoạt động (Activity Diagrams) ở phần tiếp theo.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    add_page_break(doc)

    print("  ✓ CHƯƠNG 2: PHÂN TÍCH HỆ THỐNG")
    print("    ✓ 2.1. Yêu cầu người dùng / Danh sách chức năng (28 Use Cases)")
    write_chapter2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm Chương 2 (2.1) vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
