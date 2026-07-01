"""
Script tạo CHƯƠNG 1: GIỚI THIỆU - Phần 1.1 Bối cảnh - Hiện trạng
Dựa trên phân tích hệ thống thực tế và bố cục mẫu SD-38/SD-17/SD-29
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
TEN_DU_AN = "Hệ thống hỗ trợ vận hành thông minh cho doanh nghiệp vừa và nhỏ hỗ trợ quản lý đa cấp và đưa ra quyết định bằng AI"
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    """Tiêu đề chương (CHƯƠNG X: ...) - căn giữa, in đậm, viết hoa"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    run = p.add_run(text.upper())
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    """Đoạn văn bản - căn đều hai lề"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_para_mixed(doc, parts, indent=True, space_before=3, space_after=3):
    """Đoạn văn bản có nhiều run với format khác nhau.
    parts: list of (text, bold, italic)
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    """Dòng danh sách (bullet point)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def add_dash_bullet(doc, text, bold_prefix=""):
    """Dòng gạch đầu dòng"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.9)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("- " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("- " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    run = p.add_run(caption)
    set_font(run, size=12, italic=True)


# ===================== NỘI DUNG CHƯƠNG 1 =====================

def write_chapter1(doc):
    """CHƯƠNG 1: GIỚI THIỆU - 1.1. Bối cảnh - Hiện trạng"""

    add_chapter_title(doc, "CHƯƠNG 1: GIỚI THIỆU")

    # ============================================================
    # 1.1. BỐI CẢNH - HIỆN TRẠNG
    # ============================================================
    add_heading1(doc, "1.1. Bối cảnh - Hiện trạng")

    add_heading2(doc, "1.1.1. Nhu cầu thực tế và lý do chọn đề tài")

    add_para(doc,
        "Trong bối cảnh nền kinh tế số phát triển mạnh mẽ và cuộc Cách mạng Công nghiệp "
        "4.0 đang thay đổi sâu sắc phương thức vận hành doanh nghiệp, việc ứng dụng công "
        "nghệ thông tin vào quản trị hiệu suất đã trở thành yêu cầu tất yếu. Đặc biệt, đối "
        "với các doanh nghiệp vừa và nhỏ (SMEs) – lực lượng chiếm hơn 97% tổng số doanh "
        "nghiệp tại Việt Nam – áp lực cạnh tranh ngày càng lớn đòi hỏi phải có những công "
        "cụ quản trị hiện đại để tối ưu hóa nguồn lực và nâng cao năng suất lao động."
    )

    add_para(doc,
        "Tuy nhiên, thực tế cho thấy phần lớn các doanh nghiệp vừa và nhỏ hiện nay vẫn "
        "đang quản lý hiệu suất nhân viên và vận hành nội bộ theo phương thức truyền thống, "
        "mang tính thủ công cao. Quy trình thiết lập mục tiêu chiến lược, giao chỉ tiêu KPI "
        "(Key Performance Indicator) và theo dõi OKR (Objectives and Key Results) thường "
        "được thực hiện bằng bảng tính Excel rời rạc, email hoặc thậm chí ghi chép trên giấy. "
        "Điều này dẫn đến hàng loạt những khó khăn và bất cập nghiêm trọng trong quản lý "
        "doanh nghiệp:"
    )

    # Các vấn đề cốt lõi
    add_bullet(doc,
        "Dữ liệu KPI/OKR phân tán trên nhiều bảng tính, nhiều phòng ban, "
        "không có hệ thống tập trung để đồng bộ. Khi cần đánh giá hiệu suất tổng thể, "
        "quản lý phải thu thập thủ công từ nhiều nguồn khác nhau, dễ sai sót và tốn thời gian. "
        "Đặc biệt, mối liên kết giữa mục tiêu chiến lược cấp công ty với chỉ tiêu cấp phòng "
        "ban và cá nhân gần như không được theo dõi xuyên suốt.",
        bold_prefix="Sự rời rạc và thiếu đồng bộ trong dữ liệu hiệu suất"
    )

    add_bullet(doc,
        "Việc đánh giá nhân viên dựa chủ yếu vào cảm tính và nhận xét chủ quan "
        "của quản lý trực tiếp. Không có hệ thống xếp hạng tự động dựa trên dữ liệu "
        "thực tế (target vs. achieved), dẫn đến tình trạng đánh giá không công bằng, thiếu "
        "minh bạch, gây bất mãn trong đội ngũ nhân sự. Quy trình duyệt đánh giá đa cấp "
        "(Manager → Director) hầu như không tồn tại hoặc chỉ mang tính hình thức.",
        bold_prefix="Thiếu minh bạch và khách quan trong đánh giá hiệu suất"
    )

    add_bullet(doc,
        "Các doanh nghiệp thiếu công cụ phân tích dữ liệu hiệu suất theo thời "
        "gian thực. Khi một KPI có nguy cơ không đạt mục tiêu, không có cơ chế cảnh báo "
        "sớm (early warning) để quản lý can thiệp kịp thời. Việc ra quyết định về nhân sự, "
        "phân bổ nguồn lực, điều chỉnh chiến lược vẫn dựa trên \"linh cảm\" thay vì dữ liệu "
        "phân tích cụ thể.",
        bold_prefix="Không có công cụ hỗ trợ ra quyết định dựa trên dữ liệu"
    )

    add_bullet(doc,
        "Mục tiêu chiến lược hàng năm (Sứ mệnh, Tầm nhìn, Yearly Goals) "
        "được ban hành nhưng không có cơ chế phân rã xuống cấp phòng ban và cá nhân "
        "một cách có hệ thống. OKR cấp công ty không được liên kết với KPI cấp nhân viên, "
        "dẫn đến tình trạng \"mạnh ai nấy làm\", hoạt động thực thi lệch hướng so với "
        "mục tiêu chiến lược tổng thể.",
        bold_prefix="Khó khăn trong việc liên kết chiến lược với thực thi"
    )

    add_bullet(doc,
        "Nhân viên không chủ động được tiến độ thực hiện KPI vì thiếu giao diện "
        "check-in trực quan. Quản lý không thể theo dõi tiến độ check-in của nhân viên "
        "theo thời gian thực, dẫn đến tình trạng phát hiện vấn đề quá muộn khi kỳ đánh "
        "giá đã kết thúc. Việc nhắc nhở deadline check-in hoàn toàn thủ công qua email "
        "hoặc tin nhắn.",
        bold_prefix="Quy trình check-in và theo dõi tiến độ thủ công"
    )

    add_bullet(doc,
        "Việc tính toán thưởng theo hiệu suất, xuất báo cáo đánh giá tổng hợp "
        "theo phòng ban/kỳ đánh giá đều phải thực hiện thủ công trên Excel. Với doanh "
        "nghiệp có vài trăm nhân viên, khối lượng công việc này là rất lớn và dễ xảy ra "
        "sai sót, đặc biệt khi cần đối soát dữ liệu giữa nhiều phòng ban.",
        bold_prefix="Báo cáo và tính toán thưởng phức tạp, dễ sai sót"
    )

    add_bullet(doc,
        "Hầu hết các giải pháp quản lý KPI/OKR trên thị trường (15Five, "
        "Lattice, Culture Amp) đều hướng đến doanh nghiệp lớn với chi phí cao, giao diện "
        "phức tạp và không hỗ trợ tiếng Việt. Các doanh nghiệp SME tại Việt Nam thiếu "
        "một giải pháp phù hợp với quy mô, ngân sách và đặc thù văn hóa quản lý của mình.",
        bold_prefix="Thiếu giải pháp phù hợp cho doanh nghiệp Việt Nam"
    )

    # Kết luận phần bối cảnh
    add_para(doc, "", space_before=6, space_after=0, indent=False)

    add_para_mixed(doc, [
        ("Nhận thức rõ những thách thức trên, nhóm đã đề xuất và thực hiện đề tài ", False, False),
        (f"\"{TEN_DU_AN}\"", True, False),
        (". Đây là một giải pháp công nghệ toàn diện, được thiết kế đặc biệt cho doanh "
         "nghiệp vừa và nhỏ tại Việt Nam, nhằm:", False, False),
    ], space_before=6)

    add_dash_bullet(doc,
        "Số hóa toàn bộ quy trình quản lý hiệu suất từ thiết lập chiến lược "
        "(Sứ mệnh, Tầm nhìn, Mục tiêu chiến lược) → phân rã OKR đa cấp → giao KPI "
        "→ check-in theo dõi → đánh giá xếp hạng → tính thưởng tự động.",
        bold_prefix="Tự động hóa end-to-end"
    )

    add_dash_bullet(doc,
        "Tích hợp trí tuệ nhân tạo (AI Gemini) vào quy trình vận hành: chatbot "
        "tư vấn KPI/OKR, gợi ý KPI thông minh, phân tích hiệu suất tự động, cảnh báo "
        "rủi ro sớm – giúp quản lý đưa ra quyết định nhanh chóng, chính xác dựa trên dữ liệu.",
        bold_prefix="Hỗ trợ quyết định bằng AI"
    )

    add_dash_bullet(doc,
        "Phân quyền chi tiết theo vai trò (Admin, Director, Manager, HR, "
        "Employee) với 60 permissions, đảm bảo mỗi cấp quản lý chỉ truy cập dữ liệu "
        "trong phạm vi được ủy quyền, phù hợp với cơ cấu tổ chức phân cấp của doanh nghiệp.",
        bold_prefix="Quản lý đa cấp linh hoạt"
    )

    add_dash_bullet(doc,
        "Giao diện tiếng Việt, chi phí triển khai thấp, dễ sử dụng, phù "
        "hợp với quy mô và ngân sách của SME. Hỗ trợ mô hình SaaS cho phép quản lý "
        "đa công ty trên cùng một nền tảng.",
        bold_prefix="Phù hợp doanh nghiệp Việt Nam"
    )

    # ============================================================
    # 1.1.2. HIỆN TRẠNG QUY TRÌNH QUẢN LÝ HIỆN TẠI
    # ============================================================
    add_heading2(doc, "1.1.2. Hiện trạng quy trình quản lý hiện tại")

    add_para(doc,
        "Qua quá trình khảo sát và phân tích thực tế tại nhiều doanh nghiệp vừa và nhỏ "
        "tại Việt Nam, nhóm đã nhận diện các quy trình quản lý hiện tại còn nhiều bất cập "
        "cần được cải thiện. Dưới đây là bảng tổng hợp so sánh giữa quy trình hiện tại "
        "(thủ công) và giải pháp mà hệ thống cung cấp:"
    )

    # Bảng so sánh hiện trạng vs giải pháp
    headers = ["STT", "Quy trình", "Hiện trạng (Thủ công)", "Giải pháp hệ thống"]
    rows = [
        ["1", "Thiết lập chiến lược",
         "Sứ mệnh, tầm nhìn, mục tiêu chiến lược được lưu trên văn bản Word/PDF riêng lẻ, không kết nối với hoạt động thực thi",
         "Module MissionVision quản lý tập trung, liên kết trực tiếp với OKR và KPI"],
        ["2", "Quản lý OKR đa cấp",
         "OKR cấp công ty/phòng ban/cá nhân theo dõi trên Excel, không có cơ chế phân rã và liên kết giữa các cấp",
         "OKR 3 cấp (Công ty → Phòng ban → Cá nhân) với Key Results, tự động tính tiến độ"],
        ["3", "Giao và theo dõi KPI",
         "KPI giao bằng email/meeting, theo dõi bằng bảng tính, thiếu deadline và nhắc nhở tự động",
         "KPI workflow (Nháp → Duyệt → Thực hiện → Hoàn thành), giao theo phòng ban/cá nhân, nhắc nhở tự động"],
        ["4", "Check-in tiến độ",
         "Báo cáo tiến độ qua email hoặc họp định kỳ, không có dữ liệu lịch sử",
         "Check-in online với giá trị cụ thể, auto-calculate %, review queue cho Manager/Director"],
        ["5", "Đánh giá hiệu suất",
         "Đánh giá chủ quan, không có hệ thống xếp hạng tự động, thiếu dữ liệu so sánh",
         "So sánh Target vs Achieved, xếp hạng tự động 7 bậc (S→D), workflow duyệt đa cấp"],
        ["6", "Tính thưởng",
         "Tính toán thủ công trên Excel, dễ sai sót khi số lượng nhân viên lớn",
         "BonusRules tự động theo rank, dự toán thưởng realtime, export Excel"],
        ["7", "Báo cáo tổng hợp",
         "Tổng hợp thủ công từ nhiều nguồn, mất nhiều thời gian và công sức",
         "Dashboard trực quan (ApexCharts), báo cáo theo phòng ban/kỳ, export Excel tự động"],
        ["8", "Phân tích & dự báo",
         "Không có công cụ phân tích xu hướng, cảnh báo rủi ro",
         "AI Gemini phân tích hiệu suất, cảnh báo rủi ro sớm, gợi ý KPI thông minh"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=12, align=align)

    # Set widths
    widths = [Cm(1.2), Cm(3.0), Cm(5.5), Cm(6.0)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    add_table_caption(doc, "Bảng 5: So sánh hiện trạng và giải pháp hệ thống")

    # ============================================================
    # 1.1.3. TỔNG KẾT LÝ DO CHỌN ĐỀ TÀI
    # ============================================================
    add_heading2(doc, "1.1.3. Tổng kết lý do chọn đề tài")

    add_para(doc,
        "Dựa trên phân tích bối cảnh thực tế và hiện trạng quản lý tại các doanh nghiệp "
        "vừa và nhỏ, nhóm tổng kết các lý do cốt lõi cho việc lựa chọn đề tài này:"
    )

    add_bullet(doc,
        "Hơn 97% doanh nghiệp Việt Nam là SME, nhưng phần lớn vẫn quản lý "
        "hiệu suất thủ công bằng Excel. Nhu cầu số hóa quy trình quản lý KPI/OKR là "
        "rất lớn và cấp thiết.",
        bold_prefix="Nhu cầu thị trường lớn"
    )

    add_bullet(doc,
        "Hệ thống tích hợp AI Gemini để gợi ý KPI, phân tích hiệu suất, "
        "cảnh báo rủi ro – đây là xu hướng công nghệ tiên tiến, mang lại giá trị khác biệt "
        "so với các giải pháp truyền thống.",
        bold_prefix="Ứng dụng AI vào quản trị"
    )

    add_bullet(doc,
        "Từ thiết lập chiến lược → OKR đa cấp → KPI → Check-in → "
        "Đánh giá → Thưởng. Đây là quy trình quản lý hiệu suất end-to-end hiếm có ở "
        "các sản phẩm dành cho SME tại Việt Nam.",
        bold_prefix="Giải pháp toàn diện end-to-end"
    )

    add_bullet(doc,
        "Xây dựng trên ASP.NET 10 MVC + EF Core + SQL Server, "
        "đảm bảo hiệu năng, bảo mật và khả năng mở rộng. Phân quyền RBAC với 60 "
        "permissions, phù hợp với cơ cấu tổ chức đa cấp.",
        bold_prefix="Nền tảng công nghệ hiện đại"
    )

    add_bullet(doc,
        "Dự án là cơ hội để nhóm ứng dụng kiến thức về phân tích, thiết "
        "kế hệ thống, lập trình web, tích hợp AI và quản lý dự án theo phương pháp "
        "Agile vào một sản phẩm có tính ứng dụng thực tiễn cao.",
        bold_prefix="Giá trị học thuật"
    )

    # Kết luận chương
    add_para(doc, "", space_before=6, space_after=0, indent=False)

    add_para_mixed(doc, [
        ("Kết luận: ", True, False),
        (f"Dự án \"{TEN_DU_AN}\" không chỉ là một website quản lý KPI/OKR thông thường, "
         "mà là một ", False, False),
        ("hệ sinh thái quản trị hiệu suất doanh nghiệp toàn diện", True, False),
        (", được xây dựng dựa trên những \"nỗi đau\" thực tế và nhu cầu cấp thiết của "
         "các doanh nghiệp vừa và nhỏ tại Việt Nam trong thời đại chuyển đổi số. "
         "Hệ thống kết hợp quản lý đa cấp chặt chẽ với trí tuệ nhân tạo, giúp doanh nghiệp "
         "không chỉ theo dõi hiệu suất mà còn ra quyết định thông minh, nâng cao năng lực "
         "cạnh tranh bền vững.", False, False),
    ], space_before=6, space_after=12)


def write_section_1_2(doc):
    """1.2. Mục tiêu - Phạm vi"""

    add_heading1(doc, "1.2. Mục tiêu - Phạm vi")

    # ============================================================
    # 1.2.1. MỤC TIÊU CỦA ĐỀ TÀI
    # ============================================================
    add_heading2(doc, "1.2.1. Mục tiêu của đề tài")

    add_heading3(doc, "a) Mục tiêu tổng quát")

    add_para(doc,
        f"Mục tiêu chính của đề tài là xây dựng một hệ thống quản trị hiệu suất doanh "
        "nghiệp toàn diện, tích hợp trí tuệ nhân tạo (AI), phục vụ cho các doanh nghiệp "
        "vừa và nhỏ (SMEs) tại Việt Nam. Hệ thống cho phép số hóa toàn bộ quy trình "
        "quản lý hiệu suất từ thiết lập chiến lược, phân rã mục tiêu OKR đa cấp, giao và "
        "theo dõi KPI, check-in tiến độ định kỳ, đến đánh giá xếp hạng và tính thưởng "
        "tự động. Đặc biệt, hệ thống tích hợp AI Gemini để hỗ trợ quản lý đưa ra quyết "
        "định dựa trên dữ liệu phân tích, mang lại giá trị vượt trội so với các phương thức "
        "quản lý truyền thống."
    )

    add_heading3(doc, "b) Mục tiêu cụ thể")

    add_para(doc, "Đề tài hướng đến các mục tiêu cụ thể sau:", indent=True)

    objectives = [
        ("Quản lý chiến lược tập trung",
         "Cho phép Ban lãnh đạo thiết lập và quản lý Sứ mệnh (Mission), Tầm nhìn "
         "(Vision), Mục tiêu chiến lược hàng năm (Yearly Goals) với các chỉ tiêu tài chính "
         "cụ thể. Từ đó liên kết trực tiếp với hệ thống OKR và KPI, đảm bảo mọi hoạt "
         "động thực thi đều hướng về mục tiêu chung."),
        ("Quản lý OKR đa cấp",
         "Xây dựng hệ thống OKR 3 cấp (Công ty → Phòng ban → Cá nhân) với các Key "
         "Results đo lường được. Hỗ trợ liên kết OKR với Mission và phân bổ OKR theo "
         "phòng ban/nhân viên. Tự động tính tiến độ hoàn thành dựa trên giá trị "
         "mục tiêu/thực tế của Key Results, bao gồm cả chỉ số nghịch (IsInverse)."),
        ("Quản lý KPI và quy trình giao chỉ tiêu",
         "Thiết kế hệ thống KPI đa dạng với 3 loại (Định lượng, Định tính, Hành vi) và "
         "3 thuộc tính (Tăng trưởng, Ổn định, Giảm thiểu). Mỗi KPI có Target, ngưỡng "
         "Đạt/Không đạt (Pass/Fail Threshold), đơn vị đo, tần suất check-in và deadline "
         "cụ thể. Hỗ trợ giao KPI theo phòng ban hoặc cá nhân với trọng số (weight), "
         "liên kết với OKR và Key Result tương ứng. Workflow trạng thái rõ ràng: "
         "Bản nháp → Chờ duyệt → Đang thực hiện → Hoàn thành/Không đạt."),
        ("Check-in và theo dõi tiến độ thời gian thực",
         "Cho phép nhân viên báo cáo tiến độ KPI định kỳ (check-in) với giá trị cụ thể. "
         "Hệ thống tự động tính toán tiến độ phần trăm (%), giá trị kỳ vọng tại deadline, "
         "tiến độ theo lịch (schedule progress). Manager/Director có review queue để "
         "duyệt/từ chối check-in, chấm điểm và nhận xét. Tích hợp nhắc nhở tự động "
         "khi deadline sắp đến hoặc quá hạn."),
        ("Đánh giá hiệu suất và tính thưởng tự động",
         "So sánh Target vs Achieved cho từng KPI, tự động xếp hạng theo 7 bậc "
         "(S/A+/A/B+/B/C/D) với ngưỡng điểm cấu hình được. Quy trình đánh giá "
         "đa cấp (Draft → Submitted → Director Reviewed). Tính thưởng tự động "
         "theo BonusRules (% lương + cố định), dự toán thưởng realtime theo kỳ "
         "đánh giá. Export báo cáo Excel (EPPlus) theo phòng ban."),
        ("Tích hợp AI hỗ trợ quyết định",
         "Tích hợp Google Gemini AI với các tính năng: Chatbot tư vấn KPI/OKR "
         "context-aware (dựa trên dữ liệu thực); Gợi ý KPI thông minh theo OKR, "
         "phòng ban, nhân viên; Phân tích hiệu suất theo kỳ/phòng ban/cá nhân; "
         "Phân khúc khách hàng cho Sales; Cảnh báo rủi ro tự động (Smart Alerts); "
         "Hỗ trợ viết nhận xét đánh giá (AI Review). Lưu trữ lịch sử AI với "
         "auto-cleanup."),
        ("Quản lý dự án và công việc (Kanban)",
         "Hỗ trợ quản lý dự án (WorkProject) liên kết với OKR, phân bổ theo phòng ban, "
         "theo dõi tiến độ tổng thể. Quản lý công việc (WorkItem) theo Kanban "
         "(Backlog → Todo → InProgress → Review → Done → Blocked), gắn với KPI và "
         "OKR Key Result, có độ ưu tiên và trọng số ảnh hưởng KPI."),
        ("Phân quyền đa cấp và bảo mật",
         "Hệ thống RBAC với 5 vai trò (Admin, Director, Manager, HR, Employee) và "
         "60 permissions chi tiết. Data Scope đảm bảo mỗi cấp chỉ truy cập dữ liệu "
         "trong phạm vi ủy quyền. Cookie Authentication + Google OAuth2, Anti-CSRF, "
         "OTP email cho quên mật khẩu, Audit Logs ghi nhận mọi thao tác."),
        ("Dashboard trực quan và báo cáo",
         "Dashboard tổng quan với biểu đồ ApexCharts (Line, Donut, Bar) hiển thị "
         "thống kê KPI/OKR/nhân viên theo thời gian thực. Tìm kiếm toàn cục "
         "across entities. Export báo cáo Excel tự động. SEO tối ưu với meta tags, "
         "sitemap.xml, robots.txt."),
        ("Hỗ trợ mô hình SaaS",
         "Cho phép quản lý đa công ty trên cùng nền tảng với các gói dịch vụ "
         "(SaaSPackage) phân biệt theo số người dùng, tính năng OKR nâng cao "
         "và AI Insight. Hỗ trợ thanh toán và đăng ký trực tuyến."),
    ]

    for i, (title, desc) in enumerate(objectives, 1):
        add_bullet(doc, desc, bold_prefix=f"Mục tiêu {i} – {title}")

    # ============================================================
    # 1.2.2. PHẠM VI CỦA ĐỀ TÀI
    # ============================================================
    add_heading2(doc, "1.2.2. Phạm vi của đề tài")

    add_para(doc,
        "Phạm vi của đề tài tập trung vào việc phát triển một ứng dụng web quản trị "
        "hiệu suất doanh nghiệp, xây dựng trên nền tảng ASP.NET 10 MVC + Entity "
        "Framework Core + SQL Server. Hệ thống phục vụ 5 nhóm đối tượng người dùng "
        "chính với phạm vi chức năng cụ thể như sau:"
    )

    # --- Phạm vi phía quản trị ---
    add_heading3(doc, "a) Phạm vi phía quản trị (Admin / Director / HR)")

    admin_features = [
        ("Quản lý tổ chức",
         "CRUD phòng ban (cây phân cấp), chức vụ (12 chức danh với RankLevel), "
         "nhân viên (import Excel, auto-gen mã EMP001), gán phòng ban/chức vụ. "
         "Quản lý vai trò và phân quyền 60 permissions chi tiết."),
        ("Quản lý chiến lược",
         "Thiết lập Sứ mệnh/Tầm nhìn/Mục tiêu chiến lược hàng năm. "
         "Quản lý OKR cấp công ty với Key Results, phân bổ OKR cho phòng ban."),
        ("Quản lý KPI toàn diện",
         "Tạo kỳ đánh giá (Quý/Năm), thiết lập KPI với đầy đủ tham số "
         "(Target, Threshold, đơn vị đo, tần suất, deadline). Giao KPI cho phòng ban/"
         "cá nhân. Duyệt workflow KPI (Nháp → Duyệt → Thực hiện)."),
        ("Đánh giá và thưởng",
         "Xem kết quả đánh giá toàn công ty, duyệt đánh giá (Director Reviewed). "
         "Cấu hình GradingRanks (7 bậc) và BonusRules. Xem dự toán thưởng realtime. "
         "Export báo cáo đánh giá tổng hợp (Excel)."),
        ("Dashboard và báo cáo",
         "Dashboard tổng quan với biểu đồ thống kê. Tìm kiếm toàn cục. "
         "Nhật ký hoạt động (Audit Logs). Quản lý tham số hệ thống."),
        ("AI quản trị",
         "Chatbot tư vấn chiến lược, gợi ý KPI thông minh, phân tích hiệu suất "
         "toàn công ty, phân khúc khách hàng, cảnh báo rủi ro Smart Alerts."),
        ("Quản lý SaaS",
         "Quản lý gói dịch vụ, người dùng đăng ký, thanh toán. "
         "Cấu hình tính năng theo gói (OKR nâng cao, AI Insight)."),
    ]

    for title, desc in admin_features:
        add_bullet(doc, desc, bold_prefix=title)

    # --- Phạm vi phía Manager ---
    add_heading3(doc, "b) Phạm vi phía quản lý trung gian (Manager)")

    manager_features = [
        ("Quản lý OKR phòng ban",
         "Tạo và quản lý OKR cấp phòng ban, liên kết với OKR cấp công ty. "
         "Phân bổ OKR cho nhân viên trong phòng ban. Theo dõi tiến độ Key Results."),
        ("Giao và theo dõi KPI",
         "Tạo KPI cho nhân viên phòng ban quản lý, thiết lập target và deadline. "
         "Theo dõi tiến độ check-in của nhân viên. Review queue để duyệt/từ chối "
         "check-in, chấm điểm và nhận xét."),
        ("Đánh giá nhân viên",
         "Thực hiện đánh giá hiệu suất nhân viên phòng ban (Draft → Submitted). "
         "Xem kết quả xếp hạng và so sánh Target vs Achieved."),
        ("Quản lý dự án",
         "Tạo và quản lý dự án phòng ban, phân công công việc (WorkItem) theo Kanban. "
         "Liên kết dự án với OKR, theo dõi tiến độ tổng thể."),
        ("Họp 1-on-1",
         "Lên lịch và quản lý cuộc họp riêng Manager-Employee. "
         "Ghi nhận nội dung trao đổi và kế hoạch hành động."),
        ("Bình luận và phản hồi",
         "Bình luận trên KPI/Check-in của nhân viên (GoalComments). "
         "Theo dõi thông báo nhắc nhở deadline và cảnh báo từ AI."),
    ]

    for title, desc in manager_features:
        add_bullet(doc, desc, bold_prefix=title)

    # --- Phạm vi phía nhân viên ---
    add_heading3(doc, "c) Phạm vi phía nhân viên (Employee)")

    employee_features = [
        ("Xem KPI/OKR được giao",
         "Xem danh sách KPI được giao, chi tiết target, deadline, trọng số. "
         "Xem OKR cá nhân/phòng ban/công ty và tiến độ Key Results."),
        ("Check-in tiến độ KPI",
         "Thực hiện check-in định kỳ với giá trị cụ thể. Theo dõi tiến độ phần trăm "
         "tự động tính toán. Xem lịch sử check-in và nhận xét từ quản lý."),
        ("Xem đánh giá cá nhân",
         "Xem kết quả đánh giá hiệu suất, xếp hạng, điểm số. "
         "Xem dự toán thưởng dựa trên hiệu suất thực tế."),
        ("Quản lý công việc",
         "Xem và cập nhật trạng thái công việc được giao trên Kanban board. "
         "Bình luận trên công việc (WorkItemComment)."),
        ("Thông báo và nhắc nhở",
         "Nhận thông báo deadline check-in sắp đến, cảnh báo quá hạn. "
         "Nhận cảnh báo AI về hiệu suất cá nhân."),
        ("AI cá nhân",
         "Sử dụng chatbot AI để được tư vấn cải thiện KPI, hỏi đáp về "
         "quy trình và mục tiêu. Xem gợi ý KPI từ AI."),
    ]

    for title, desc in employee_features:
        add_bullet(doc, desc, bold_prefix=title)

    # --- Phạm vi kỹ thuật ---
    add_heading3(doc, "d) Phạm vi kỹ thuật")

    add_para(doc, "Về mặt kỹ thuật, hệ thống được triển khai với các thành phần sau:")

    tech_scope = [
        ("Backend", "ASP.NET 10 MVC (.NET 10.0), Entity Framework Core 10, SQL Server 2019+"),
        ("Frontend", "Razor Views + Bootstrap 5 + Vanilla JavaScript, ApexCharts.js cho biểu đồ"),
        ("AI Engine", "Google Gemini API (gemini-2.5-flash) với rate limiting (15 req/min, 1500 req/day)"),
        ("Authentication", "Cookie-based Authentication + Google OAuth2 + OTP Email (quên mật khẩu)"),
        ("Email", "SMTP Gmail cho thông báo, nhắc nhở deadline, xác thực OTP"),
        ("Export", "EPPlus (NonCommercial License) cho xuất báo cáo Excel"),
        ("Database", "45 entities chia thành 7 nhóm: Foundation, Organization, OKR, KPI, Check-in, Evaluation, System"),
        ("Background Services", "AIHistoryCleanupService tự động dọn dẹp lịch sử AI cũ"),
    ]

    for title, desc in tech_scope:
        add_bullet(doc, desc, bold_prefix=title)

    # --- Ngoài phạm vi ---
    add_heading3(doc, "e) Những chức năng nằm ngoài phạm vi")

    add_para(doc,
        "Trong khuôn khổ đồ án tốt nghiệp với thời gian và nguồn lực hạn chế, "
        "một số chức năng sau nằm ngoài phạm vi triển khai của dự án hiện tại:"
    )

    out_of_scope = [
        "Ứng dụng di động (Mobile App) – hệ thống hiện chỉ hỗ trợ giao diện web responsive, "
        "chưa có ứng dụng native cho iOS/Android.",
        "Tích hợp thanh toán trực tuyến qua cổng thanh toán (VNPay, MoMo) – "
        "module SaaS hiện chỉ quản lý gói dịch vụ và đăng ký, chưa tích hợp cổng thanh toán thực tế.",
        "Tích hợp với hệ thống HRM/ERP bên ngoài (SAP, Oracle HCM, BambooHR) – "
        "hệ thống hoạt động độc lập, chưa hỗ trợ API mở cho tích hợp bên thứ ba.",
        "Multi-language (đa ngôn ngữ) – giao diện hiện chỉ hỗ trợ tiếng Việt.",
        "Real-time collaboration (cộng tác thời gian thực) – hệ thống chưa tích hợp "
        "WebSocket/SignalR cho tính năng cộng tác đồng thời nhiều người dùng.",
        "Tính năng phân tích nâng cao (BI Dashboard) – biểu đồ phân tích hiện "
        "ở mức cơ bản, chưa hỗ trợ drill-down, pivot table hoặc tùy chỉnh báo cáo nâng cao.",
    ]

    for item in out_of_scope:
        add_dash_bullet(doc, item)

    # Kết luận phạm vi
    add_para(doc, "", space_before=6, space_after=0, indent=False)

    add_para_mixed(doc, [
        ("Tóm lại, ", True, False),
        ("phạm vi đề tài tập trung vào việc xây dựng một hệ thống web hoàn chỉnh với "
         "đầy đủ chức năng quản lý hiệu suất end-to-end, phục vụ 5 nhóm đối tượng "
         "(Admin, Director, Manager, HR, Employee) trên cùng một nền tảng thống nhất. "
         "Hệ thống đảm bảo tính toàn vẹn dữ liệu, bảo mật phân quyền đa cấp và tích "
         "hợp AI để hỗ trợ ra quyết định – tất cả đều được thiết kế phù hợp với đặc thù "
         "quản lý của doanh nghiệp vừa và nhỏ tại Việt Nam.", False, False),
    ], space_before=6, space_after=12)


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    add_page_break(doc)

    print("  ✓ CHƯƠNG 1: GIỚI THIỆU")
    print("    ✓ 1.1. Bối cảnh - Hiện trạng")
    write_chapter1(doc)

    print("    ✓ 1.2. Mục tiêu - Phạm vi")
    print("      ✓ 1.2.1. Mục tiêu (Tổng quát + 10 mục tiêu cụ thể)")
    print("      ✓ 1.2.2. Phạm vi (Admin/Manager/Employee/Kỹ thuật/Ngoài phạm vi)")
    write_section_1_2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm Chương 1 (1.1 + 1.2) vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()

