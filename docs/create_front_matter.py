"""
Script tạo Phần 2: Phần đầu tài liệu - Báo cáo Dự án Tốt nghiệp
Bao gồm: Mục lục, Danh mục hình ảnh, Danh mục bảng biểu, Theo dõi phiên bản,
Quy ước tài liệu, Bảng chú giải thuật ngữ, Danh sách thành viên, GVHD,
Lời cảm ơn, Lời mở đầu, Tóm tắt nội dung dự án
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = 14
FONT_SIZE_TITLE = 20
FONT_SIZE_HEADING2 = 14
FONT_SIZE_FOOTER = 12

TEN_DU_AN = "Hệ thống hỗ trợ vận hành thông minh cho doanh nghiệp vừa và nhỏ hỗ trợ quản lý đa cấp và đưa ra quyết định bằng AI"
TEN_DU_AN_UPPER = "HỆ THỐNG HỖ TRỢ VẬN HÀNH THÔNG MINH CHO DOANH NGHIỆP VỪA VÀ NHỎ HỖ TRỢ QUẢN LÝ ĐA CẤP VÀ ĐƯA RA QUYẾT ĐỊNH BẰNG AI"
NHOM = "NEXTGEN"
GIANG_VIEN = "Phan Hoàng Khải"

SINH_VIEN = [
    ("Phạm Trần Anh Quân", "TB01758", ""),
    ("Phạm Trần An An", "TB01817", ""),
    ("Bùi Nguyễn Anh Như", "TB01785", ""),
    ("Trần Thanh Phong", "TB01649", ""),
    ("Nguyễn Thế Bảo", "TB01573", ""),
    ("Đoàn Quốc Khánh", "TB01544", ""),
    ("Vũ Hoàng Huy Nhật", "TB01605", ""),
]

# Header/Footer màu xanh nhạt
HEADER_BG_COLOR = "1F4E79"


# ===================== HELPER FUNCTIONS =====================

def set_font(run, size=FONT_SIZE_NORMAL, bold=False, italic=False, font_name=FONT_NAME, color=None):
    """Thiết lập font cho run"""
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    """Thêm ngắt trang"""
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)
    # Xóa paragraph format
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_section_title(doc, title, font_size=FONT_SIZE_TITLE):
    """Thêm tiêu đề phần (căn giữa, in hoa, đậm)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    run = p.add_run(title.upper())
    set_font(run, size=font_size, bold=True)
    return p


def add_paragraph_text(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_before=3, space_after=3, first_line_indent=True, font_size=FONT_SIZE_NORMAL):
    """Thêm đoạn văn bản"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(20)
    if first_line_indent:
        pf.first_line_indent = Cm(1.27)

    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, italic=italic)
    return p


def add_heading_text(doc, text, level=1, font_size=FONT_SIZE_NORMAL):
    """Thêm heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)

    run = p.add_run(text)
    set_font(run, size=font_size, bold=True)
    return p


def set_cell_shading(cell, color):
    """Đặt màu nền cho ô bảng"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, font_size=FONT_SIZE_NORMAL, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  color=None):
    """Thiết lập nội dung cho ô bảng"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    """Tạo bảng với header và dữ liệu"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header_text, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG_COLOR)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_text(cell, str(cell_text), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def add_table_caption(doc, caption):
    """Thêm chú thích bảng"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    run = p.add_run(caption)
    set_font(run, size=12, italic=True)


# ===================== CÁC PHẦN NỘI DUNG =====================

def add_toc_placeholder(doc):
    """Mục lục - placeholder (cần cập nhật bằng Word)"""
    add_section_title(doc, "MỤC LỤC")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    # Tạo TOC field code
    run = p.add_run()
    set_font(run, size=FONT_SIZE_NORMAL)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run("(Nhấn chuột phải → Update Field để cập nhật mục lục trong Word)")
    set_font(run4, size=12, italic=True, color=RGBColor(128, 128, 128))

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


def add_list_of_figures(doc):
    """Danh mục hình ảnh - placeholder"""
    add_section_title(doc, "DANH MỤC HÌNH ẢNH")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run()
    set_font(run, size=FONT_SIZE_NORMAL)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\h \\z \\c "Hình" '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run("(Nhấn chuột phải → Update Field để cập nhật danh mục hình ảnh trong Word)")
    set_font(run4, size=12, italic=True, color=RGBColor(128, 128, 128))

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


def add_list_of_tables(doc):
    """Danh mục bảng biểu - placeholder"""
    add_section_title(doc, "DANH MỤC BẢNG BIỂU")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run()
    set_font(run, size=FONT_SIZE_NORMAL)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\h \\z \\c "Bảng" '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run("(Nhấn chuột phải → Update Field để cập nhật danh mục bảng biểu trong Word)")
    set_font(run4, size=12, italic=True, color=RGBColor(128, 128, 128))

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


def add_version_tracking(doc):
    """Theo dõi phiên bản tài liệu"""
    add_section_title(doc, "THEO DÕI PHIÊN BẢN TÀI LIỆU")

    headers = ["TÊN", "NGÀY", "LÝ DO THAY ĐỔI", "PHIÊN BẢN"]
    rows = [
        [TEN_DU_AN, "01/07/2026", "Khởi tạo tài liệu", "1.0"],
    ]
    create_table(doc, headers, rows, col_widths=[5.5, 2.5, 5.0, 2.5])
    add_table_caption(doc, "Bảng 1: Theo dõi phiên bản tài liệu")


def add_document_conventions(doc):
    """Quy ước tài liệu"""
    add_section_title(doc, "QUY ƯỚC TÀI LIỆU")

    headers = ["Mục", "Thuộc tính", "Giá trị"]
    rows = [
        ["Font chữ", "", "Times New Roman"],
        ["Kích thước", "Tiêu đề", "Cỡ chữ 14"],
        ["", "Chữ thường", "Cỡ chữ 14"],
        ["", "Chân trang", "Cỡ chữ 12"],
        ["Tiêu đề", "Tiêu đề chính", "Cỡ chữ 14, In đậm, Viết hoa"],
        ["", "Tiêu đề cấp 2", "Cỡ chữ 14, In đậm"],
        ["", "Tiêu đề cấp 3", "Cỡ chữ 14, In đậm"],
        ["", "Tiêu đề thường", "Cỡ chữ 14, In thường"],
        ["Bảng", "Định dạng", "Tự khớp Window"],
        ["", "Căn chỉnh", "Chính giữa / Sang trái"],
        ["", "Tiêu đề", "In đậm, Nền xanh đậm, Chữ trắng"],
        ["", "Đường viền", "1.0pt, Màu đen"],
        ["Đoạn văn", "", "Căn đều hai lề, Khoảng cách dòng 1.5"],
    ]

    table = create_table(doc, headers, rows, col_widths=[3.5, 4.0, 8.0])

    # Merge cells cho cột "Mục" khi giá trị rỗng
    # (Cần xử lý bằng tay trong Word cho chính xác)

    add_table_caption(doc, "Bảng 2: Quy ước tài liệu")


def add_glossary(doc):
    """Bảng chú giải thuật ngữ"""
    add_section_title(doc, "BẢNG CHÚ GIẢI THUẬT NGỮ")

    headers = ["STT", "Thuật ngữ", "Giải thích"]
    rows = [
        ["1", "KPI", "Key Performance Indicator – Chỉ số đo lường hiệu suất công việc, giúp đánh giá mức độ hoàn thành mục tiêu."],
        ["2", "OKR", "Objectives and Key Results – Phương pháp quản lý mục tiêu bằng cách thiết lập mục tiêu (Objective) và các kết quả then chốt (Key Results) đo lường được."],
        ["3", "RBAC", "Role-Based Access Control – Kiểm soát truy cập dựa trên vai trò, phân quyền người dùng theo chức danh."],
        ["4", "MVC", "Model-View-Controller – Mô hình kiến trúc phần mềm phân tách ứng dụng thành 3 thành phần: Model (dữ liệu), View (giao diện), Controller (logic điều khiển)."],
        ["5", "EF Core", "Entity Framework Core – ORM (Object-Relational Mapping) của Microsoft, ánh xạ đối tượng C# sang bảng cơ sở dữ liệu."],
        ["6", "API", "Application Programming Interface – Giao diện lập trình ứng dụng, cho phép các hệ thống giao tiếp với nhau."],
        ["7", "CRUD", "Create, Read, Update, Delete – Bốn thao tác cơ bản trên dữ liệu."],
        ["8", "SQL Server", "Hệ quản trị cơ sở dữ liệu quan hệ của Microsoft, lưu trữ và quản lý dữ liệu hệ thống."],
        ["9", "AI / Gemini", "Trí tuệ nhân tạo (Artificial Intelligence). Gemini là mô hình AI của Google được tích hợp trong dự án để gợi ý KPI, phân tích hiệu suất."],
        ["10", "Dashboard", "Bảng điều khiển tổng quan, hiển thị dữ liệu trực quan bằng biểu đồ và chỉ số thống kê."],
        ["11", "Check-in", "Hành động cập nhật tiến độ thực hiện KPI/OKR định kỳ (hàng tuần/tháng)."],
        ["12", "ERD", "Entity Relationship Diagram – Sơ đồ quan hệ thực thể, mô tả cấu trúc và mối quan hệ giữa các bảng trong CSDL."],
        ["13", "Use Case", "Kịch bản sử dụng, mô tả sự tương tác giữa người dùng và hệ thống cho một chức năng cụ thể."],
        ["14", "Sprint", "Chu kỳ làm việc ngắn (1-4 tuần) trong phương pháp Agile Scrum."],
        ["15", "Leader", "Trưởng nhóm, người phụ trách quản lý tiến độ và điều phối công việc trong nhóm dự án."],
        ["16", "Developer", "Lập trình viên, người phát triển mã nguồn và xây dựng các chức năng phần mềm."],
        ["17", "Tester", "Kiểm thử viên, người kiểm tra chất lượng sản phẩm và phát hiện lỗi (bug)."],
        ["18", "SMTP", "Simple Mail Transfer Protocol – Giao thức gửi email, được sử dụng cho chức năng thông báo và nhắc nhở."],
        ["19", "SaaS", "Software as a Service – Mô hình phân phối phần mềm dưới dạng dịch vụ đám mây, người dùng trả phí theo gói."],
        ["20", "Agile Scrum", "Khung làm việc phát triển phần mềm linh hoạt, tập trung vào các chu kỳ lặp ngắn (Sprint) và cải tiến liên tục."],
    ]

    create_table(doc, headers, rows, col_widths=[1.5, 3.0, 11.0])
    add_table_caption(doc, "Bảng 3: Chú giải thuật ngữ")


def add_member_list(doc):
    """Danh sách thành viên"""
    add_section_title(doc, "DANH SÁCH THÀNH VIÊN")

    headers = ["STT", "HỌ VÀ TÊN", "MÃ SV", "EMAIL"]
    rows = []
    for i, (name, msv, email) in enumerate(SINH_VIEN):
        rows.append([str(i + 1), name, msv, email if email else ""])

    create_table(doc, headers, rows, col_widths=[1.5, 5.0, 3.0, 6.0])
    add_table_caption(doc, "Bảng 4: Danh sách thành viên nhóm " + NHOM)

    # Cam kết
    add_paragraph_text(doc, "", space_before=12, space_after=0, first_line_indent=False)
    add_paragraph_text(doc,
        "Nhóm dự án cam kết rằng tất cả thông tin, báo cáo, và nội dung được cung cấp "
        "trong tài liệu này là một sự biểu đạt trung thực của quan điểm và công việc của "
        "tất cả các thành viên. Nhóm xác nhận rằng không sao chép, không vi phạm bản "
        "quyền, hoặc không thể hiện các vấn đề liên quan đến việc sao chép từ nguồn bên "
        "ngoài. Trong trường hợp vi phạm bản quyền hoặc sự sao chép không cẩn thận, "
        "nhóm sẽ chịu trách nhiệm và đồng ý hủy bỏ kết quả của dự án này.",
        italic=True, space_before=6, space_after=6
    )


def add_gvhd_page(doc):
    """Giảng viên hướng dẫn"""
    add_section_title(doc, "GIẢNG VIÊN HƯỚNG DẪN")

    add_paragraph_text(doc, f"Họ và tên: {GIANG_VIEN}", bold=False, first_line_indent=False, space_before=12)
    add_paragraph_text(doc, "Cơ quan công tác: Trường Cao Đẳng FPT Polytechnic", first_line_indent=False)
    add_paragraph_text(doc, "Điện thoại: .......................................", first_line_indent=False)
    add_paragraph_text(doc, "Email: .......................................", first_line_indent=False)

    add_paragraph_text(doc, "", space_before=12, space_after=0, first_line_indent=False)
    add_paragraph_text(doc, "Ý kiến nhận xét, đánh giá của giảng viên hướng dẫn:", bold=True, first_line_indent=False)

    # Dòng chấm chờ nhận xét
    for _ in range(12):
        add_paragraph_text(doc,
            ".......................................................................................................................",
            first_line_indent=False, space_before=0, space_after=0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

    # Chữ ký
    add_paragraph_text(doc, "", space_before=24, space_after=0, first_line_indent=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.right_indent = Cm(2.0)
    run = p.add_run("Giảng viên hướng dẫn")
    set_font(run, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2f = p2.paragraph_format
    p2f.space_before = Pt(2)
    p2f.space_after = Pt(0)
    p2f.right_indent = Cm(2.0)
    run2 = p2.add_run("(Ký và ghi rõ họ tên)")
    set_font(run2, italic=True, size=12)

    # Khoảng trống cho chữ ký
    add_paragraph_text(doc, "", space_before=36, space_after=0, first_line_indent=False)
    add_paragraph_text(doc, "", space_before=12, space_after=0, first_line_indent=False)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3f = p3.paragraph_format
    p3f.right_indent = Cm(2.0)
    run3 = p3.add_run(GIANG_VIEN)
    set_font(run3, bold=True)


def add_loi_cam_on(doc):
    """Lời cảm ơn"""
    add_section_title(doc, "LỜI CẢM ƠN")

    add_paragraph_text(doc,
        f"Bản báo cáo dự án tốt nghiệp \"{TEN_DU_AN}\" là kết quả của một quá trình "
        "nghiên cứu nghiêm túc, ứng dụng những kiến thức công nghệ hiện đại và nỗ lực "
        "làm việc tập trung của cả nhóm. Dự án này không chỉ đánh dấu sự trưởng thành "
        "về mặt kỹ năng chuyên môn của chúng em mà còn là minh chứng cho sự hỗ trợ, "
        "đồng hành và định hướng quý báu từ phía nhà trường, thầy cô cùng gia đình."
    )

    add_paragraph_text(doc,
        "Chúng em xin bày tỏ lòng biết ơn sâu sắc nhất tới toàn thể quý Thầy/Cô thuộc "
        "Bộ môn Ứng dụng phần mềm tại trường Cao đẳng FPT Polytechnic, những người "
        "đã truyền dạy nền tảng tri thức vững chắc và tư duy lập trình thực chiến trong "
        "suốt thời gian qua."
    )

    add_paragraph_text(doc,
        f"Đặc biệt, chúng em xin gửi lời cảm ơn chân thành đến Thầy {GIANG_VIEN}, "
        "giảng viên hướng dẫn trực tiếp Dự án tốt nghiệp của nhóm. Thầy đã dành thời "
        "gian theo sát tiến độ, cung cấp những định hướng nghiệp vụ chuyên sâu và các "
        "góp ý kỹ thuật mang tính thực tế cao, giúp nhóm tháo gỡ những bài toán phức "
        "tạp về quản trị và vận hành hệ thống. Những chỉ dẫn rõ ràng và chuyên nghiệp "
        "của Thầy không chỉ giúp chất lượng dự án được nâng cao mà còn giúp chúng em "
        "hình thành phong cách làm việc của một kỹ sư phần mềm thực thụ."
    )

    add_paragraph_text(doc,
        "Bên cạnh đó, chúng em cũng xin tri ân gia đình và bạn bè đã luôn tạo điều kiện "
        "tốt nhất, đóng góp những ý kiến phản hồi quý giá để hệ thống ngày càng hoàn "
        "thiện hơn."
    )

    add_paragraph_text(doc,
        "Do được thực hiện trong một khoảng thời gian khá hạn hẹp với khối lượng nghiệp "
        "vụ đồ sộ, chúng em tự nhận thức rằng dự án vẫn còn những khía cạnh có thể phát "
        "triển và tối ưu thêm. Vì vậy, chúng em rất mong nhận được những nhận xét, đánh "
        "giá công tâm và quý báu từ phía quý Thầy/Cô trong Hội đồng bảo vệ. Những ý "
        "kiến đóng góp đó sẽ là hành trang vô giá giúp chúng em mở rộng tầm nhìn, hoàn "
        "thiện kỹ năng và vững vàng hơn trên con đường phát triển sự nghiệp sau này."
    )

    add_paragraph_text(doc,
        "Một lần nữa, chúng em xin kính chúc quý Thầy/Cô sức khỏe và tiếp tục gặt hái "
        "được nhiều thành công trong sứ mệnh truyền đạt kiến thức cho thế hệ tương lai.",
        space_after=12
    )


def add_loi_mo_dau(doc):
    """Lời mở đầu"""
    add_section_title(doc, "LỜI MỞ ĐẦU")

    add_paragraph_text(doc,
        "Trong bối cảnh nền kinh tế số phát triển mạnh mẽ và cuộc Cách mạng Công nghiệp "
        "4.0 đang len lỏi vào mọi lĩnh vực, việc chuyển đổi số trong quản trị doanh nghiệp "
        "đã trở thành yêu cầu tất yếu để nâng cao năng lực cạnh tranh. Đặc biệt, trong bối "
        "cảnh các doanh nghiệp ngày càng chú trọng đến hiệu suất làm việc và sự liên kết "
        "giữa mục tiêu chiến lược với hoạt động thực thi, nhu cầu về một hệ thống quản lý "
        "hiệu suất toàn diện, minh bạch và dễ sử dụng là vô cùng cấp thiết."
    )

    add_paragraph_text(doc,
        "Tuy nhiên, thực tế cho thấy nhiều doanh nghiệp vừa và nhỏ vẫn đang quản lý hiệu "
        "suất nhân viên theo phương thức truyền thống: theo dõi KPI bằng bảng tính Excel rời "
        "rạc, đánh giá hiệu suất dựa trên cảm tính, thiếu công cụ kết nối giữa mục tiêu cấp "
        "công ty với cấp phòng ban và cá nhân. Điều này dẫn đến nhiều hệ lụy như:"
    )

    # Danh sách vấn đề
    issues = [
        "Sự rời rạc trong dữ liệu hiệu suất: Dữ liệu KPI/OKR phân tán trên nhiều bảng "
        "tính, thiếu tính đồng bộ và khó theo dõi tiến độ tổng thể.",
        "Thiếu minh bạch trong đánh giá: Không có hệ thống xếp hạng tự động, quy trình "
        "đánh giá phụ thuộc nhiều vào đánh giá chủ quan của quản lý.",
        "Không có công cụ hỗ trợ ra quyết định: Thiếu phân tích dữ liệu, cảnh báo sớm "
        "khi KPI có nguy cơ không đạt mục tiêu.",
        "Khó khăn trong việc liên kết chiến lược: Mục tiêu cấp công ty không được phân "
        "rã và theo dõi xuyên suốt đến cấp phòng ban và cá nhân."
    ]

    for issue in issues:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(20)
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-0.63)
        run = p.add_run("- ")
        set_font(run, bold=True)
        run2 = p.add_run(issue)
        set_font(run2)

    add_paragraph_text(doc,
        f"Nhận thức rõ những thách thức đó, dự án \"{TEN_DU_AN}\" đã được đề xuất và "
        "thực hiện. Mục tiêu cốt lõi của dự án là xây dựng một nền tảng quản trị hiệu suất "
        "thông minh, đồng bộ hóa các hoạt động từ thiết lập chiến lược đến theo dõi và "
        "đánh giá kết quả:", space_before=6
    )

    solutions = [
        "Về phía Quản lý: Cung cấp bộ công cụ thiết lập mục tiêu chiến lược, giao KPI/OKR "
        "cho phòng ban và cá nhân, theo dõi tiến độ thời gian thực với biểu đồ trực quan và "
        "xếp hạng tự động (S/A+/A/B+/B/C/D).",
        "Về phía Nhân viên: Giao diện check-in KPI trực quan, theo dõi tiến độ cá nhân, "
        "nhận thông báo nhắc nhở deadline và nhận gợi ý từ AI để cải thiện hiệu suất.",
        "Về phía Hệ thống: Tích hợp AI Gemini để phân tích hiệu suất, cảnh báo rủi ro, "
        "gợi ý KPI thông minh, hỗ trợ ra quyết định dựa trên dữ liệu."
    ]

    for sol in solutions:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(20)
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-0.63)
        run = p.add_run("- ")
        set_font(run, bold=True)
        run2 = p.add_run(sol)
        set_font(run2)

    add_paragraph_text(doc,
        "Thông qua dự án này, nhóm mong muốn tạo ra một giải pháp công nghệ có tính ứng "
        "dụng thực tiễn cao, giúp các doanh nghiệp tối ưu hóa quy trình quản lý hiệu suất "
        "và tăng hiệu quả vận hành. Đồng thời, đây cũng là cơ hội để chúng em củng cố kiến "
        "thức về lập trình hệ thống, xử lý logic nghiệp vụ và rèn luyện tư duy giải quyết vấn "
        "đề trong môi trường thực tế.", space_before=6
    )

    add_paragraph_text(doc, "Tài liệu bao gồm 6 chương:", space_before=6)

    chapters = [
        "Chương 1: Giới thiệu về bối cảnh, mục tiêu, nguồn lực, kế hoạch dự án và khảo sát.",
        "Chương 2: Phân tích các thực thể, chức năng và sơ đồ quan hệ thực thể.",
        "Chương 3: Thiết kế cơ sở dữ liệu và sơ đồ giao diện.",
        "Chương 4: Thực thi mã nguồn, thư viện sử dụng và đặc tả chức năng.",
        "Chương 5: Kiểm thử với tiêu chí, chiến lược và kết quả kiểm thử.",
        "Chương 6: Hướng dẫn sử dụng hệ thống theo các vai trò người dùng."
    ]

    for ch in chapters:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(20)
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-0.63)
        run = p.add_run("- ")
        set_font(run, bold=True)
        run2 = p.add_run(ch)
        set_font(run2)

    add_paragraph_text(doc,
        "Từ đó, chúng tôi mong rằng người đọc sẽ hiểu được nội dung, kế hoạch, mục tiêu "
        "dự án cũng như cách thức vận hành và sử dụng hệ thống.",
        space_before=6, space_after=12
    )


def add_tom_tat(doc):
    """Tóm tắt nội dung dự án"""
    add_section_title(doc, "TÓM TẮT NỘI DUNG DỰ ÁN")

    add_paragraph_text(doc,
        f"Dự án \"{TEN_DU_AN}\" được hình thành nhằm giải quyết những bất cập trong việc "
        "quản lý hiệu suất doanh nghiệp theo phương thức truyền thống, nơi mà các chỉ tiêu "
        "KPI và mục tiêu OKR thường được theo dõi rời rạc trên bảng tính, thiếu tính liên kết "
        "và khó đánh giá toàn diện."
    )

    add_paragraph_text(doc,
        "Mục tiêu cốt lõi của hệ thống là xây dựng một nền tảng quản trị hiệu suất thông "
        "minh, toàn diện, tạo ra sự kết nối liền mạch giữa mục tiêu chiến lược cấp công ty, "
        "phòng ban và cá nhân."
    )

    add_paragraph_text(doc,
        "Về phía Quản lý (Admin/HR): Cung cấp bộ công cụ quản trị tập trung để thiết lập "
        "sứ mệnh, tầm nhìn, mục tiêu chiến lược hàng năm. Hệ thống hỗ trợ giao KPI cho "
        "phòng ban và cá nhân, theo dõi tiến độ check-in định kỳ, xếp hạng tự động "
        "(S/A+/A/B+/B/C/D) và tính thưởng dựa trên hiệu suất. Đặc biệt, tích hợp AI Gemini "
        "để gợi ý KPI phù hợp, phân tích xu hướng hiệu suất và cảnh báo rủi ro sớm."
    )

    add_paragraph_text(doc,
        "Về phía Nhân viên: Cung cấp giao diện check-in KPI trực quan, nhận thông báo "
        "deadline, xem tiến độ cá nhân và sử dụng chatbot AI để được tư vấn cải thiện hiệu "
        "suất. Nhân viên có thể tự theo dõi quá trình đánh giá của mình một cách minh bạch."
    )

    add_paragraph_text(doc,
        "Về mặt kỹ thuật, hệ thống được xây dựng trên nền tảng ASP.NET 10 MVC kết hợp "
        "Entity Framework Core và SQL Server, đảm bảo tính ổn định và khả năng mở rộng. "
        "Kiến trúc RBAC phân quyền chi tiết theo vai trò (Admin, Manager, Employee), hệ "
        "thống thông báo SMTP tự động nhắc nhở deadline check-in, và tích hợp mô hình SaaS "
        "cho phép quản lý đa công ty trên cùng một nền tảng."
    )

    add_paragraph_text(doc,
        "Tổng hòa các tính năng này, dự án không chỉ mang lại sự tiện lợi, minh bạch cho "
        "quá trình quản lý hiệu suất mà còn là bước tiến quan trọng trong việc hiện đại hóa "
        "hoạt động quản trị nhân sự của các doanh nghiệp.", space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    # Thêm page break sau trang bìa
    add_page_break(doc)

    print("  ✓ Mục lục")
    add_toc_placeholder(doc)
    add_page_break(doc)

    print("  ✓ Danh mục hình ảnh")
    add_list_of_figures(doc)
    add_page_break(doc)

    print("  ✓ Danh mục bảng biểu")
    add_list_of_tables(doc)
    add_page_break(doc)

    print("  ✓ Theo dõi phiên bản tài liệu")
    add_version_tracking(doc)
    add_page_break(doc)

    print("  ✓ Quy ước tài liệu")
    add_document_conventions(doc)
    add_page_break(doc)

    print("  ✓ Bảng chú giải thuật ngữ")
    add_glossary(doc)
    add_page_break(doc)

    print("  ✓ Danh sách thành viên")
    add_member_list(doc)
    add_page_break(doc)

    print("  ✓ Giảng viên hướng dẫn")
    add_gvhd_page(doc)
    add_page_break(doc)

    print("  ✓ Lời cảm ơn")
    add_loi_cam_on(doc)
    add_page_break(doc)

    print("  ✓ Lời mở đầu")
    add_loi_mo_dau(doc)
    add_page_break(doc)

    print("  ✓ Tóm tắt nội dung dự án")
    add_tom_tat(doc)

    # Lưu file
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã tạo xong Phần 2: Phần đầu tài liệu!")
    print(f"   📄 {OUTPUT_PATH}")
    print(f"\n📝 Lưu ý:")
    print(f"   - Mở file Word, nhấn Ctrl+A → F9 để cập nhật Mục lục/Danh mục")
    print(f"   - Bổ sung SĐT, Email thành viên và GVHD nếu cần")


if __name__ == '__main__':
    main()
