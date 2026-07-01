"""
Script tạo phần 4.2: Công nghệ và thư viện sử dụng cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0, 1]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 4.2 =====================

def write_section_4_2(doc):
    """4.2. Công nghệ và thư viện sử dụng"""

    add_heading1(doc, "4.2. Công nghệ và thư viện sử dụng")

    add_para(doc,
        "Để đảm bảo hệ thống vận hành ổn định, bảo mật cao và có khả năng mở rộng tốt, "
        "nhóm đã chọn lựa kỹ lưỡng hệ sinh thái công nghệ của Microsoft kết hợp với các thư viện mã nguồn mở uy tín. "
        "Dưới đây là bảng tổng hợp các công nghệ, framework và thư viện sử dụng trong hệ thống cùng mục đích cụ thể:"
    )

    # Bảng 28: Công nghệ sử dụng
    headers_tech = ["STT", "Công nghệ / Thư viện", "Phiên bản", "Mục đích sử dụng chính trong dự án"]
    rows_tech = [
        ["1", ".NET 10.0 / ASP.NET MVC", "10.0.0", "Môi trường thực thi & Web framework chính để xây dựng ứng dụng theo mô hình MVC, định tuyến router và render views."],
        ["2", "EF Core 10.0", "10.0.0", "Lớp ORM truy cập cơ sở dữ liệu SQL Server, thực thi các truy vấn LINQ và quản lý Migrations database tự động."],
        ["3", "SQL Server 2019+", "RDBMS", "Hệ quản trị cơ sở dữ liệu quan hệ lưu trữ 45 bảng dữ liệu thực thể của hệ thống."],
        ["4", "Google Gemini API", "gemini-2.5-flash", "Trí tuệ nhân tạo (Gemini Service) cung cấp gợi ý chỉ tiêu KPI, phân tích hiệu suất cuối kỳ và chatbot widget."],
        ["5", "EPPlus", "7.1.0", "Đọc và xuất các tệp báo cáo Excel, hỗ trợ import danh sách hàng loạt nhân viên và export bảng đánh giá, bảng lương thưởng HR."],
        ["6", "MailKit & MimeKit", "4.3.0", "Gửi email thông báo tự động (thông báo deadline check-in, xác thực mã OTP khi quên mật khẩu qua Gmail SMTP)."],
        ["7", "DotNetEnv", "3.0.0", "Đọc file cấu hình môi trường (.env) nhằm bảo mật thông tin nhạy cảm (Connection String, API Key, Client Secret)."],
        ["8", "Google OAuth 2.0", "2.0.0", "Tích hợp đăng nhập bên thứ ba, cho phép nhân viên đăng nhập hệ thống bằng tài khoản Google Workspace công ty."],
        ["9", "Bootstrap", "5.3.2", "Framework CSS chính xây dựng giao diện responsive tương thích hoàn toàn thiết bị di động, cung cấp modals, forms, grids."],
        ["10", "jQuery & AJAX", "3.7.1", "Tương tác DOM và gửi yêu cầu không đồng bộ AJAX lên AIController cho tính năng chat AI mượt mà, không tải lại trang."],
        ["11", "ApexCharts.js", "3.44.0", "Thư viện JavaScript để vẽ biểu đồ trực quan hóa dữ liệu hiệu suất (biểu đồ đường, cột, donut) trên dashboard."],
        ["12", "Cookie Auth Middleware", "Tích hợp sẵn", "Quản lý phiên đăng nhập (session), lưu trữ claims phân quyền người dùng phía client an toàn."],
    ]

    create_table(doc, headers_tech, rows_tech, col_widths=[1.0, 3.5, 2.0, 9.5])
    add_table_caption(doc, "Bảng 28: Danh sách các công nghệ và thư viện sử dụng trong hệ thống")

    # Chi tiết một số công nghệ cốt lõi
    add_heading2(doc, "4.2.2. Phân tích vai trò của các công nghệ cốt lõi")
    
    add_para(doc, "Trong số các công nghệ sử dụng, có 4 thành phần đóng vai trò quyết định đến kiến trúc của dự án:")

    add_bullet(doc, "Mang lại hiệu năng thực thi vượt trội, cơ chế Dependency Injection "
                   "tích hợp sẵn rất mạnh mẽ. Sử dụng tính năng nâng cao của .NET 10 giúp tối ưu hóa bộ nhớ và "
                   "hỗ trợ chạy đa luồng không đồng bộ (async/await) mượt mà cho các dịch vụ background.", bold_prefix="ASP.NET Core 10.0")

    add_bullet(doc, "Cho phép giao tiếp với CSDL hoàn toàn bằng mã C# mà không cần viết SQL thủ công, "
                   "giúp giảm thiểu tối đa các lỗi cú pháp. EF Core 10 hỗ trợ cấu hình Fluent API chi tiết cho 45 thực thể, "
                   "tự động xử lý cascade delete và duy trì tính toàn vẹn dữ liệu.", bold_prefix="Entity Framework Core 10.0")

    add_bullet(doc, "Mô hình ngôn ngữ lớn tốc độ cao và chi phí thấp của Google, "
                   "hỗ trợ context window lớn. Việc sử dụng Gemini API giúp hệ thống tự động hóa các khâu lập kế hoạch "
                   "và phân tích – điều mà các phần mềm truyền thống không thể thực hiện được.", bold_prefix="Google Gemini 2.5 Flash")

    add_bullet(doc, "Thư viện Excel tốt nhất cho .NET, hoạt động hiệu năng cao "
                   "mà không cần cài đặt Microsoft Office trên máy chủ. EPPlus giúp HR xuất nhanh hàng trăm dòng dữ liệu "
                   "đánh giá phòng ban ra tệp Excel định dạng chuẩn trong thời gian dưới 1 giây.", bold_prefix="EPPlus Library")

    # Kết luận mục
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Việc kết hợp đồng bộ giữa một backend .NET mạnh mẽ, một DB SQL Server tin cậy, giao diện Bootstrap "
        "trực quan và trí tuệ nhân tạo Gemini API đã cung cấp một bệ đỡ công nghệ vững chắc, đảm bảo "
        "hệ thống đáp ứng tốt các yêu cầu phi chức năng về hiệu năng và bảo mật.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 4.2. Công nghệ và thư viện sử dụng")
    print("      ✓ 4.2.1. Danh sách công nghệ (Bảng 28)")
    print("      ✓ 4.2.2. Phân tích công nghệ cốt lõi")
    write_section_4_2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 4.2 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
