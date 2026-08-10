"""
Script tạo phần 4.3: Đặc tả chức năng / Sequence diagram cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 4.3 =====================

def write_section_4_3(doc):
    """4.3. Đặc tả chức năng / Sequence diagram"""

    add_heading1(doc, "4.3. Đặc tả chức năng / Sequence diagram")

    add_para(doc,
        "Đặc tả chức năng đi sâu vào mô tả chi tiết logic xử lý nghiệp vụ ở lớp backend, cấu trúc dữ liệu "
        "đầu vào/đầu ra (Input/Output) và các quy tắc ràng buộc nghiệp vụ chính. Để minh họa trực quan "
        "quy trình giao tiếp liên lớp, mục này trình bày sơ đồ tuần tự (Sequence Diagram) cho quy trình "
        "yêu cầu gợi ý chỉ tiêu KPI thông minh qua model gateway có kiểm soát nguồn."
    )

    # ============================================================
    # 4.3.1. SƠ ĐỒ TUẦN TỰ QUY TRÌNH AI SUGGESTION
    # ============================================================
    add_heading2(doc, "4.3.1. Sơ đồ tuần tự quy trình gợi ý KPI có nguồn")
    
    add_para(doc,
        "Sơ đồ dưới đây mô tả luồng tương tác không đồng bộ (Asynchronous call) từ trình duyệt web của người dùng, "
        "đi qua lớp bảo mật phân quyền của Controller, nạp dữ liệu ngữ cảnh thực tế từ Database thông qua DbContext, "
        "gửi snapshot tối thiểu qua model gateway, kiểm tra strict JSON/citation, dựng lại source fingerprint rồi mới trả bản nháp cho người dùng:"
    )

    script_dir = os.path.dirname(os.path.abspath(__file__))
    seq_path = os.path.join(script_dir, "sequence_ai.png")

    p_seq = doc.add_paragraph()
    p_seq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(seq_path):
        p_seq.add_run().add_picture(seq_path, width=Cm(15.0))
    else:
        p_seq.add_run("[SƠ ĐỒ SEQUENCE AI - LỖI HÌNH ẢNH]")
    add_figure_caption(doc, "Hình 13: Sơ đồ tuần tự quy trình yêu cầu gợi ý KPI qua model gateway")

    # ============================================================
    # 4.3.2. MÔ TẢ XỬ LÝ BACKEND (BACKEND LOGIC)
    # ============================================================
    add_heading2(doc, "4.3.2. Mô tả xử lý backend")

    add_para(doc,
        "Logic xử lý ở backend được phân rã thành các bước xử lý tuần tự trong 3 thành phần chính "
        "đăng ký Dependency Injection trong Program.cs:"
    )

    add_bullet(doc, "Action 'SuggestKPI' tiếp nhận các ID phạm vi, yêu cầu anti-forgery và "
                   "HasPermissionAttribute kiểm tra quyền 'KPIS_CREATE'. Nếu hợp lệ, action gọi KPI Suggestion Advisor; "
                   "mọi lỗi quyền, source conflict, timeout và schema sai được ánh xạ thành response an toàn.", bold_prefix="1. AIController.cs")

    add_bullet(doc, "Lớp xử lý nạp ngữ cảnh thực tế từ database. "
                   "Sử dụng Entity Framework Core 10 truy vấn các bảng: OKRs, KPIs, Employees. "
                   "Hệ thống biên dịch thành snapshot tối thiểu gồm kỳ đang mở, OKR/Key Result, loại KPI, KPI mẫu và chức danh; "
                   "không đưa tên, mã, email hoặc số điện thoại nhân viên vào context model.", bold_prefix="2. AIDataService.cs")

    add_bullet(doc, "Gọi IAIModelClient ở nhiệt độ 0 và chỉ nhận strict JSON 3-5 bản nháp hoặc danh sách rỗng. "
                   "Advisor kiểm tra source ID, đơn vị, chiều KPI và quan hệ target/pass/fail, sau đó dựng lại snapshot trong "
                   "transaction Serializable. Chỉ AgentRun/citation metadata được lưu; prompt và raw response không được lưu.", bold_prefix="3. KpiSuggestionAdvisor.cs")

    # ============================================================
    # 4.3.3. MÔ TẢ INPUT / OUTPUT (CẤU TRÚC DỮ LIỆU)
    # ============================================================
    add_heading2(doc, "4.3.3. Mô tả Input và Output của chức năng")
    
    add_para(doc, "Cấu trúc dữ liệu trao đổi giữa Client và API Backend được quy định chặt chẽ như sau:")

    # --- Input ---
    add_heading3(doc, "a) Dữ liệu đầu vào (Input - JSON Request)")
    add_para(doc, "Khi client gửi Ajax Request lên endpoint `/AI/SuggestKPI`, gói tin JSON bao gồm:")
    add_bullet(doc, "ID của kỳ đánh giá hiện tại để lọc OKR liên quan.", bold_prefix="PeriodId [int]")
    add_bullet(doc, "ID phòng ban của nhân viên để lọc OKR phòng ban.", bold_prefix="DepartmentId [int]")
    add_bullet(doc, "ID của nhân viên để lọc chức danh công việc.", bold_prefix="EmployeeId [int]")
    add_bullet(doc, "ID Objective và Key Result liên kết; Key Result bắt buộc phải thuộc đúng Objective.", bold_prefix="OkrId / OkrKeyResultId [int]")

    # --- Output ---
    add_heading3(doc, "b) Dữ liệu đầu ra (Output - JSON Response)")
    add_para(doc, "Sau khi xử lý thành công, backend phản hồi gói tin JSON có cấu trúc:")
    add_bullet(doc, "Giá trị true/false xác định cuộc gọi API thành công hay lỗi.", bold_prefix="Success [bool]")
    add_bullet(doc, "Danh sách bản nháp có name, targetValue, unit, passThreshold, failThreshold, isInverse, rationale và sourceIds.", bold_prefix="Suggestions [array]")
    add_bullet(doc, "Mã AgentRun và danh sách citation metadata tương ứng với sourceIds đã dùng.", bold_prefix="AgentRunId / Citations")
    add_bullet(doc, "Cảnh báo abstain hoặc thông báo an toàn nếu success = false.", bold_prefix="Warnings [array]")

    # ============================================================
    # 4.3.4. QUY TẮC NGHIỆP VỤ CHÍNH (BUSINESS RULES)
    # ============================================================
    add_heading2(doc, "4.3.4. Quy tắc nghiệp vụ chính (Business Rules)")

    add_para(doc, "Chức năng tích hợp AI được ràng buộc bởi các quy tắc nghiệp vụ nghiêm ngặt sau:")

    rules = [
        ("Ràng buộc Access Scope bảo mật dữ liệu",
         "Người dùng chỉ được yêu cầu gợi ý trong tenant/phạm vi được cấp quyền và phải có KPIS_CREATE. "
         "Backend xác minh lại employee, department, kỳ, OKR và Key Result trước lẫn sau model call; nguồn đổi thì từ chối bản nháp cũ."),
        
        ("Giới hạn input và retry hữu hạn",
         "Snapshot gửi model bị giới hạn 24.000 ký tự, response tối đa 30.000 ký tự và chỉ retry schema một lần. "
         "Request hủy/timeout được dừng và ánh xạ thành lỗi an toàn, không giữ request model chạy vô hạn."),
        
        ("Strict schema, citation và validator KPI",
         "Model chỉ được dùng source ID và đơn vị do server cấp. Server từ chối field thừa, nguồn giả, KPI trùng tên, "
         "số lượng ngoài 3-5 và ngưỡng trái chiều KPI; danh sách rỗng được coi là abstain hợp lệ."),
        
        ("AI đề xuất, con người quyết định",
         "Advisor không ghi KPIs/KPIDetails và không lưu prompt/raw output. Nút áp dụng chỉ điền form; thao tác POST tạo KPI "
         "vẫn chạy toàn bộ validator quyền, kỳ, OKR, ngưỡng và phân bổ. Endpoint refine không nguồn đã được loại bỏ."),
    ]

    for title, desc in rules:
        add_bullet(doc, desc, bold_prefix=title)

    # Kết luận mục
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Việc đặc tả chi tiết logic backend, cấu trúc dữ liệu trao đổi và các quy tắc nghiệp vụ trên giúp "
        "quy trình tích hợp AI qua model gateway có thể kiểm toán, đồng thời bảo vệ an toàn thông tin và kiểm soát tốt "
        "chi phí tài nguyên hệ thống trong suốt quá trình vận hành lâu dài.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 4.3. Đặc tả chức năng / Sequence diagram")
    print("      ✓ 4.3.1. Sơ đồ tuần tự AI (Hình 13)")
    print("      ✓ 4.3.2. Mô tả xử lý backend")
    print("      ✓ 4.3.3. Mô tả Input/Output dữ liệu")
    print("      ✓ 4.3.4. Ràng buộc nghiệp vụ AI (Rate Limit, Scope)")
    write_section_4_3(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 4.3 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
