"""
Script tạo phần 6.2: Hướng dẫn sử dụng cho vai trò Giám đốc (Director) cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG 6.2 =====================

def write_section_6_2(doc):
    """6.2. Hướng dẫn sử dụng cho vai trò Giám đốc (Director)"""

    add_heading1(doc, "6.2. Hướng dẫn sử dụng cho vai trò Giám đốc (Director)")

    add_para(doc,
        "Vai trò Giám đốc (Director) tập trung vào giám sát chiến lược vĩ mô của doanh nghiệp, "
        "phê duyệt các chỉ tiêu OKR/KPI cấp công ty, đánh giá hiệu suất cuối kỳ của các Trưởng phòng "
        "và sử dụng Trợ lý AI để phân tích hiệu suất tổng thể của doanh nghiệp."
    )

    # 6.2.1. Tổng quan các nhiệm vụ chính
    add_heading2(doc, "6.2.1. Tổng quan nhiệm vụ của Giám đốc")
    add_para(doc, "Giám đốc thực hiện các nghiệp vụ quản trị chiến lược sau:")
    add_bullet(doc, "Theo dõi các biểu đồ phân phối điểm KPI phòng ban, tỷ lệ hoàn thành mục tiêu chiến lược và chạy AI phân tích.", bold_prefix="Giám sát Dashboard tổng quan")
    add_bullet(doc, "Thiết lập Sứ mệnh, Tầm nhìn năm và xây dựng cây mục tiêu OKR cấp công ty.", bold_prefix="Định hướng chiến lược (OKRs)")
    add_bullet(doc, "Xem xét và phê duyệt hoặc từ chối các chỉ tiêu KPI do Trưởng phòng đề xuất cho phòng ban.", bold_prefix="Phê duyệt chỉ tiêu KPI")
    add_bullet(doc, "Duyệt chốt điểm số đánh giá hiệu suất cuối kỳ của toàn bộ nhân viên, đưa ra quyết định lương thưởng.", bold_prefix="Duyệt chốt Đánh giá & Tính thưởng")

    # 6.2.2. Hướng dẫn thao tác chi tiết
    add_heading2(doc, "6.2.2. Hướng dẫn các thao tác quản trị chính")

    # Thao tác 1
    add_heading3(doc, "a) Giám sát Dashboard & Gọi AI Phân tích (URL: /Dashboard)")
    add_para(doc,
        "Màn hình Dashboard cung cấp cho Giám đốc cái nhìn toàn cảnh về hiệu suất doanh nghiệp:\n"
        "1. Đăng nhập hệ thống với tài khoản Giám đốc. Hệ thống tự động hiển thị trang chủ Dashboard (URL: `/Dashboard`).\n"
        "2. Quan sát các biểu đồ: Biểu đồ đường thể hiện tiến độ check-in trung bình, biểu đồ donut thể hiện tỷ lệ trạng thái KPI. "
        "Theo dõi widget 'Dự toán quỹ thưởng thực tế' (Expected Bonus Fund) tăng giảm theo thời gian thực dựa trên tiến độ KPI.\n"
        "3. Sử dụng tính năng AI phân tích: Click nút 'AI Performance Analysis'. Hệ thống sẽ tự động tổng hợp toàn bộ context "
        "về tiến độ OKRs công ty, KPIs chậm hạn của các phòng ban gửi sang Gemini API. "
        "Sau khoảng 3 giây, widget AI sẽ hiển thị báo cáo đánh giá chuyên sâu bằng tiếng Việt kèm các gợi ý giải pháp cải thiện cụ thể.",
        italic=False
    )

    # Thao tác 2
    add_heading3(doc, "b) Thiết lập OKR Công ty & xem Sơ đồ phân cấp (URL: /OKRs)")
    add_para(doc,
        "Giám đốc trực tiếp xây dựng mục tiêu chiến lược và kết quả then chốt KR cho công ty:\n"
        "1. Truy cập mục 'Mục tiêu OKRs' trên thanh Menu (URL: `/OKRs`).\n"
        "2. Click 'Tạo mới OKR', chọn phân loại 'Cấp công ty', nhập Tên mục tiêu (Ví dụ: 'Mở rộng thị phần miền Nam Q2-2026'), "
        "chọn Kỳ đánh giá và click 'Lưu'.\n"
        "3. Click vào OKR vừa tạo, chọn 'Thêm Kết quả then chốt (Key Result)', nhập tên KR (Ví dụ: 'Đạt doanh số 10 tỷ đồng'), "
        "giá trị mục tiêu Target (10,000,000,000) và đơn vị tính (VND). Click 'Lưu'.\n"
        "4. Hệ thống sẽ tự động kết xuất sơ đồ cây OKR. Tiến độ của OKR công ty sẽ tự động cập nhật tăng dần khi các phòng ban "
        "hoàn thành các KPI liên kết trực thuộc.",
        italic=False
    )

    # Thao tác 3
    add_heading3(doc, "c) Phê duyệt chốt kết quả đánh giá cuối kỳ (URL: /EvaluationResults)")
    add_para(doc,
        "Cuối mỗi kỳ đánh giá, Giám đốc phê duyệt bảng xếp hạng hiệu suất và thưởng để bộ phận Kế toán chi trả:\n"
        "1. Truy cập mục 'Đánh giá hiệu suất' -> 'Danh sách chờ duyệt' (URL: `/EvaluationResults`).\n"
        "2. Hệ thống hiển thị danh sách bảng đánh giá do các Trưởng phòng gửi lên (trạng thái 'Submitted'). "
        "Giám đốc click 'Xem chi tiết' để kiểm tra điểm số KPI đạt được và nhận xét của Trưởng phòng đối với nhân viên.\n"
        "3. Nhập ý kiến chỉ đạo tại ô 'Director Review Comment'. Click 'Phê duyệt (Approve)' để chốt kết quả đánh giá. "
        "Trạng thái bản ghi chuyển sang 'Closed', hệ thống khóa dữ liệu và tự động gửi thông báo điểm số cùng xếp hạng "
        "cho nhân viên qua email và hệ thống.",
        italic=False
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 6.2. Vai trò Giám đốc (Director)")
    print("      ✓ 6.2.1. Tổng quan nhiệm vụ")
    print("      ✓ 6.2.2. Thao tác chính (Dashboard, OKRs, Phê duyệt đánh giá)")
    write_section_6_2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 6.2 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
