"""
Script tạo phần 6.3: Hướng dẫn sử dụng cho vai trò Quản lý (Manager/Trưởng phòng) cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG 6.3 =====================

def write_section_6_3(doc):
    """6.3. Hướng dẫn sử dụng cho vai trò Quản lý (Manager)"""

    add_heading1(doc, "6.3. Hướng dẫn sử dụng cho vai trò Quản lý (Trưởng phòng / Manager)")

    add_para(doc,
        "Vai trò Quản lý (Manager / Trưởng phòng) đóng vai trò trung gian kết nối giữa ban giám đốc "
        "và đội ngũ nhân viên. Quản lý chịu trách nhiệm phân rã mục tiêu phòng ban, giao chỉ tiêu KPI, "
        "phê duyệt check-in tiến độ định kỳ, vận hành dự án Kanban và đánh giá hiệu suất của nhân viên cuối kỳ."
    )

    # 6.3.1. Tổng quan các nhiệm vụ chính
    add_heading2(doc, "6.3.1. Tổng quan nhiệm vụ của Quản lý")
    add_para(doc, "Quản lý chịu trách nhiệm thực hiện các quy trình nghiệp vụ sau:")
    add_bullet(doc, "Giao chỉ tiêu KPI kèm target, deadline và phân bổ trọng số cho nhân viên trực thuộc (sử dụng AI gợi ý).", bold_prefix="Thiết lập & Giao KPI")
    add_bullet(doc, "Kiểm duyệt tiến độ báo cáo định kỳ của nhân viên tại Review Queue, ghi nhận phản hồi và đánh giá điểm check-in.", bold_prefix="Duyệt check-in tiến độ")
    add_bullet(doc, "Quản lý dự án của phòng ban thông qua bảng Kanban board, tạo và giao các WorkItems liên kết trực tiếp với chỉ tiêu KPI.", bold_prefix="Vận hành Kanban & Dự án")
    add_bullet(doc, "Lên lịch họp 1-on-1 định kỳ để lắng nghe rào cản và thảo luận định hướng hành động với nhân viên.", bold_prefix="Họp trao đổi 1-on-1")
    add_bullet(doc, "Đánh giá điểm số KPI trung bình, đề xuất xếp hạng hiệu suất cuối kỳ gửi lên Giám đốc phê duyệt.", bold_prefix="Đánh giá hiệu suất nhân viên")

    # 6.3.2. Hướng dẫn thao tác chi tiết
    add_heading2(doc, "6.3.2. Hướng dẫn các thao tác chính")

    # Thao tác 1
    add_heading3(doc, "a) Giao KPI & Gọi AI gợi ý (URL: /KPIs)")
    add_para(doc,
        "Quy trình giao chỉ tiêu KPI cho nhân viên kết hợp AI sinh chỉ tiêu thông minh:\n"
        "1. Đăng nhập vai trò Quản lý. Vào mục 'Quản lý KPIs' (URL: `/KPIs`).\n"
        "2. Click nút 'Giao KPI mới'. Chọn Kỳ đánh giá, chọn Nhân viên trực thuộc, liên kết OKR Key Result tương ứng.\n"
        "3. Sử dụng AI gợi ý: Click nút 'AI KPI Generator'. Hệ thống tự động phân tích mục tiêu của phòng ban "
        "và đề xuất 3 chỉ tiêu KPI định lượng phù hợp. Quản lý có thể chọn 'Áp dụng nhanh' để tự động điền thông tin.\n"
        "4. Điền TargetValue (Giá trị mục tiêu), Unit (Đơn vị tính), Deadline và Trọng số (Weight %). "
        "Click 'Gửi yêu cầu'. KPI sẽ ở trạng thái 'Chờ duyệt' gửi lên Giám đốc.",
        italic=False
    )

    # Thao tác 2
    add_heading3(doc, "b) Duyệt tiến độ check-in trong Review Queue (URL: /KPICheckIns)")
    add_para(doc,
        "Khi nhân viên gửi check-in báo cáo tiến độ, hệ thống sẽ đẩy vào Review Queue của Quản lý để phê duyệt:\n"
        "1. Truy cập mục 'Duyệt Check-in' (URL: `/KPICheckIns`).\n"
        "2. Hệ thống hiển thị danh sách các bản ghi check-in ở trạng thái 'Pending'. Click 'Xem chi tiết'.\n"
        "3. Đọc báo cáo của nhân viên bao gồm: Giá trị thực tế đạt được, giải trình khó khăn (Barriers) "
        "và cờ báo chậm tiến độ do hệ thống quét.\n"
        "4. Điền điểm đánh giá lần check-in này (ReviewScore - thang điểm 100), nhập nhận xét (ReviewComment). "
        "Click 'Phê duyệt' (Approve) hoặc 'Yêu cầu check-in lại' (Reject). "
        "Khi được duyệt, hệ thống tự động cộng dồn tiến độ của KPI và OKR liên quan.",
        italic=False
    )

    # Thao tác 3
    add_heading3(doc, "c) Quản lý dự án Kanban & Giao việc liên kết KPI (URL: /WorkProjects)")
    add_para(doc,
        "Quản lý công việc chi tiết của phòng ban và liên kết trực tiếp với mục tiêu hiệu suất:\n"
        "1. Truy cập mục 'Dự án & Kanban' (URL: `/WorkProjects`). Click chọn dự án cần quản lý.\n"
        "2. Hệ thống hiển thị bảng Kanban board với 4 cột trạng thái: Backlog, To Do, In Progress, Done. "
        "Để giao việc: Click 'Thêm thẻ công việc' tại cột To Do.\n"
        "3. Nhập tên đầu việc, người thực hiện (Assignee), và đặc biệt chọn KPI liên kết (KPIId) "
        "để công việc này đóng góp trực tiếp vào tiến độ KPI của nhân viên đó. Click 'Lưu'.\n"
        "4. Quản lý và nhân viên có thể kéo thả các thẻ công việc giữa các cột để cập nhật trạng thái thời gian thực.",
        italic=False
    )

    # Thao tác 4
    add_heading3(doc, "d) Đánh giá hiệu suất cuối kỳ cho nhân viên (URL: /EvaluationResults)")
    add_para(doc,
        "Cuối kỳ đánh giá, Quản lý lập bảng tổng hợp hiệu suất gửi lên ban Giám đốc:\n"
        "1. Truy cập mục 'Đánh giá cuối kỳ' (URL: `/EvaluationResults`). Chọn nhân viên cần đánh giá.\n"
        "2. Hệ thống tự động tính điểm KPI trung bình cộng (TotalScore) dựa trên kết quả duyệt check-in thực tế. "
        "Hệ thống đưa ra gợi ý xếp hạng bậc (S->D) dựa theo điểm số.\n"
        "3. Quản lý điền nhận xét tổng hợp hiệu suất trong kỳ (ReviewComment). Click 'Gửi đánh giá' (Submit to Director).\n"
        "4. Trạng thái bản ghi chuyển sang 'Submitted' và khóa tính năng sửa đổi của Quản lý để chờ Giám đốc duyệt chốt.",
        italic=False
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 6.3. Vai trò Quản lý (Manager)")
    print("      ✓ 6.3.1. Tổng quan nhiệm vụ")
    print("      ✓ 6.3.2. Thao tác chính (Giao KPI + AI, Duyệt check-in, Kanban, Đánh giá nhân viên)")
    write_section_6_3(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 6.3 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
