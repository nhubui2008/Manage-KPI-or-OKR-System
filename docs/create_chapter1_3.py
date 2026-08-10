"""
Script tạo phần 1.3: Nguồn lực - Kế hoạch cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=12, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 1.3 =====================

def write_section_1_3(doc):
    """1.3. Nguồn lực - Kế hoạch"""

    add_heading1(doc, "1.3. Nguồn lực - Kế hoạch")

    # ============================================================
    # 1.3.1. NGUỒN LỰC
    # ============================================================
    add_heading2(doc, "1.3.1. Nguồn lực")

    add_para(doc,
        "Để hiện thực hóa đề tài và đảm bảo tiến độ cũng như chất lượng của dự án tốt nghiệp, "
        "nhóm NEXTGEN đã phân bổ nguồn lực cụ thể bao gồm đội ngũ nhân sự và các tài nguyên "
        "công nghệ. Dưới đây là cơ cấu vai trò và nhiệm vụ chi tiết của các thành viên:"
    )

    # Bảng 1.2: Nguồn lực nhân sự
    headers_member = ["STT", "Họ và tên", "Mã SV", "Vai trò", "Nhiệm vụ chính trong dự án"]
    rows_member = [
        ["1", "Phạm Trần Anh Quân", "TB01758", "Leader, AI Specialist",
         "Quản lý tiến độ dự án, chịu trách nhiệm kiến trúc AI-native, model gateway và các rào chắn hỗ trợ quyết định."],
        ["2", "Bùi Nguyễn Anh Như", "TB01785", "Frontend Developer (KPI, OKR)",
         "Xây dựng và phát triển giao diện người dùng cho các module quản lý OKR, thiết lập KPI và dashboard tổng quan."],
        ["3", "Phạm Trần An An", "TB01817", "Backend Developer (KPI, OKR)",
         "Thiết kế cơ sở dữ liệu và xây dựng các API, nghiệp vụ xử lý dữ liệu cho module KPI và OKR đa cấp."],
        ["4", "Vũ Hoàng Huy Nhật", "TB01605", "Frontend Developer (Vận hành & Kanban)",
         "Phát triển giao diện check-in tiến độ, phòng họp 1-on-1, và bảng Kanban quản lý công việc (Order/Kanban)."],
        ["5", "Nguyễn Thế Bảo", "TB01573", "Backend Developer (Vận hành & Kanban)",
         "Lập trình logic backend duyệt check-in, tính lương thưởng tự động, và các API kéo thả công việc bảng Kanban."],
        ["6", "Trần Thanh Phong", "TB01649", "Fullstack Developer (Nền tảng & Hệ thống)",
         "Xây dựng các chức năng nền tảng: Auth & Google OAuth, phân quyền Roles, quản lý SaaS Tenants, Audit Logs và MailKit SMTP."],
        ["7", "Đoàn Quốc Khánh", "TB01544", "Quality Assurance & Tester",
         "Thiết lập kịch bản kiểm thử (Test Cases), thực hiện kiểm thử chức năng và hiệu năng, đảm bảo chất lượng hệ thống."],
    ]

    create_table(doc, headers_member, rows_member, col_widths=[1.0, 4.0, 2.2, 3.5, 6.0])
    add_table_caption(doc, "Bảng 6: Phân bổ nguồn lực nhân sự nhóm NEXTGEN")

    add_para(doc, "Bên cạnh nguồn lực nhân sự, nhóm cũng xác định các tài nguyên công nghệ cốt lõi phục vụ dự án:")
    add_bullet(doc, "Hệ điều hành Windows, IDE Visual Studio 2022, SQL Server Management Studio (SSMS), Git/GitHub.", bold_prefix="Công cụ phát triển")
    add_bullet(doc, "ASP.NET Core MVC (.NET 10.0), Entity Framework Core 10.0, SQL Server 2019+ làm DB.", bold_prefix="Công nghệ Backend")
    add_bullet(doc, "Razor Views, CSS (Bootstrap 5), JavaScript (Vanilla JS), ApexCharts.js cho Dashboard đồ họa.", bold_prefix="Công nghệ Frontend")
    add_bullet(doc, "Model gateway qua IAIModelClient cho các advisor strict-schema, citation và abstain.", bold_prefix="Tích hợp AI")
    add_bullet(doc, "EPPlus cho xuất báo cáo Excel, DotNetEnv quản lý cấu hình bảo mật, MailKit cho SMTP Email.", bold_prefix="Thư viện phụ trợ")

    # ============================================================
    # 1.3.2. KẾ HOẠCH
    # ============================================================
    add_heading2(doc, "1.3.2. Kế hoạch")

    add_para(doc,
        "Kế hoạch thực hiện dự án tốt nghiệp của nhóm NEXTGEN được chia thành các giai đoạn rõ ràng "
        "theo mô hình phát triển phần mềm linh hoạt (Agile Scrum), đảm bảo các mốc bàn giao "
        "bám sát lịch trình của nhà trường. Dưới đây là bảng timeline chi tiết các đầu việc chính:"
    )

    # Bảng 1.3: Kế hoạch thực hiện
    headers_plan = ["STT", "Tên công việc", "Ngày bắt đầu", "Ngày kết thúc", "Người thực hiện", "Hoàn thành"]
    rows_plan = [
        ["1", "Chọn đề tài & lập nhóm", "05/05/2026", "08/05/2026", "Cả nhóm", "100%"],
        ["2", "Khảo sát hiện trạng & Nghiên cứu bối cảnh", "09/05/2026", "13/05/2026", "Cả nhóm", "100%"],
        ["3", "Xác định mục tiêu & Phạm vi đề tài", "14/05/2026", "16/05/2026", "Cả nhóm", "100%"],
        ["4", "Phân tích yêu cầu nghiệp vụ (SRS)", "17/05/2026", "21/05/2026", "Anh Quân, Cả nhóm", "100%"],
        ["5", "Thiết kế cơ sở dữ liệu (ERD, Schema)", "22/05/2026", "25/05/2026", "An An, Thế Bảo", "100%"],
        ["6", "Thiết kế Use Case & Luồng xử lý nghiệp vụ", "22/05/2026", "25/05/2026", "Anh Quân, Cả nhóm", "100%"],
        ["7", "Thiết kế giao diện UI/UX (Figma mockup)", "24/05/2026", "28/05/2026", "Anh Như, Huy Nhật", "100%"],
        ["8", "Khởi tạo Project & Cấu hình bảo mật (RBAC)", "29/05/2026", "31/05/2026", "Anh Quân, Thanh Phong", "100%"],
        ["9", "Lập trình Module 1: Tổ chức & Phân quyền", "01/06/2026", "05/06/2026", "Anh Quân, Thanh Phong, An An", "100%"],
        ["10", "Lập trình Module 2: Thiết lập chiến lược & OKR", "06/06/2026", "10/06/2026", "Anh Như, An An, Thế Bảo", "100%"],
        ["11", "Lập trình Module 3: Giao KPI & Quản lý KPI", "11/06/2026", "15/06/2026", "Huy Nhật, An An, Thế Bảo", "100%"],
        ["12", "Lập trình Module 4: Check-in & Duyệt tiến độ", "16/06/2026", "20/06/2026", "Anh Như, Huy Nhật, An An", "100%"],
        ["13", "Lập trình Module 5: Tích hợp AI-native", "19/06/2026", "24/06/2026", "Anh Quân", "100%"],
        ["14", "Lập trình Module 6: Quản lý công việc (Order/Kanban)", "20/06/2026", "25/06/2026", "Huy Nhật, Thế Bảo", "100%"],
        ["15", "Tích hợp hệ thống & Tối ưu hóa Database", "25/06/2026", "27/06/2026", "Cả nhóm", "100%"],
        ["16", "Viết kịch bản test & Thực hiện kiểm thử (QA)", "26/06/2026", "28/06/2026", "Quốc Khánh (Tester)", "100%"],
        ["17", "Sửa lỗi (Bug fixing) & Cấu hình môi trường", "28/06/2026", "29/06/2026", "Developer, Cả nhóm", "100%"],
        ["18", "Viết báo cáo & Chuẩn bị Slide bảo vệ", "29/06/2026", "30/06/2026", "Cả nhóm", "100%"],
        ["19", "Hoàn thiện sản phẩm & Nghiệm thu", "30/06/2026", "01/07/2026", "Cả nhóm", "100%"],
    ]

    create_table(doc, headers_plan, rows_plan, col_widths=[1.0, 5.5, 2.2, 2.2, 3.8, 2.0])
    add_table_caption(doc, "Bảng 7: Kế hoạch và tiến độ thực hiện dự án tốt nghiệp")

    # Đoạn tổng kết chương 1
    add_para(doc, "", space_before=12, space_after=0, indent=False)
    add_para(doc,
        "Thông qua việc hoạch định nguồn lực chi tiết và thiết lập một kế hoạch thực hiện rõ ràng, "
        "nhóm NEXTGEN cam kết phân bổ thời gian hợp lý cho từng giai đoạn, từ phân tích thiết kế "
        "đến triển khai và kiểm thử sản phẩm. Kế hoạch này là cơ sở để nhóm theo dõi chặt chẽ tiến "
        "độ hàng tuần, giảm thiểu rủi ro chậm trễ và đảm bảo sản phẩm đạt chất lượng cao nhất khi "
        "đưa ra hội đồng bảo vệ.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 1.3. Nguồn lực - Kế hoạch")
    print("      ✓ 1.3.1. Nguồn lực (Bảng thành viên, vai trò, công nghệ)")
    print("      ✓ 1.3.2. Kế hoạch (Bảng timeline chi tiết)")
    write_section_1_3(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 1.3 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
