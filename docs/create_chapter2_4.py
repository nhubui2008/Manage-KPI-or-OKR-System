"""
Script tạo phần 2.4: ERD / Quan hệ thực thể cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 2.4 =====================

def write_section_2_4(doc):
    """2.4. ERD / Quan hệ thực thể"""

    add_heading1(doc, "2.4. ERD / Quan hệ thực thể")

    add_para(doc,
        "Thiết kế cơ sở dữ liệu quan hệ là nền tảng quyết định độ ổn định và hiệu năng của toàn bộ hệ thống. "
        "Với cấu trúc 45 bảng dữ liệu thực tế được tổ chức chặt chẽ trong DbContext của ASP.NET 10 MVC, "
        "dưới đây là đặc tả danh mục thực thể, sơ đồ quan hệ thực thể (ERD) và phân tích các quan hệ cốt lõi."
    )

    # ============================================================
    # 2.4.1. DANH SÁCH THỰC THỂ DỮ LIỆU
    # ============================================================
    add_heading2(doc, "2.4.1. Danh sách thực thể dữ liệu")
    
    add_para(doc,
        "Để quản trị toàn diện hiệu suất doanh nghiệp kết hợp AI và dự án Kanban, cơ sở dữ liệu của hệ thống "
        "được phân rã thành 7 phân nhóm chức năng chính như sau:"
    )

    # Bảng 15: Phân nhóm bảng
    headers_db = ["STT", "Nhóm thực thể", "Các bảng chính", "Mô tả chức năng lưu trữ"]
    rows_db = [
        ["1", "Foundation\n(Nền tảng)", "SystemParameters, Roles, Permissions, Role_Permissions, Statuses",
         "Lưu trữ các tham số vận hành hệ thống động và phân quyền chi tiết (60 permissions)."],
        ["2", "Organization\n(Tổ chức)", "Departments, Positions, SystemUsers, Employees, EmployeeAssignments, GradingRanks",
         "Quản lý cây phòng ban đa cấp, hồ sơ nhân sự và phân bổ chức vụ."],
        ["3", "Strategy & OKR\n(Chiến lược)", "MissionVisions, OKRTypes, OKRs, OKRKeyResults, OKR Mappings",
         "Lưu trữ Sứ mệnh, Tầm nhìn, OKR 3 cấp và các kết quả then chốt KR."],
        ["4", "KPI setup\n(Chỉ tiêu KPI)", "EvaluationPeriods, KPITypes, KPIProperties, KPIs, KPIDetails, KPI Assignments",
         "Quản lý thiết lập các kỳ đánh giá hiệu suất, giao KPI kèm target, deadline và trọng số."],
        ["5", "Check-in\n(Thực thi)", "CheckInStatuses, FailReasons, KPICheckIns, CheckInDetails, GoalComments, OneOnOneMeetings",
         "Ghi nhận lịch sử check-in tiến độ của nhân viên và queue phê duyệt của quản lý."],
        ["6", "Evaluation\n(Đánh giá)", "EvaluationResults, KPIAdjustmentHistories, BonusRules, RealtimeExpectedBonuses",
         "Tính điểm hiệu suất cuối kỳ, tự động xếp hạng bậc S->D và dự toán lương thưởng tương ứng."],
        ["7", "System & AI\n(Hệ thống)", "SystemAlerts, AuditLogs, AIGenerationHistories",
         "Ghi nhận cảnh báo tiến độ, nhật ký Audit Logs hoạt động và lưu trữ lịch sử chat của Gemini AI."],
    ]

    create_table(doc, headers_db, rows_db, col_widths=[1.0, 3.0, 5.5, 6.5])
    add_table_caption(doc, "Bảng 15: Phân nhóm các thực thể dữ liệu trong hệ thống")

    # ============================================================
    # 2.4.2. SƠ ĐỒ QUAN HỆ THỰC THỂ (ERD)
    # ============================================================
    add_heading2(doc, "2.4.2. Sơ đồ quan hệ thực thể tổng thể (Core ERD)")

    add_para(doc,
        "Để đảm bảo tính trực quan và tập trung vào luồng xử lý chính, sơ đồ dưới đây thể hiện mối quan hệ "
        "giữa các thực thể cốt lõi nhất của hệ thống, bao gồm nhân sự, phòng ban, OKR/KPI, check-in, đánh giá "
        "và các dự án công việc:"
    )

    script_dir = os.path.dirname(os.path.abspath(__file__))
    img_path = os.path.join(script_dir, "erd_diagram.png")

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        p_img.add_run().add_picture(img_path, width=Cm(15.5))
    else:
        p_img.add_run("[SƠ ĐỒ ERD CỐT LÕI - LỖI HÌNH ẢNH]")
    add_table_caption(doc, "Hình 5: Sơ đồ quan hệ thực thể cốt lõi (Core ERD) của hệ thống")

    # ============================================================
    # 2.4.3. GIẢI THÍCH CÁC MỐI QUAN HỆ CHÍNH
    # ============================================================
    add_heading2(doc, "2.4.3. Giải thích các mối quan hệ chính")

    add_para(doc, "Dựa trên sơ đồ ERD cốt lõi, các mối quan hệ thực thể được thiết lập và ràng buộc chặt chẽ như sau:")

    relations_desc = [
        ("Quan hệ SystemUser – Employee (1 – 1)",
         "Một tài khoản đăng nhập (SystemUser) liên kết duy nhất với một hồ sơ nhân sự (Employee) thông qua "
         "khoá ngoại SystemUserId. Mối quan hệ này giúp hệ thống xác định danh tính và tự động áp dụng phạm vi dữ liệu "
         "(Access Scope) của nhân viên khi đăng nhập thành công."),
        
        ("Quan hệ Department – Employee (1 – N)",
         "Một phòng ban (Department) có thể chứa nhiều nhân viên (Employee). Tuy nhiên, mỗi nhân viên tại một thời "
         "điểm chỉ thuộc một phòng ban chính thông qua DepartmentId. Bảng Department còn có quan hệ đệ quy (ParentId) "
         "để xây dựng cấu trúc cây phòng ban phân cấp đa tầng."),
        
        ("Quan hệ Position – Employee (1 – N)",
         "Một chức vụ (Position) có thể được gán cho nhiều nhân viên, nhưng một nhân viên chỉ giữ một chức vụ chính "
         "qua PositionId. Thuộc tính RankLevel của Position là cơ sở để hệ thống đối chiếu với bảng BonusRules để tính thưởng."),
        
        ("Quan hệ OKR – OKRKeyResult (1 – N)",
         "Một mục tiêu OKR (Objective) có thể phân rã thành nhiều kết quả then chốt (Key Results) qua OKRId. "
         "Tiến độ hoàn thành của OKR được tính toán tự động dựa trên trung bình cộng tiến độ của các Key Results trực thuộc."),
        
        ("Quan hệ OKRKeyResult – KPI (1 – N)",
         "Để liên kết giữa chiến lược và thực thi, một Key Result có thể được cụ thể hóa bằng nhiều chỉ tiêu KPI khác nhau. "
         "Mối quan hệ này giúp tiến độ KPI khi được duyệt sẽ tự động đồng bộ hóa ngược lên Key Result tương ứng."),
        
        ("Quan hệ KPI – KPIDetail (1 – 1)",
         "Một KPI liên kết duy nhất với một bản ghi KPIDetail qua KPIId để lưu trữ các tham số cấu hình nâng cao như TargetValue, "
         "PassThreshold và CheckInFrequencyDays, nhằm tối ưu hóa kích thước bảng KPIs chính."),
        
        ("Quan hệ KPI – KPICheckIn (1 – N)",
         "Một chỉ tiêu KPI sẽ có nhiều lần báo cáo check-in tiến độ định kỳ trong suốt kỳ đánh giá qua KPIId. Bảng KPICheckIn liên kết "
         "với bảng Employee qua EmployeeId để xác định người thực hiện báo cáo."),
        
        ("Quan hệ Employee – EvaluationResult (1 – N)",
         "Một nhân viên có nhiều kết quả đánh giá (EvaluationResult) qua các kỳ đánh giá khác nhau (PeriodId). Bảng này lưu trữ điểm "
         "hiệu suất tổng hợp, xếp hạng rank cuối cùng (S->D) và số tiền thưởng thực lĩnh."),
        
        ("Quan hệ OKR – WorkProject (1 – N)",
         "Một mục tiêu OKR có thể được thực thi thông qua nhiều dự án (WorkProject) liên kết qua LinkedOKRId, giúp tổ chức theo dõi "
         "tiến độ dự án đóng góp vào OKR vĩ mô."),
        
        ("Quan hệ WorkProject – WorkItem (1 – N)",
         "Một dự án chứa nhiều thẻ công việc (WorkItem) được quản lý qua WorkProjectId. WorkItem liên kết với Employee qua AssigneeId "
         "và liên kết trực tiếp với KPIId để ghi nhận kết quả đóng góp công việc trực tiếp vào chỉ tiêu KPI."),
    ]

    for title, desc in relations_desc:
        add_bullet(doc, desc, bold_prefix=title)

    # Kết luận
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Việc thiết lập cơ sở dữ liệu quan hệ chuẩn hóa (3NF) giúp hệ thống đảm bảo tính toàn vẹn dữ liệu, "
        "tránh dư thừa và tối ưu hóa tốc độ truy vấn. Đây là nền tảng kỹ thuật quan trọng để nhóm triển khai "
        "giai đoạn thiết kế giao diện và thực thi mã nguồn ở các chương tiếp theo.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 2.4. ERD / Quan hệ thực thể")
    print("      ✓ 2.4.1. Danh sách thực thể dữ liệu (Bảng 15)")
    print("      ✓ 2.4.2. Sơ đồ ERD cốt lõi (Hình 5)")
    print("      ✓ 2.4.3. Giải thích mối quan hệ chính")
    write_section_2_4(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 2.4 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
