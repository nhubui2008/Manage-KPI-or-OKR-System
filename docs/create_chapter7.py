"""
Script tạo CHƯƠNG 7: TỔNG KẾT VÀ ĐÁNH GIÁ cho báo cáo tốt nghiệp
Dựa trên cấu trúc chuẩn của báo cáo tốt nghiệp FPT Polytechnic
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    """Tiêu đề chương (CHƯƠNG X: ...) - căn giữa, in đậm, viết hoa (Heading 1)"""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG CHƯƠNG 7 =====================

def write_chapter7(doc):
    """CHƯƠNG 7: TỔNG KẾT VÀ ĐÁNH GIÁ"""

    add_page_break(doc)
    add_chapter_title(doc, "CHƯƠNG 7: TỔNG KẾT VÀ ĐÁNH GIÁ")

    # ============================================================
    # 7.1. Kết quả thực hiện dự án
    # ============================================================
    add_heading1(doc, "7.1. Kết quả thực hiện dự án")

    # 7.1.1. Thời gian phát triển
    add_heading2(doc, "7.1.1. Thời gian phát triển")
    add_para(doc,
        "Dự án được triển khai và hoàn thiện trong tổng thời gian 8 tuần (từ ngày 05/05/2026 đến 01/07/2026). "
        "Nhóm NEXTGEN đã bám sát chặt chẽ bảng kế hoạch thực hiện theo từng giai đoạn (chọn đề tài, khảo sát, "
        "phân tích yêu cầu, thiết kế cơ sở dữ liệu, lập trình mã nguồn, kiểm thử và đóng gói sản phẩm), "
        "đảm bảo hoàn thành đúng thời hạn nghiệm thu tốt nghiệp."
    )

    # 7.1.2. Mức độ hoàn thành
    add_heading2(doc, "7.1.2. Mức độ hoàn thành")
    add_para(doc,
        "Tính đến thời điểm hiện tại, dự án đã đạt mức độ hoàn thành 98% khối lượng công việc đề ra. "
        "Hệ thống đã chạy thử nghiệm thực tế ổn định, giải quyết triệt để bài toán quản lý hiệu suất "
        "OKR/KPI kết hợp AI của các doanh nghiệp vừa và nhỏ, sẵn sàng đưa vào vận hành thực tế."
    )

    # 7.1.3. Những chức năng đã hoàn thiện
    add_heading2(doc, "7.1.3. Những chức năng đã hoàn thiện")
    add_para(doc, "Hệ thống đã xây dựng và kiểm thử thành công các phân hệ chức năng cốt lõi:")
    add_bullet(doc, "Cổng đăng nhập an toàn, phân quyền động Claims-based với 60 permissions, phân cấp dữ liệu Access Scope và tích hợp đăng nhập Google OAuth2.", bold_prefix="Phân hệ Xác thực & Phân quyền")
    add_bullet(doc, "Xây dựng cây cơ cấu tổ chức phòng ban đa cấp, hồ sơ nhân sự (CRUD, import/export dữ liệu Excel).", bold_prefix="Phân hệ Tổ chức & Nhân sự")
    add_bullet(doc, "Thiết lập mục tiêu OKR công ty, phòng ban, cá nhân và kết quả then chốt KR; tự động tính toán tiến độ liên cấp.", bold_prefix="Phân hệ Chiến lược & OKRs")
    add_bullet(doc, "Giao chỉ tiêu KPI kèm target, deadline, trọng số; quy trình check-in tiến độ định kỳ, hàng đợi duyệt check-in Review Queue của quản lý.", bold_prefix="Phân hệ Vận hành & Check-in KPI")
    add_bullet(doc, "Bảng Kanban board kéo thả quản lý công việc chi tiết (WorkItems) liên kết trực tiếp với chỉ tiêu KPI.", bold_prefix="Phân hệ Dự án & Kanban")
    add_bullet(doc, "Lên lịch họp và ghi nhận biên bản họp 1-on-1 giữa quản lý và nhân viên.", bold_prefix="Phân hệ Trao đổi & Họp 1-on-1")
    add_bullet(doc, "HR đóng kỳ đánh giá, hệ thống tự động tính điểm KPI trung bình, xếp hạng Rank S->D và dự toán thưởng dựa trên kết quả đạt được thực tế.", bold_prefix="Phân hệ Đánh giá & Thưởng")
    add_bullet(doc, "Widget tư vấn có nguồn, bản nháp KPI/KR và phân tích hiệu suất tham khảo; dữ liệu chính thức vẫn do con người xác nhận.", bold_prefix="Phân hệ AI-native có kiểm soát")

    # 7.1.4. Những hạn chế / Chức năng chưa hoàn thiện
    add_heading2(doc, "7.1.4. Những chức năng chưa hoàn thiện và hạn chế của hệ thống")
    add_para(doc, "Mặc dù hệ thống đã đáp ứng hầu hết các yêu cầu nghiệp vụ, dự án vẫn còn một số hạn chế cần cải thiện:")
    add_bullet(doc, "Hệ thống chưa tích hợp sâu với các hệ thống ERP lớn có sẵn của doanh nghiệp (như SAP, Oracle, Odoo) để tự động lấy dữ liệu doanh số thực tế, hiện tại vẫn dựa vào việc nhân viên tự nhập check-in thủ công.", bold_prefix="Tự động hóa kết nối dữ liệu ngoại vi")
    add_bullet(doc, "Hệ thống thanh toán của gói SaaS hiện tại mới dừng ở mức mô phỏng giao dịch (Mock Transactions), chưa liên kết trực tiếp với các cổng thanh toán ngân hàng (VNPAY, PayOS) để tự động kích hoạt tài khoản.", bold_prefix="Cổng thanh toán SaaS")

    # ============================================================
    # 7.2. Khó khăn và Giải pháp đã áp dụng
    # ============================================================
    add_heading1(doc, "7.2. Khó khăn gặp phải và Giải pháp áp dụng")

    # Khó khăn
    add_heading2(doc, "7.2.1. Khó khăn gặp phải")
    add_bullet(doc, "Model provider và RAG có thể giới hạn tần suất hoặc tăng độ trễ; hệ thống cần timeout, retry hữu hạn, quota và theo dõi staging.", bold_prefix="Tích hợp AI")
    add_bullet(doc, "Việc tính toán tiến độ OKR liên cấp (OKR roll-up progress) dễ gặp xung đột dữ liệu (Race Conditions) khi nhiều nhân viên check-in đồng thời.", bold_prefix="Đồng bộ hóa tiến độ OKR/KPI")
    add_bullet(doc, "Hệ thống claims phân quyền lớn (60 permissions) gây chậm hiệu năng tải trang web khi phải kiểm tra quyền liên tục trên mỗi request.", bold_prefix="Hiệu năng phân quyền Claims")

    # Giải pháp
    add_heading2(doc, "7.2.2. Giải pháp đã áp dụng")
    add_bullet(doc, "Thiết lập cơ chế Rate Limiting ở backend (tối đa 15 calls/phút/user), nạp ngữ cảnh tối giản bằng XML để giảm kích thước gói tin gửi đi, và sử dụng spinner loading trực quan ở client.", bold_prefix="Tối ưu hóa API AI")
    add_bullet(doc, "Áp dụng cơ chế khóa giao dịch DB Transaction và sử dụng Service background xếp hàng đợi xử lý tuần tự các yêu cầu cập nhật tiến độ.", bold_prefix="Khóa đồng bộ tiến độ")
    add_bullet(doc, "Áp dụng IMemoryCache lưu trữ tạm thời claims phân quyền của tài khoản trong phiên làm việc, giảm 90% số lần truy vấn database kiểm tra quyền.", bold_prefix="Caching phân quyền")

    # ============================================================
    # 7.3. Bài học rút ra và Định hướng phát triển
    # ============================================================
    add_heading1(doc, "7.3. Bài học kinh nghiệm và Định hướng phát triển")

    # Bài học
    add_heading2(doc, "7.3.1. Bài học kinh nghiệm")
    add_para(doc, "Qua quá trình làm việc nhóm thiết kế và phát triển dự án, nhóm đã rút ra được những bài học quý báu:")
    add_bullet(doc, "Việc tuân thủ mô hình kiến trúc phân tầng (Layered Architecture) và giao tiếp qua Interfaces giúp nhóm dev dễ dàng code song song mà không bị xung đột mã nguồn.", bold_prefix="Tầm quan trọng của kiến trúc")
    add_bullet(doc, "Cơ sở dữ liệu chuẩn hóa 3NF là cực kỳ quan trọng để đảm bảo dữ liệu không bị dư thừa và dễ viết các truy vấn LINQ phức tạp.", bold_prefix="Chuẩn hóa CSDL")
    add_bullet(doc, "Kiểm thử sớm (Test Early) giúp phát hiện hơn 80% lỗi giao diện và bảo mật ngay từ giai đoạn phát triển, tránh dồn ứ lỗi về cuối kỳ tốt nghiệp.", bold_prefix="Quy trình kiểm thử")

    # Định hướng
    add_heading2(doc, "7.3.2. Định hướng phát triển tương lai")
    add_para(doc, "Để đưa sản phẩm thương mại hóa rộng rãi trên thị trường, định hướng phát triển tiếp theo của nhóm là:")
    add_bullet(doc, "Xây dựng ứng dụng di động native trên hệ điều hành iOS và Android để nhân viên tiện lợi check-in tiến độ mọi lúc mọi nơi.", bold_prefix="Phát triển Mobile App")
    add_bullet(doc, "Hỗ trợ kết nối API trực tiếp với các nguồn dữ liệu kế toán, doanh thu và Google Sheets để tự động cập nhật kết quả KPI.", bold_prefix="Kết nối dữ liệu tự động")
    add_bullet(doc, "Áp dụng cơ chế lưu trữ Vector Database và RAG (Retrieval-Augmented Generation) giúp Trợ lý AI có thể đọc hiểu các tệp quy chế nội bộ của doanh nghiệp để tư vấn nhân sự chính xác hơn.", bold_prefix="Nâng cấp Trí tuệ nhân tạo")

    # Lời kết luận của báo cáo
    add_para(doc, "", space_before=12, space_after=0, indent=False)
    add_para(doc,
        "Dự án 'Hệ thống hỗ trợ vận hành thông minh cho doanh nghiệp vừa và nhỏ hỗ trợ quản lý đa cấp và đưa ra quyết định bằng AI' "
        "đã hoàn thành trọn vẹn các mục tiêu nghiên cứu và ứng dụng công nghệ. Đây là minh chứng cho sự kết hợp "
        "hiệu quả giữa lý thuyết quản trị doanh nghiệp hiện đại và các công nghệ lập trình tiên tiến nhất hiện nay. "
        "Nhóm NEXTGEN hy vọng sản phẩm sẽ đóng góp một phần nhỏ vào tiến trình chuyển đổi số và nâng cao năng suất "
        "cho các doanh nghiệp Việt Nam.",
        bold=True, italic=True, space_before=12, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("  ✓ CHƯƠNG 7: TỔNG KẾT VÀ ĐÁNH GIÁ")
    print("    ✓ 7.1. Kết quả (Thời gian, Mức độ hoàn thành, Chức năng, Hạn chế)")
    print("    ✓ 7.2. Khó khăn & Giải pháp")
    print("    ✓ 7.3. Bài học & Định hướng tương lai")
    write_chapter7(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm Chương 7 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
