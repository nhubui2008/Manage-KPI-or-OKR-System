#!/usr/bin/env python3
"""Build the submission-ready graduation report from the existing DOCX template.

The source document is never modified.  Its cover, styles, section settings,
headers and footers are reused; the report body is rebuilt from verified
repository facts so stale provider-specific prose is not carried forward.

Requires python-docx 1.2+.
"""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT_NAME = "Times New Roman"
BODY_SIZE = 14
CAPTION_SIZE = 12
HEADER_COLOR = "1F4E79"
LIGHT_COLOR = "D9EAF7"
ACCENT_COLOR = "5B3FD6"


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_styles(doc: DocumentType):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    specs = {
        "Heading 1": (14, WD_ALIGN_PARAGRAPH.CENTER, 24, 18),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT, 18, 8),
        "Heading 3": (13, WD_ALIGN_PARAGRAPH.LEFT, 12, 6),
        "Heading 4": (13, WD_ALIGN_PARAGRAPH.LEFT, 10, 4),
    }
    for name, (size, alignment, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT_NAME
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    caption.font.size = Pt(CAPTION_SIZE)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False


def trim_after_cover(doc: DocumentType):
    marker = next(
        (p for p in doc.paragraphs if p.text.strip().upper() == "MỤC LỤC"),
        None,
    )
    if marker is None:
        raise RuntimeError("Không tìm thấy mốc MỤC LỤC trong file nguồn.")
    body = doc._element.body
    started = False
    for child in list(body):
        if child is marker._p:
            started = True
        if started and child.tag != qn("w:sectPr"):
            body.remove(child)


def compact_cover(doc: DocumentType):
    """Keep the seven-member list and footer on the original one-page cover."""
    members = [
        ("Phạm Trần Anh Quân", "TB01758"),
        ("Phạm Trần An An", "TB01817"),
        ("Bùi Nguyễn Anh Như", "TB01785"),
        ("Trần Thanh Phong", "TB01649"),
        ("Nguyễn Thế Bảo", "TB01573"),
        ("Đoàn Quốc Khánh", "TB01544"),
        ("Vũ Hoàng Huy Nhật", "TB01605"),
    ]
    first = next(
        (i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Sinh viên thực hiện:")),
        None,
    )
    if first is None:
        return
    for offset, (name, student_id) in enumerate(members):
        p = doc.paragraphs[first + offset]
        p.clear()
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(Cm(6.5), WD_TAB_ALIGNMENT.LEFT)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(13.3), WD_TAB_ALIGNMENT.LEFT)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_together = True
        text = (
            f"Sinh viên thực hiện:  {name}\t{student_id}"
            if offset == 0
            else f"\t{name}\t{student_id}"
        )
        run = p.add_run(text)
        set_run_font(run, size=11.5, bold=offset == 0)
    for p in doc.paragraphs[first + len(members) : first + len(members) + 3]:
        if not p.text.strip():
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(1)
            for run in p.runs:
                set_run_font(run, size=1)


def add_page_break(doc: DocumentType):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: DocumentType, text: str, level: int, page_break=False):
    if page_break:
        add_page_break(doc)
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper() if level == 1 else text)
    set_run_font(run, 14 if level <= 2 else 13, bold=True)
    return p


def add_para(
    doc: DocumentType,
    text: str,
    *,
    bold=False,
    italic=False,
    indent=True,
    center=False,
    space_before=3,
    space_after=4,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent and not center:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc: DocumentType, title: str, text: str, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.space_after = Pt(3)
    lead = p.add_run("• " + title + (": " if title else ""))
    set_run_font(lead, bold=bool(title))
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, value, *, bold=False, color=None, align=None, size=11):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(value))
    set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc: DocumentType, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, value in enumerate(headers):
        shade_cell(header.cells[idx], HEADER_COLOR)
        set_cell_text(
            header.cells[idx],
            value,
            bold=True,
            color=(255, 255, 255),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=11,
        )
    for ridx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if ridx % 2:
            for cell in row.cells:
                shade_cell(cell, "F3F6FA")
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], value, align=align, size=10.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    return table


def add_field(paragraph, instruction, placeholder="Cập nhật trường khi mở tài liệu"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run)


def add_toc(doc: DocumentType, instruction="TOC \\o \"1-3\" \\h \\z \\u"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_field(p, instruction)


def add_caption(doc: DocumentType, label: str, description: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + " ")
    set_run_font(r1, size=CAPTION_SIZE, italic=True)
    add_field(p, f"SEQ {label} \\* ARABIC", "1")
    r2 = p.add_run(": " + description)
    set_run_font(r2, size=CAPTION_SIZE, italic=True)
    return p


def add_picture(doc: DocumentType, path: Path, description: str, width_cm=16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, "Hình", description)


def add_table_with_caption(doc, headers, rows, caption, widths=None):
    table = add_table(doc, headers, rows, widths)
    add_caption(doc, "Bảng", caption)
    return table


def set_update_fields(doc: DocumentType):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def infer_test_count(status_path: Path) -> str:
    try:
        text = status_path.read_text(encoding="utf-8")
    except OSError:
        return "toàn bộ"
    patterns = [
        r"toàn bộ\s+(\d+)\s+test",
        r"full suite[^\d]*(\d+)\s*/\s*\1",
        r"(\d+)\s*/\s*\1\s+(?:tests?|kiểm thử)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1)
    return "toàn bộ"


def write_front_matter(doc: DocumentType, test_count: str):
    add_page_break(doc)
    add_heading(doc, "MỤC LỤC", 1)
    add_toc(doc)

    add_heading(doc, "DANH MỤC HÌNH ẢNH", 1, page_break=True)
    add_toc(doc, 'TOC \\h \\z \\c "Hình"')

    add_heading(doc, "DANH MỤC BẢNG BIỂU", 1, page_break=True)
    add_toc(doc, 'TOC \\h \\z \\c "Bảng"')

    add_heading(doc, "THEO DÕI PHIÊN BẢN TÀI LIỆU", 1, page_break=True)
    add_table_with_caption(
        doc,
        ["Phiên bản", "Ngày", "Nội dung cập nhật", "Trạng thái"],
        [
            ["1.0", "01/07/2026", "Khởi tạo hồ sơ báo cáo dự án tốt nghiệp.", "Hoàn thành"],
            [
                "2.0",
                "11/08/2026",
                "Chuẩn hóa báo cáo theo kiến trúc AI-native, cập nhật bằng chứng kiểm thử và phạm vi bàn giao đã xác minh.",
                "Bản nộp",
            ],
        ],
        "Theo dõi phiên bản tài liệu",
        [2.2, 3.0, 8.2, 2.5],
    )

    add_heading(doc, "BẢNG CHÚ GIẢI THUẬT NGỮ", 1, page_break=True)
    glossary = [
        ["1", "KPI", "Key Performance Indicator – chỉ số đo lường hiệu suất."],
        ["2", "OKR", "Objectives and Key Results – phương pháp quản trị mục tiêu và kết quả then chốt."],
        ["3", "KR", "Key Result – kết quả then chốt có giá trị mục tiêu và cách đo cụ thể."],
        ["4", "Tenant", "Không gian dữ liệu độc lập của một doanh nghiệp trong hệ thống SaaS."],
        ["5", "RLS", "Row-Level Security – lớp kiểm soát truy cập dữ liệu tại SQL Server."],
        ["6", "RAG", "Retrieval-Augmented Generation – sinh nội dung có bổ sung bằng chứng truy xuất."],
        ["7", "ACL", "Access Control List – danh sách quyền theo người dùng, vai trò và phòng ban."],
        ["8", "AgentRun", "Bản ghi vòng đời thực thi của một tác vụ AI-native."],
        ["9", "Proposal", "Đề xuất có nguồn và phiên bản; không phải dữ liệu nghiệp vụ chính thức."],
        ["10", "Citation", "Tham chiếu đến nguồn dữ liệu được server cấp quyền và xác minh."],
        ["11", "Abstain", "Cơ chế chủ động không kết luận khi bằng chứng không đạt ngưỡng."],
        ["12", "Outbox", "Hàng đợi bền vững được ghi cùng transaction nghiệp vụ."],
        ["13", "RowVersion", "Dấu phiên bản dùng phát hiện cập nhật đồng thời trên SQL Server."],
        ["14", "Idempotency", "Tính chất cho phép retry mà không tạo tác động lặp."],
        ["15", "Rubric", "Bộ tiêu chí đánh giá có phiên bản và ngưỡng confidence."],
    ]
    add_table_with_caption(doc, ["STT", "Thuật ngữ", "Giải thích"], glossary, "Chú giải thuật ngữ", [1.5, 3.5, 12.0])

    add_heading(doc, "DANH SÁCH THÀNH VIÊN", 1, page_break=True)
    members = [
        ["1", "Phạm Trần Anh Quân", "TB01758", "Trưởng nhóm, kiến trúc và AI-native"],
        ["2", "Phạm Trần An An", "TB01817", "Phân tích nghiệp vụ và kiểm thử"],
        ["3", "Bùi Nguyễn Anh Như", "TB01785", "Giao diện và trải nghiệm người dùng"],
        ["4", "Trần Thanh Phong", "TB01649", "Backend và dữ liệu"],
        ["5", "Nguyễn Thế Bảo", "TB01573", "Nghiệp vụ KPI/OKR"],
        ["6", "Đoàn Quốc Khánh", "TB01544", "Kiểm thử và tài liệu"],
        ["7", "Vũ Hoàng Huy Nhật", "TB01605", "Tích hợp và trình bày sản phẩm"],
    ]
    add_table_with_caption(doc, ["STT", "Họ và tên", "Mã sinh viên", "Vai trò"], members, "Danh sách thành viên nhóm NEXTGEN", [1.3, 5.2, 3.0, 7.5])

    add_heading(doc, "GIẢNG VIÊN HƯỚNG DẪN", 1, page_break=True)
    add_para(doc, "Giảng viên hướng dẫn: Phan Hoàng Khải", bold=True, center=True, indent=False)
    add_para(doc, "Chuyên ngành: Phát triển phần mềm", center=True, indent=False)

    add_heading(doc, "LỜI CẢM ƠN", 1, page_break=True)
    add_para(
        doc,
        "Nhóm NEXTGEN trân trọng cảm ơn Trường Cao đẳng FPT Polytechnic, giảng viên hướng dẫn Phan Hoàng Khải và các thầy cô chuyên ngành Phát triển phần mềm đã cung cấp nền tảng kiến thức, góp ý nghiệp vụ và phương pháp kiểm chứng sản phẩm. Những phản hồi trong quá trình thực hiện giúp nhóm hoàn thiện một hệ thống có luồng nghiệp vụ liên kết, dữ liệu có kiểm soát và phần AI đóng vai trò hỗ trợ thay vì thay thế quyết định của con người.",
    )
    add_para(
        doc,
        "Nhóm cũng cảm ơn các thành viên đã phối hợp trong phân tích, thiết kế, lập trình, kiểm thử và chuẩn hóa tài liệu. Báo cáo này ghi nhận kết quả của phiên bản bàn giao đã được đối chiếu với mã nguồn, migration và bộ kiểm thử tại ngày 11/08/2026.",
    )

    add_heading(doc, "LỜI MỞ ĐẦU", 1, page_break=True)
    add_para(
        doc,
        "Doanh nghiệp vừa và nhỏ thường quản lý mục tiêu, chỉ tiêu và công việc trên nhiều bảng tính hoặc kênh trao đổi rời rạc. Khi mục tiêu chiến lược không liên kết trực tiếp với KPI, check-in và đầu việc, người quản lý khó theo dõi nguyên nhân chậm tiến độ và nhân viên thiếu một nguồn dữ liệu thống nhất để phối hợp.",
    )
    add_para(
        doc,
        "Đề tài xây dựng hệ thống quản trị KPI/OKR đa cấp trên ASP.NET Core MVC và SQL Server. Sản phẩm kết nối chuỗi chiến lược → OKR/KR → KPI → công việc → check-in → đánh giá; đồng thời tích hợp các advisor AI-native có nguồn, có kiểm tra quyền, có cơ chế từ chối khi thiếu bằng chứng và có bước con người xác nhận trước mọi thay đổi nghiệp vụ.",
    )
    add_para(
        doc,
        f"Phiên bản báo cáo này tập trung vào phạm vi đã hoàn tất và được kiểm chứng. Bằng chứng kỹ thuật gồm build sạch, {test_count} kiểm thử tự động đạt, migration SQL Server có diễn tập Up/Down và các kiểm thử cô lập tenant, cạnh tranh đồng thời, stale source cùng durable outbox.",
    )

    add_heading(doc, "TÓM TẮT NỘI DUNG DỰ ÁN", 1, page_break=True)
    add_para(
        doc,
        "Hệ thống hỗ trợ quản trị mục tiêu và hiệu suất cho doanh nghiệp theo mô hình SaaS nhiều tenant. Các phân hệ chính gồm xác thực và phân quyền, cơ cấu tổ chức, OKR/KR, KPI, check-in, dự án Kanban, đánh giá cuối kỳ, thưởng, cảnh báo và nhật ký kiểm toán. Quan hệ dữ liệu được bảo vệ bằng validation nghiệp vụ, transaction, unique constraint, row-version, query filter và SQL Server Row-Level Security.",
    )
    add_para(
        doc,
        "Khối AI-native gồm Goal Planning Agent, Check-in AI Evaluator, advisor cho KR, gợi ý KPI/KR, phân tích hiệu suất, phân khúc khách hàng, soạn bản nháp nhận xét và chat có nguồn. Mọi luồng đều dùng snapshot do server dựng trong phạm vi được phép, strict JSON, source fingerprint, citation metadata và cơ chế human-in-the-loop. Dữ liệu chính thức về điểm, trạng thái duyệt, xếp hạng và thưởng chỉ thay đổi qua quy trình nghiệp vụ do người có quyền thực hiện.",
    )


def write_chapter_1(doc: DocumentType):
    add_heading(doc, "CHƯƠNG 1: GIỚI THIỆU", 1, page_break=True)
    add_heading(doc, "1.1. Bối cảnh và lý do chọn đề tài", 2)
    add_heading(doc, "1.1.1. Bài toán quản trị hiệu suất", 3)
    add_para(
        doc,
        "KPI đo lường kết quả vận hành, trong khi OKR giúp tổ chức tập trung vào mục tiêu và kết quả then chốt. Trong thực tế, hai hệ thống thường được theo dõi tách rời khỏi công việc hằng ngày. Điều này tạo ra độ trễ cập nhật, số liệu không đồng nhất và khó truy vết trách nhiệm khi kết quả không đạt kỳ vọng.",
    )
    add_para(
        doc,
        "Một hệ thống phù hợp cần đồng thời giải quyết ba lớp: liên kết chiến lược với thực thi, bảo vệ dữ liệu nhiều doanh nghiệp và hỗ trợ người dùng ra quyết định từ bằng chứng. AI chỉ tạo giá trị khi hoạt động bên trong các ràng buộc đó; một câu trả lời trôi nổi không có nguồn hoặc tự ghi dữ liệu sẽ làm tăng rủi ro thay vì giảm tải cho người quản lý.",
    )

    add_heading(doc, "1.1.2. Giá trị của giải pháp", 3)
    add_bullet(doc, "Một chuỗi dữ liệu thống nhất", "Mục tiêu, KR, KPI, dự án, công việc, check-in và đánh giá được kết nối bằng khóa ngoại và quy tắc nghiệp vụ.")
    add_bullet(doc, "Phân quyền theo phạm vi", "Người dùng chỉ thấy dữ liệu thuộc tenant, vai trò, phòng ban và quan hệ công việc được cấp.")
    add_bullet(doc, "AI có kiểm soát", "Advisor tạo bản nháp có nguồn; server xác minh schema, nguồn, quyền và độ mới trước khi hiển thị hoặc áp dụng.")
    add_bullet(doc, "Khả năng kiểm toán", "Run, proposal, approval, citation metadata, source version và audit log cho phép giải thích một đề xuất được tạo và quyết định như thế nào.")

    add_heading(doc, "1.2. Mục tiêu của đề tài", 2)
    add_heading(doc, "1.2.1. Mục tiêu tổng quát", 3)
    add_para(
        doc,
        "Xây dựng một ứng dụng web quản trị KPI/OKR có thể sử dụng cho quy trình nội bộ của doanh nghiệp vừa và nhỏ, trong đó dữ liệu hiệu suất được kết nối xuyên suốt và AI đóng vai trò đồng hành có bằng chứng, không thay thế thẩm quyền của con người.",
    )
    add_heading(doc, "1.2.2. Mục tiêu cụ thể", 3)
    goals = [
        ("Số hóa chiến lược", "Quản lý sứ mệnh, tầm nhìn, OKR nhiều cấp và các KR đo lường được."),
        ("Quản lý KPI", "Thiết lập kỳ, chỉ tiêu, target, ngưỡng, trọng số và phân công cho cá nhân/phòng ban."),
        ("Theo dõi thực thi", "Ghi nhận check-in, hàng đợi review, nhận xét và lịch sử thay đổi."),
        ("Liên kết công việc", "Quản lý WorkProject/WorkItem trên Kanban và đồng bộ tiến độ có trọng số sang check-in."),
        ("Đánh giá có kiểm soát", "Tính điểm từ dữ liệu đã duyệt, xếp hạng và thưởng theo quy trình phê duyệt."),
        ("AI-native", "Cung cấp advisor có strict schema, nguồn, confidence/abstain, source version và bước phê duyệt."),
        ("An toàn đa tenant", "Kết hợp tenant context, query filter, write guard và RLS tại SQL Server."),
        ("Kiểm chứng được", "Bao phủ các quy tắc trọng yếu bằng unit/integration test và diễn tập migration trên SQL Server."),
    ]
    for title, text in goals:
        add_bullet(doc, title, text)

    add_heading(doc, "1.3. Phạm vi bàn giao", 2)
    add_para(doc, "Báo cáo trình bày các phân hệ đã có mã nguồn, dữ liệu, giao diện hoặc kiểm thử đối chiếu trong repository:")
    scopes = [
        ("Nền tảng doanh nghiệp", "tenant, membership, tài khoản, vai trò, permission, phòng ban, chức vụ, nhân viên và audit."),
        ("Quản trị mục tiêu", "Mission/Vision, OKR, KR, kỳ đánh giá, KPI, cấu hình KPI và phân bổ mục tiêu."),
        ("Vận hành", "check-in, Review Queue, bình luận, cuộc họp 1-on-1, dự án và công việc Kanban."),
        ("Đánh giá", "kết quả cuối kỳ, rank, quy tắc thưởng, báo cáo và xuất dữ liệu."),
        ("AI-native", "chín advisor nghiệp vụ, durable check-in evaluator, rubric có phiên bản, RAG ingestion và trang vận hành."),
        ("An toàn và độ tin cậy", "xác thực cookie/Google, CSRF, reset link một lần, rate limiting, transaction, idempotency, row-version và RLS."),
    ]
    for title, text in scopes:
        add_bullet(doc, title, text)

    add_heading(doc, "1.4. Phương pháp thực hiện", 2)
    add_bullet(doc, "Phân tích domain", "Mô hình hóa tổ chức, mục tiêu, chỉ tiêu, thực thi và đánh giá thành các aggregate có quan hệ rõ ràng.")
    add_bullet(doc, "Ponytail-style minimalism", "Tái sử dụng helper, validator, service và convention ASP.NET Core hiện có; mỗi thay đổi giữ diff tập trung.")
    add_bullet(doc, "Deterministic-first", "Công thức điểm, phân loại, quyền và trạng thái do server quyết định; model chỉ giải thích hoặc tạo bản nháp trong schema giới hạn.")
    add_bullet(doc, "Test theo rủi ro", "Ưu tiên tenant isolation, concurrency, stale source, idempotency, transaction và migration lifecycle.")


def write_chapter_2(doc: DocumentType):
    add_heading(doc, "CHƯƠNG 2: PHÂN TÍCH HỆ THỐNG", 1, page_break=True)
    add_heading(doc, "2.1. Tác nhân hệ thống", 2)
    actors = [
        ["1", "Admin", "Quản trị nền tảng/tenant, tài khoản, vai trò, permission, cấu hình và nhật ký."],
        ["2", "Director", "Thiết lập chiến lược, theo dõi toàn doanh nghiệp và phê duyệt kết quả cuối kỳ."],
        ["3", "HR", "Quản lý nhân sự, kỳ đánh giá, rank, quy tắc thưởng và báo cáo."],
        ["4", "Manager", "Giao KPI, quản lý dự án/công việc, review check-in và đánh giá nhân viên trong phạm vi."],
        ["5", "Employee", "Theo dõi KPI cá nhân, cập nhật công việc, gửi check-in và sử dụng advisor trong phạm vi."],
        ["6", "Background Worker", "Claim outbox/ingestion job, retry có lease và ghi trạng thái bền vững theo tenant."],
    ]
    add_table_with_caption(doc, ["STT", "Tác nhân", "Trách nhiệm"], actors, "Tác nhân và trách nhiệm trong hệ thống", [1.3, 3.5, 13.0])

    add_heading(doc, "2.2. Yêu cầu chức năng", 2)
    requirements = [
        ["FR01", "Xác thực và phiên", "Đăng nhập, Google OAuth, đổi/reset mật khẩu bằng liên kết một lần và vô hiệu hóa phiên cũ."],
        ["FR02", "Tổ chức", "Quản lý tenant, phòng ban đa cấp, chức vụ, nhân viên và assignment."],
        ["FR03", "OKR/KR", "Tạo mục tiêu nhiều cấp, kết quả then chốt và phân bổ cho phòng ban/cá nhân."],
        ["FR04", "KPI", "Thiết lập KPI, target, ngưỡng, chiều thuận/nghịch, trọng số và kỳ đánh giá."],
        ["FR05", "Check-in", "Nhân viên gửi tiến độ; quản lý review với optimistic concurrency và audit."],
        ["FR06", "Kanban", "Quản lý dự án, task, trạng thái, assignee, deadline và liên kết KPI/KR."],
        ["FR07", "Đánh giá", "Tổng hợp dữ liệu đã duyệt, quản lý workflow và xếp hạng phía server."],
        ["FR08", "Thưởng", "Áp dụng BonusRule sau quyết định đánh giá theo đúng thẩm quyền."],
        ["FR09", "Goal Planning", "Sinh đúng ba kế hoạch task có assignee, dependency, rủi ro, data gap và citation."],
        ["FR10", "AI Evaluator", "Mô phỏng check-in, giữ baseline chính thức, confidence breakdown và rubric định tính có phiên bản."],
        ["FR11", "Advisor", "Gợi ý KPI/KR, chat, phân tích hiệu suất, phân khúc và bản nháp nhận xét có nguồn."],
        ["FR12", "RAG", "Quản lý tài liệu, phiên bản, ingestion job, chunk, embedding, index và ACL."],
        ["FR13", "Vận hành AI", "Theo dõi run, proposal, citation, dead-letter, retry và trạng thái xử lý theo tenant."],
        ["FR14", "Audit", "Ghi nhận thao tác người dùng và metadata quyết định AI cần thiết cho truy vết."],
    ]
    add_table_with_caption(doc, ["Mã", "Nhóm", "Yêu cầu"], requirements, "Danh sách yêu cầu chức năng của phiên bản bàn giao", [1.7, 3.5, 12.5])

    add_heading(doc, "2.3. Yêu cầu phi chức năng", 2)
    nfrs = [
        ("Bảo mật", "Cookie authentication, antiforgery, permission attribute, scope filter, secret qua cấu hình môi trường và thông báo lỗi không lộ chi tiết nội bộ."),
        ("Cô lập dữ liệu", "Tenant context xuyên request/worker, global query filter, SaveChanges guard và SQL Server RLS."),
        ("Toàn vẹn", "Transaction, FK, unique index, row-version, idempotency key và source fingerprint."),
        ("Khả năng giải thích", "Mọi nhận định AI dùng source ID hợp lệ; UI hiển thị citation, confidence và data gap."),
        ("An toàn AI", "Strict JSON, timeout/retry hữu hạn, abstain khi thiếu bằng chứng và rollout fail-closed bằng kill switch, Shadow/Pilot."),
        ("Khả năng phục hồi", "Outbox/ingestion worker có atomic claim, lease heartbeat, backoff, dead-letter và recovery."),
        ("Khả năng bảo trì", "ASP.NET Core MVC, EF Core, service interface và test project nằm trong solution."),
        ("Khả dụng", "Giao diện Razor/Bootstrap nhất quán, trạng thái loading rõ ràng và thao tác chính có phản hồi lỗi cụ thể."),
    ]
    for title, text in nfrs:
        add_bullet(doc, title, text)

    add_heading(doc, "2.4. Các luồng nghiệp vụ trọng tâm", 2)
    add_heading(doc, "2.4.1. Chuỗi chiến lược đến đánh giá", 3)
    add_para(doc, "Director thiết lập OKR/KR → Manager phân rã và giao KPI → nhân viên thực hiện task/check-in → Manager review → dịch vụ tính điểm tổng hợp dữ liệu đã duyệt → Director quyết định kết quả cuối kỳ → quy tắc thưởng được áp dụng theo workflow.")
    add_heading(doc, "2.4.2. Goal Planning có phê duyệt", 3)
    add_para(doc, "Người có quyền chọn KPI/OKR/KR/project → server dựng snapshot và tìm nguồn trong ACL → model trả ba phương án strict JSON → server validate/critic và tính fit → người dùng xem, chỉnh và xác nhận → domain validator tạo task trong transaction → tiến độ task được tổng hợp thành check-in Pending và durable outbox.")
    add_heading(doc, "2.4.3. Check-in AI Evaluator", 3)
    add_para(doc, "Check-in được ghi thành công cùng outbox → rollout gate kiểm tra kill switch/mode/tenant/phòng ban → worker claim job theo tenant → evaluator nạp baseline Approved, candidate và rubric đang hiệu lực → server tính projected progress/classification/confidence → model chỉ chấm tiêu chí định tính đủ nguồn → proposal được lưu. Shadow mode chỉ phục vụ quan sát; Pilot/General Availability mới cho quản lý áp dụng vào bản nháp, sửa và tự gửi review.")


def write_chapter_3(doc: DocumentType, docs_dir: Path):
    add_heading(doc, "CHƯƠNG 3: THIẾT KẾ SẢN PHẨM", 1, page_break=True)
    add_heading(doc, "3.1. Kiến trúc tổng thể", 2)
    layers = [
        ["1", "Presentation", "Razor Views, Bootstrap, JavaScript và ViewModel; hiển thị theo permission/scope."],
        ["2", "Controller", "Nhận request, antiforgery, authorization, mapping DTO và điều phối workflow."],
        ["3", "Domain/Application", "Validator, calculator, advisor, persistence, queue và policy dùng chung."],
        ["4", "Infrastructure", "EF Core/SQL Server, HttpClient provider, private storage, Azure Search và worker."],
        ["5", "Cross-cutting", "Tenant context, RLS session context, audit, rate limit, row-version và telemetry."],
    ]
    add_table_with_caption(doc, ["STT", "Lớp", "Trách nhiệm"], layers, "Phân lớp kiến trúc ứng dụng", [1.3, 3.5, 13.0])
    add_para(doc, "Thiết kế tuân theo nguyên tắc controller mỏng, logic dùng chung nằm trong service/validator và dữ liệu chỉ được ghi qua EF Core trong transaction. Các adapter AI, embedding, search và parser được đăng ký qua dependency injection để có thể kiểm thử độc lập.")

    add_heading(doc, "3.2. Thiết kế dữ liệu", 2)
    entities = [
        ["Tenant", "TenantMembership", "Xác định doanh nghiệp, thành viên và vai trò hoạt động."],
        ["SystemUser", "Role/Permission", "Danh tính đăng nhập và quyền truy cập động."],
        ["Department", "Employee/Assignment", "Cơ cấu tổ chức, hồ sơ và phạm vi quản lý."],
        ["OKR", "OKRKeyResult", "Mục tiêu và kết quả then chốt nhiều cấp."],
        ["KPI", "KPIDetail/Assignment", "Chỉ tiêu, công thức, ngưỡng và phân công."],
        ["KPICheckIn", "CheckInDetail", "Lần báo cáo tiến độ và dữ liệu chi tiết."],
        ["WorkProject", "WorkItem", "Dự án Kanban, task, assignee, deadline và liên kết KPI/KR."],
        ["EvaluationResult", "BonusRule", "Kết quả đánh giá và quy tắc thưởng theo workflow."],
        ["AgentRun", "AgentDraftAction/Approval", "Vòng đời AI, bản nháp và quyết định của người dùng."],
        ["AiEvaluationProposal", "CriterionResult/Evidence", "Đề xuất check-in có rubric, confidence và citation metadata."],
        ["EvaluationRubric", "EvaluationCriterion", "Bộ tiêu chí bất biến theo phiên bản và ngưỡng chấm."],
        ["KnowledgeDocument", "Version/Chunk/Job", "Tài liệu, phiên bản, đoạn chỉ mục và ingestion bền vững."],
        ["CheckInAiEvaluationOutbox", "Lease/SourceVersion", "Hàng đợi đánh giá chống trùng và phục hồi được."],
    ]
    add_table_with_caption(doc, ["Thực thể chính", "Thực thể liên quan", "Vai trò"], entities, "Nhóm thực thể cốt lõi và AI-native", [4.0, 5.0, 9.0])

    add_heading(doc, "3.3. Thiết kế bảo mật đa tenant", 2)
    add_bullet(doc, "Tầng request", "Middleware xác định tenant/membership hiện hành và tạo principal chuẩn từ vai trò đang hoạt động.")
    add_bullet(doc, "Tầng truy vấn", "Global query filter tự thêm TenantId vào truy vấn entity nghiệp vụ.")
    add_bullet(doc, "Tầng ghi", "SaveChanges guard gắn TenantId và từ chối tham chiếu chéo tenant.")
    add_bullet(doc, "Tầng SQL", "Interceptor thiết lập SESSION_CONTEXT; RLS filter/block predicate bảo vệ truy vấn và thao tác ghi trực tiếp.")
    add_bullet(doc, "Tầng AI/RAG", "Filter tenant/ACL do server sinh, source được tái kiểm tra trước model và trước khi dùng kết quả.")

    add_heading(doc, "3.4. Thiết kế giao diện", 2)
    add_para(doc, "Giao diện sử dụng layout sidebar, khu vực nội dung theo card và các CTA nhất quán. Trạng thái rủi ro, tiến độ, review và AI proposal được phân biệt bằng màu, nhãn và thông tin nguồn; thao tác áp dụng AI chỉ điền bản nháp để người dùng tiếp tục chỉnh sửa.")
    screenshots = [
        ("screenshot_login.png", "Cổng đăng nhập và lựa chọn phương thức xác thực"),
        ("screenshot_dashboard.png", "Dashboard tổng quan theo vai trò"),
        ("screenshot_okrs.png", "Màn hình quản lý mục tiêu OKR"),
        ("screenshot_kpis.png", "Màn hình quản lý và giao KPI"),
        ("screenshot_checkin.png", "Màn hình theo dõi check-in KPI"),
        ("screenshot_ai.png", "Bizen AI Assistant trong ngữ cảnh check-in"),
    ]
    for filename, caption in screenshots:
        path = docs_dir / filename
        if path.exists():
            add_picture(doc, path, caption)


def write_chapter_4(doc: DocumentType, docs_dir: Path):
    add_heading(doc, "CHƯƠNG 4: THỰC THI MÃ NGUỒN", 1, page_break=True)
    add_heading(doc, "4.1. Công nghệ sử dụng", 2)
    tech = [
        [".NET / ASP.NET Core MVC", "10", "Web framework, routing, Razor, authentication, authorization và DI."],
        ["Entity Framework Core", "10", "ORM, migration, transaction, concurrency token và query filter."],
        ["SQL Server", "Relational", "CSDL chính, constraint, index, row-version và Row-Level Security."],
        ["Bootstrap + JavaScript", "Repo-native", "Giao diện responsive, modal, loading và tương tác form."],
        ["xUnit", "2.9", "Unit/integration/security/SQL Server regression tests."],
        ["EPPlus", "7.7", "Nhập/xuất dữ liệu Excel cho nghiệp vụ nhân sự và báo cáo."],
        ["DeepSeek adapter", "IAIModelClient", "Cổng mô hình strict JSON qua typed HttpClient."],
        ["MinerU + BGE-M3", "HTTP adapters", "Trích xuất tài liệu và embedding 1.024 chiều."],
        ["Azure AI Search", "Hybrid/RAG", "Chỉ mục chunk, hybrid retrieval và security trimming."],
        ["ClamAV", "Private daemon", "Quét nội dung trước pipeline ingestion tài liệu."],
    ]
    add_table_with_caption(doc, ["Công nghệ", "Phiên bản/contract", "Vai trò"], tech, "Công nghệ và thành phần kỹ thuật", [4.5, 3.5, 10.0])

    add_heading(doc, "4.2. Tổ chức mã nguồn", 2)
    folders = [
        ("Controllers/", "Endpoint MVC/API, authorization và điều phối workflow."),
        ("Models/", "Entity, DTO, ViewModel và contract AI có cấu trúc."),
        ("Services/", "Nghiệp vụ KPI/OKR, validator, calculator, queue và advisor."),
        ("Services/AI/", "Model client, RAG, evaluator, proposal persistence và durable workers."),
        ("Data/", "DbContext, model configuration và tenant/RLS integration."),
        ("Views/", "Razor UI theo vai trò, permission và human review."),
        ("Migrations/", "Lịch sử schema, constraint, AI persistence và RLS policy."),
        ("tests/", "Bộ kiểm thử unit, controller, service, security và SQL Server."),
    ]
    for title, text in folders:
        add_bullet(doc, title, text)

    add_heading(doc, "4.3. Các luồng AI-native đã triển khai", 2)
    advisors = [
        ["Goal Planning Agent", "KPI/OKR/KR/project", "Ba kế hoạch task; fit server-side; xác nhận mới tạo task."],
        ["Check-in AI Evaluator", "Approved baseline + candidate", "Projected score, classification, rubric, confidence, citation và data gap."],
        ["OKR KR Advisor", "KR chính thức + giá trị dự kiến", "Đề xuất có fingerprint; accept/reject chỉ ghi metadata."],
        ["Evaluation Review Draft", "Kết quả đánh giá trong scope", "Bản nháp nhận xét có source/version; áp dụng vào form."],
        ["Customer Segment Advisor", "Dữ liệu thương mại trong scope", "Phân khúc và hành động tham khảo, không xếp hạng nhân sự."],
        ["Performance Analysis", "Check-in Approved", "Tổng quan, điểm mạnh, rủi ro và hành động có nguồn."],
        ["KPI Suggestion", "Kỳ/OKR/KR/phạm vi", "3–5 bản nháp được server kiểm tra target/ngưỡng/đơn vị."],
        ["Chat Advisor", "Snapshot KPI/OKR + RAG ACL", "Câu trả lời strict JSON, citation và abstain."],
        ["KR Suggestion", "Objective + KR hiện hữu", "Bản nháp KR không trùng, đơn vị chuẩn và nguồn hợp lệ."],
    ]
    add_table_with_caption(doc, ["Advisor", "Nguồn", "Kết quả và cổng an toàn"], advisors, "Danh mục các luồng AI-native", [4.3, 4.7, 9.0])

    add_heading(doc, "4.4. Hợp đồng AI an toàn", 2)
    add_bullet(doc, "Authorized snapshot", "Server nạp dữ liệu tối thiểu theo tenant, role, employee/department scope và trạng thái workflow.")
    add_bullet(doc, "Strict output", "Model phải trả đúng JSON schema; trường thừa, citation giả, score ngoài ngưỡng hoặc JSON lỗi bị từ chối.")
    add_bullet(doc, "Source fingerprint", "Fingerprint bao gồm entity, trạng thái, phiên bản rubric và nguồn; thay đổi làm proposal cũ thành Stale.")
    add_bullet(doc, "Recheck", "Quyền và nguồn được nạp lại sau model call và ngay trước thao tác áp dụng.")
    add_bullet(doc, "Metadata-only", "Lưu model/prompt version, latency, token, citation hash và quyết định; không lưu prompt/context/raw response cho các luồng mới.")
    add_bullet(doc, "Human-in-the-loop", "AI chỉ điền draft hoặc giải thích; form/domain command chuẩn vẫn là cổng ghi chính thức.")

    add_heading(doc, "4.5. Goal Planning và durable execution", 2)
    add_para(doc, "Goal Planning Agent duy trì vòng đời Planning → Retrieving → Generating → Validating → Critiquing → WaitingApproval → Executing → Completed/Failed. Approval token dùng một lần, row-version và idempotency ngăn xác nhận lặp. Khi task đã được xác nhận, IWorkItemCommandValidator kiểm tra assignee, phòng ban, deadline, KPI/KR và project scope giống thao tác thủ công.")
    add_para(doc, "Các task cùng KPI/nhân viên được tổng hợp theo KpiImpactWeight thành một check-in Pending trong ngày. Project, task, liên kết, audit, check-in, detail và outbox được ghi trong transaction Serializable; worker xử lý outbox bằng atomic claim, heartbeat, retry/backoff và dead-letter.")

    add_heading(doc, "4.6. Check-in Evaluator và rubric có phiên bản", 2)
    add_para(doc, "Evaluator tách official baseline từ check-in Approved và projected score của bản ghi đang xét. Phân loại định lượng do server tính từ target, chiều KPI, assignment weight, deadline và tiến độ kỳ. Confidence được phân rã 40% độ phủ, 25% thẩm quyền, 20% nhất quán và 15% độ mới. Dưới 0,60, hệ thống giữ kết quả định lượng nhưng không phát sinh điểm định tính.")
    add_para(doc, "EvaluationRubric/EvaluationCriterion tạo theo phiên bản bất biến. Khi phát hành phiên bản mới, proposal tạm thời có nguồn cũ chuyển Stale và check-in Pending được đưa lại vào outbox. Khi quản lý áp dụng proposal, row-version và source version được kiểm tra trong transaction; quyết định cuối lệch trên 10 điểm so với baseline bắt buộc có lý do.")

    sequence = docs_dir / "sequence_ai.png"
    if sequence.exists():
        add_heading(doc, "4.7. Sơ đồ tuần tự advisor có nguồn", 2)
        add_picture(doc, sequence, "Sơ đồ tuần tự KPI Suggestion Advisor có nguồn", width_cm=16.2)
        add_para(doc, "Sơ đồ thể hiện bốn cổng chính: xác thực quyền, dựng snapshot, schema/citation validation và rebuild snapshot trong transaction trước khi trả bản nháp cho browser.")

    add_heading(doc, "4.8. Pipeline tài liệu RAG", 2)
    steps = [
        ["1", "Tiếp nhận", "Kiểm tra loại/kích thước, chữ ký tệp, checksum, nguồn và ACL."],
        ["2", "An toàn", "Quét ClamAV và lưu bản gốc trong private storage."],
        ["3", "Trích xuất", "Gửi job idempotent sang MinerU; poll bằng durable lease."],
        ["4", "Chuẩn hóa", "Parse UTF-8/JSON, giới hạn chunk và gắn page/section/table metadata."],
        ["5", "Embedding", "Tạo vector BGE-M3 đúng 1.024 chiều và lưu model version."],
        ["6", "Index", "Upsert Azure Search với tenant/ACL, reliability và source version."],
        ["7", "Truy xuất", "Hybrid search dùng filter do server sinh; nguồn được đối chiếu lại với SQL."],
    ]
    add_table_with_caption(doc, ["Bước", "Giai đoạn", "Kiểm soát"], steps, "Pipeline ingestion và retrieval RAG", [1.5, 3.2, 13.5])

    add_heading(doc, "4.9. Cô lập tenant và RLS", 2)
    add_para(doc, "Migration RLS thiết lập policy cho 57 bảng thuộc nghiệp vụ, AI và tài liệu. Interceptor gắn TenantId/SystemUserId vào SESSION_CONTEXT ở mỗi connection. Worker nền không dùng bypass; mỗi lần claim chạy trong đúng tenant. Kiểm thử SQL Server chứng minh raw SQL và IgnoreQueryFilters vẫn bị lọc, thao tác ghi chéo tenant bị chặn và pooled connection được đặt lại đúng context.")


def write_chapter_5(doc: DocumentType, test_count: str):
    add_heading(doc, "CHƯƠNG 5: KIỂM THỬ VÀ NGHIỆM THU", 1, page_break=True)
    add_heading(doc, "5.1. Chiến lược kiểm thử", 2)
    add_para(doc, "Kiểm thử được tổ chức theo rủi ro thay vì chỉ theo màn hình. Quy tắc định lượng được kiểm tra ở service; quyền và workflow được kiểm tra ở controller; tenant/RLS, migration, constraint và concurrency được chạy trên SQL Server. Các luồng AI dùng fake model/retriever để kiểm tra schema, nguồn, stale source, idempotency và bảo đảm không ghi dữ liệu chính thức khi thiếu phê duyệt.")
    test_levels = [
        ["Unit", "Calculator, parser, confidence, source fingerprint và validator."],
        ["Service integration", "Goal Planning, advisor, persistence, queue, worker, timeout/retry."],
        ["Controller", "CSRF, permission, scope, row-version, idempotency và workflow."],
        ["Relational", "SQL Server constraint, transaction, concurrency, RLS và migration lifecycle."],
        ["Static", "Build, EF model drift, JavaScript syntax và git diff whitespace."],
    ]
    add_table_with_caption(doc, ["Mức", "Phạm vi"], test_levels, "Các mức kiểm thử của hệ thống", [4.0, 14.0])

    add_heading(doc, "5.2. Bằng chứng kiểm chứng bản bàn giao", 2)
    count_text = f"{test_count}/{test_count}" if test_count.isdigit() else test_count
    evidence = [
        ["Build solution", "Đạt", "0 warning, 0 error trên .NET 10 solution."],
        ["Bộ kiểm thử tự động", "Đạt", f"{count_text} kiểm thử chạy xanh tại thời điểm đóng gói báo cáo."],
        ["EF model drift", "Đạt", "Không có thay đổi model chưa được thể hiện bằng migration."],
        ["Migration lifecycle", "Đạt", "Database rỗng → latest; Down/Up/reapply nhóm migration AI trên SQL Server."],
        ["Tenant RLS", "Đạt", "Raw SQL/IgnoreQueryFilters bị lọc; INSERT/UPDATE chéo tenant bị chặn."],
        ["Pooled connection", "Đạt", "SESSION_CONTEXT được đặt lại khi connection tái sử dụng."],
        ["Outbox và rollout", "Đạt", "Atomic claim/lease/retry; kill switch, Shadow và Pilot chặn đúng enqueue, worker, UI/server apply."],
        ["Goal Planning", "Đạt", "Concurrent draft, double-confirm, stale source, token rotation và reject."],
        ["Check-in rubric", "Đạt", "Cạnh tranh writer hội tụ, version tuần tự, một active rubric và requeue đúng source."],
        ["Proposal safety", "Đạt", "Citation giả/thu hồi, malformed JSON, low-confidence và row-version bị chặn đúng contract."],
    ]
    add_table_with_caption(doc, ["Hạng mục", "Kết quả", "Bằng chứng"], evidence, "Bằng chứng kỹ thuật của phiên bản bàn giao", [4.0, 2.2, 11.8])

    add_heading(doc, "5.3. Kịch bản nghiệm thu tiêu biểu", 2)
    cases = [
        ["TC01", "Giao KPI đúng phạm vi", "KPI, detail và assignment được lưu; ngoài scope bị từ chối."],
        ["TC02", "Check-in có submission trùng", "Idempotency/constraint ngăn double-submit."],
        ["TC03", "Review đồng thời", "Row-version phát hiện xung đột; không ghi đè quyết định mới."],
        ["TC04", "Nhiều task cùng KR", "Quan hệ một KR–nhiều task hoạt động qua index không unique."],
        ["TC05", "Một OKR–nhiều project", "SourceOKRId là nguồn duy nhất; FK/index bảo vệ quan hệ."],
        ["TC06", "Goal Planning xác nhận lặp", "Một approval token chỉ thực thi một lần; task không bị nhân đôi."],
        ["TC07", "Nguồn đổi sau model call", "Fingerprint conflict; bản nháp cũ không được áp dụng."],
        ["TC08", "Check-in confidence thấp", "Giữ phân loại định lượng; phần định tính abstain."],
        ["TC09", "Rubric phát hành đồng thời", "Version tuần tự và chỉ một phiên bản active."],
        ["TC10", "Citation giả", "Parser từ chối source ID không thuộc tập server cấp."],
        ["TC11", "Truy cập chéo tenant", "Query/command/raw SQL đều không đọc hoặc ghi dữ liệu tenant khác."],
        ["TC12", "Worker hết lease", "Job được retry/dead-letter; không treo ở trạng thái Processing."],
    ]
    add_table_with_caption(doc, ["Mã", "Tình huống", "Kết quả nghiệm thu"], cases, "Kịch bản nghiệm thu nghiệp vụ và AI-native", [1.8, 6.0, 10.2])

    add_heading(doc, "5.4. Kết luận nghiệm thu", 2)
    add_para(doc, "Các luồng được lựa chọn cho phiên bản bàn giao đáp ứng tiêu chí: build sạch; toàn bộ test tự động đạt; schema khớp migration; tenant isolation và transaction được kiểm tra trên SQL Server; AI không có đường ghi trực tiếp vào điểm, phê duyệt, xếp hạng hoặc thưởng; mọi đề xuất có source version và cổng xác nhận của người dùng.")


def write_chapter_6(doc: DocumentType):
    add_heading(doc, "CHƯƠNG 6: HƯỚNG DẪN SỬ DỤNG", 1, page_break=True)
    add_heading(doc, "6.1. Khởi động môi trường trình bày", 2)
    add_para(doc, "Tại thư mục gốc dự án, cấu hình connection string và các secret provider bằng biến môi trường/User Secrets, sau đó chạy migration và ứng dụng:")
    add_bullet(doc, "Bước 1", "dotnet restore")
    add_bullet(doc, "Bước 2", "dotnet ef database update")
    add_bullet(doc, "Bước 3", "dotnet run")
    add_bullet(doc, "Bước 4", "Mở URL cục bộ được ứng dụng in ra và đăng nhập bằng tài khoản demo được cấp riêng cho buổi trình bày.")
    add_para(doc, "Mật khẩu, API key và connection string không được ghi trong báo cáo hoặc commit vào repository.", bold=True)

    add_heading(doc, "6.2. Hướng dẫn theo vai trò", 2)
    add_heading(doc, "6.2.1. Admin", 3)
    add_bullet(doc, "Quản lý tenant/tài khoản", "Kiểm tra membership, trạng thái hoạt động, role và phạm vi dữ liệu.")
    add_bullet(doc, "Phân quyền", "Cấu hình role–permission và kiểm tra Audit Logs sau thay đổi.")
    add_bullet(doc, "Vận hành AI/RAG", "Theo dõi tài liệu, ingestion job, run/outbox lỗi và chạy retry cho bản ghi thuộc tenant.")
    add_heading(doc, "6.2.2. Director", 3)
    add_bullet(doc, "Chiến lược", "Thiết lập Mission/Vision, OKR công ty và KR đo lường được.")
    add_bullet(doc, "Theo dõi", "Sử dụng Dashboard và Performance Analysis Advisor để đọc số liệu đã duyệt có nguồn.")
    add_bullet(doc, "Phê duyệt", "Xem tổng hợp đánh giá, nhận xét và quyết định kết quả cuối kỳ theo workflow.")
    add_heading(doc, "6.2.3. Manager", 3)
    add_bullet(doc, "Giao KPI", "Chọn kỳ, nhân viên/phòng ban, target, ngưỡng, trọng số và quan hệ OKR/KR.")
    add_bullet(doc, "Quản lý công việc", "Tạo project/task, gán assignee, deadline, KPI/KR và cập nhật Kanban.")
    add_bullet(doc, "Review check-in", "Mở Employee Tracking, xem baseline/proposal AI, áp dụng vào bản nháp, chỉnh sửa và tự gửi quyết định.")
    add_heading(doc, "6.2.4. Employee", 3)
    add_bullet(doc, "KPI cá nhân", "Theo dõi target, deadline, trạng thái và lịch check-in.")
    add_bullet(doc, "Check-in", "Nhập giá trị đạt được, minh chứng/ghi chú và gửi; hệ thống tự tính tiến độ theo assignment weight.")
    add_bullet(doc, "Công việc", "Cập nhật task trong phạm vi được giao; task liên kết KPI được tổng hợp theo trọng số.")

    add_heading(doc, "6.3. Hướng dẫn các tính năng AI-native", 2)
    add_heading(doc, "6.3.1. Goal Planning Agent", 3)
    add_para(doc, "Từ KPI/OKR/KR hoặc dự án, chọn tạo kế hoạch bằng AI. Kiểm tra ba phương án, citation, data gap, dependency và fit. Chọn phương án phù hợp, chỉnh assignee/deadline/đóng góp, sau đó xác nhận. Hệ thống kiểm tra lại nguồn và quyền trước khi tạo task.")
    add_heading(doc, "6.3.2. Check-in AI Evaluator", 3)
    add_para(doc, "Sau khi check-in được gửi, proposal xuất hiện tại trang theo dõi nhân viên. Đọc official baseline, projected score, phân loại định lượng, confidence breakdown, tiêu chí rubric và nguồn. Ở Shadow mode, giao diện ghi rõ chế độ quan sát và không hiển thị khả năng áp dụng. Khi Pilot/General Availability được mở đúng tenant/phòng ban, nút áp dụng mới sao chép đề xuất còn hiệu lực vào form review; quản lý chỉnh và tự quyết định Duyệt/Từ chối.")
    add_heading(doc, "6.3.3. Gợi ý KPI và KR", 3)
    add_para(doc, "Chọn đúng kỳ và phạm vi trước khi gọi advisor. Mỗi gợi ý có source ID và đã qua validator về tên, đơn vị, target, ngưỡng và quan hệ. Áp dụng bản nháp vào form, kiểm tra lại rồi gửi form nghiệp vụ chuẩn để tạo bản ghi chính thức.")
    add_heading(doc, "6.3.4. Chat và phân tích hiệu suất", 3)
    add_para(doc, "Đặt câu hỏi ngắn, gắn với mục tiêu/chỉ tiêu đang xem. Câu trả lời hiển thị citation và có thể chủ động báo thiếu dữ liệu. Performance Analysis chỉ tổng hợp check-in Approved và trả insight tham khảo; không thay đổi điểm hoặc quyết định nhân sự.")
    add_heading(doc, "6.3.5. Quản lý rubric", 3)
    add_para(doc, "Từ KPI Details, người có quyền KPIS_EDIT mở Evaluation Rubrics, xem phiên bản đang hoạt động/lịch sử và tạo phiên bản mới. Các phiên bản đã phát hành chỉ đọc. Sau phát hành, proposal cũ được đánh dấu Stale và check-in Pending được đánh giá lại theo rubric mới.")

    add_heading(doc, "6.4. Nguyên tắc sử dụng an toàn", 2)
    add_bullet(doc, "Kiểm tra nguồn", "Chỉ dùng nhận định có citation hợp lệ và đúng ngữ cảnh nghiệp vụ.")
    add_bullet(doc, "Kiểm tra trạng thái", "Phân biệt baseline chính thức, projected candidate và dữ liệu Approved/Pending.")
    add_bullet(doc, "Giữ quyền quyết định", "Người có thẩm quyền phải xem, chỉnh và gửi form chuẩn; không xem proposal là kết luận tự động.")
    add_bullet(doc, "Không đưa secret/PII vào prompt", "Sử dụng dữ liệu do server nạp trong scope, không dán mật khẩu, key hoặc dữ liệu cá nhân không cần thiết.")


def write_chapter_7(doc: DocumentType, test_count: str):
    add_heading(doc, "CHƯƠNG 7: TỔNG KẾT VÀ ĐÁNH GIÁ", 1, page_break=True)
    add_heading(doc, "7.1. Kết quả thực hiện", 2)
    add_para(doc, "Nhóm đã hoàn thiện một hệ thống ASP.NET Core MVC kết nối đầy đủ chuỗi mục tiêu và thực thi: cơ cấu tổ chức → OKR/KR → KPI → dự án/công việc → check-in → đánh giá → thưởng. Dữ liệu được tổ chức theo tenant và bảo vệ đồng thời ở application layer lẫn SQL Server RLS.")
    add_para(doc, "Khối AI-native đã chuyển từ mô hình gọi chatbot rời rạc sang các advisor nghiệp vụ có contract rõ ràng. Mỗi advisor nhận authorized snapshot, trả strict JSON, dùng citation do server cấp, phát hiện source stale và áp dụng nguyên tắc AI đề xuất – con người quyết định. Goal Planning tạo task qua validator/transaction; Check-in Evaluator giữ baseline Approved, rubric phiên bản và confidence/abstain.")

    add_heading(doc, "7.2. Giá trị nổi bật của sản phẩm", 2)
    values = [
        ("Liên kết chiến lược với công việc", "Một thay đổi task có thể phản ánh vào check-in có trọng số và tiếp tục đi qua review chính thức."),
        ("Quyết định dựa trên bằng chứng", "Advisor hiển thị source, confidence và data gap thay vì chỉ trả văn bản thuyết phục."),
        ("Không trao quyền ghi cho model", "Điểm, trạng thái, rank, thưởng và dữ liệu HR chỉ thay đổi qua domain command đã kiểm tra quyền."),
        ("Cô lập tenant nhiều lớp", "Query filter, write guard, SESSION_CONTEXT và RLS giảm rủi ro lọt dữ liệu."),
        ("Khả năng phục hồi", "Outbox/worker dùng idempotency, lease, heartbeat, backoff và dead-letter."),
        ("Khả năng kiểm chứng", f"Build sạch và {test_count} kiểm thử tự động cùng diễn tập SQL Server tạo bằng chứng có thể lặp lại."),
    ]
    for title, text in values:
        add_bullet(doc, title, text)

    add_heading(doc, "7.3. Bài học kỹ thuật", 2)
    add_bullet(doc, "Chia sẻ quy tắc", "Luồng thủ công và AI phải dùng cùng validator/calculator để tránh hai chuẩn nghiệp vụ.")
    add_bullet(doc, "Khóa đúng thứ tự", "Transaction và thứ tự lock nhất quán là điều kiện để tránh proposal/rubric/check-in đảo trạng thái khi cạnh tranh.")
    add_bullet(doc, "Nguồn là một phần của contract", "Citation ID, source version và ACL recheck cần được thiết kế ngay trong DTO/persistence, không bổ sung sau.")
    add_bullet(doc, "Confidence là chất lượng dữ liệu", "Điểm confidence phải do server tính từ coverage/authority/consistency/freshness, không lấy từ lời tự đánh giá của model.")
    add_bullet(doc, "Migration là sản phẩm", "Up/Down, preflight, constraint và rehearsal SQL Server có giá trị ngang với code application.")

    add_heading(doc, "7.4. Kết luận", 2)
    add_para(doc, "Phiên bản bàn giao vận hành trọn vẹn các luồng quản trị đã nghiệm thu của hệ thống KPI/OKR AI-native. Sản phẩm không chỉ số hóa biểu mẫu mà tạo được liên kết dữ liệu xuyên quy trình, đưa AI vào đúng điểm hỗ trợ và giữ toàn bộ quyết định chính thức trong tay người dùng có thẩm quyền.")
    add_para(doc, "Kết quả nổi bật nhất là sự kết hợp giữa tính thực dụng của ASP.NET Core/EF Core/SQL Server với các cơ chế AI hiện đại nhưng có kiểm soát: RAG theo ACL, strict schema, citation, abstain, source fingerprint, durable workflow và human approval. Đây là nền tảng kỹ thuật và nghiệp vụ nhất quán cho buổi bảo vệ dự án tốt nghiệp.")


def write_appendices(doc: DocumentType, test_count: str):
    add_heading(doc, "PHỤ LỤC A: CÁC ENDPOINT VÀ MÀN HÌNH CHÍNH", 1, page_break=True)
    routes = [
        ["/Dashboard", "Dashboard theo vai trò và Performance Analysis."],
        ["/OKRs", "Quản lý OKR/KR và advisor cho KR."],
        ["/KPIs", "Thiết lập/giao KPI và gợi ý KPI có nguồn."],
        ["/KPICheckIns", "Check-in, Employee Tracking, Review Queue và AI proposal."],
        ["/EvaluationRubrics?KpiId=...", "Quản lý phiên bản rubric của KPI."],
        ["/EvaluationResults", "Workflow đánh giá cuối kỳ và bản nháp nhận xét."],
        ["/WorkProjects", "Project, Kanban task và Goal Planning Agent."],
        ["/KnowledgeDocuments", "Quản trị tài liệu/ingestion, theo dõi và retry outbox tenant-scoped."],
        ["POST /AI/Chat", "Chat advisor strict JSON có nguồn."],
        ["POST /AI/SuggestKPI", "Sinh 3–5 bản nháp KPI trong scope."],
        ["POST /AI/CreateGoalPlanningDraft", "Tạo draft kế hoạch công việc."],
        ["POST /AI/ConfirmDecompose", "Xác nhận draft và tạo task qua domain validator."],
        ["POST /AI/EvaluateCheckInProposal", "Đánh giá check-in theo baseline/rubric."],
        ["POST /AI/DecideCheckInProposal", "Ghi quyết định metadata cho proposal."],
        ["POST /AI/AnalyzePerformance", "Phân tích check-in Approved có citation."],
    ]
    add_table_with_caption(doc, ["Route", "Mục đích"], routes, "Các route phục vụ demo và kiểm tra", [6.5, 11.5])

    add_heading(doc, "PHỤ LỤC B: HỒ SƠ DỮ LIỆU AI-NATIVE", 1, page_break=True)
    fields = [
        ["AgentRun", "TenantId, RunType, CorrelationId, State, FailureCode, ApprovalTokenHash, RequestedBySystemUserId, CreatedAtUtc, UpdatedAtUtc, RowVersion"],
        ["AgentDraftAction", "AgentRunId, EvaluationResultId, SourceEntityType/Id/Version, ActionType, Status, DraftText, CreatedAtUtc, UpdatedAtUtc, RowVersion"],
        ["AgentApproval", "AgentRunId, ApprovedBySystemUserId, Decision, IdempotencyKey, ResultEntityId, AppliedItemCount, DecidedAtUtc"],
        ["AiEvaluationProposal", "AgentRunId, KPICheckInId, EvaluationResultId, EvaluationRubricId/RubricVersion, SourceEntityType/Id/Version, Status, official/projected values, confidence breakdown, human-decision metadata, RowVersion"],
        ["AiEvaluationCriterionResult", "AiEvaluationProposalId, EvaluationCriterionId, RubricVersion, ProposedStatus/ProposedScorePercent, ConfidenceScore, CitationCount, CreatedAtUtc, RowVersion"],
        ["EvidenceReferenceMetadata", "AgentRunId/AiEvaluationProposalId, SourceType/Id/Title/VersionId/Page/Section, ObservedAtUtc, Reliability, IsDirectlyRelevant, IsCurrent"],
        ["EvaluationRubric", "KPIId, PeriodId, Version, Name, IsActive, OnTrackPercent, AtRiskPercent, MinimumConfidenceToPropose, EffectiveFromUtc, SupersededAtUtc, RowVersion"],
        ["EvaluationCriterion", "EvaluationRubricId, Ordinal, Name, Description, MeasurementType, WeightPercent, confidence/score bounds, IsActive, RowVersion"],
        ["CheckInAiEvaluationOutbox", "TenantId, CheckInId, SourceVersion, RequestedBySystemUserId, State, AttemptCount, AvailableAtUtc, LeaseId/LeaseExpiresAtUtc, LastFailureCode, Created/CompletedAtUtc, RowVersion"],
        ["KnowledgeDocument", "TenantId, Title, OwnerSystemUserId, AccessPrincipalsJson, AccessPolicyVersion, IsDeleted, Created/UpdatedAtUtc, RowVersion"],
        ["KnowledgeDocumentVersion", "DocumentId, VersionNumber, ContentSha256, SourceBlobUri, OriginalFileName, ContentType, FileSizeBytes, Status, CreatedAtUtc, RowVersion"],
        ["KnowledgeChunk", "DocumentVersionId, PipelineVersion, AccessPolicyVersion, Ordinal, ContentSha256/ContentBlobUri, SearchIndexKey, Page/Section, TokenCount, IsActive, RowVersion"],
        ["DocumentIngestionJob", "DocumentVersionId, Operation, PipelineVersion, AccessPolicyVersion, State, AttemptCount, AvailableAtUtc, lease fields, MinerUJobId, ParserResultBlobUri, LastFailureCode, timestamps, RowVersion"],
    ]
    add_table_with_caption(doc, ["Thực thể", "Thuộc tính nghiệp vụ chính"], fields, "Dữ liệu kiểm toán và vận hành AI-native", [5.0, 13.0])

    add_heading(doc, "PHỤ LỤC C: LỆNH KIỂM CHỨNG", 1, page_break=True)
    commands = [
        ["Restore", "dotnet restore"],
        ["Build", "dotnet build Manage-KPI-or-OKR-System.sln --no-restore --nologo --verbosity minimal"],
        ["Full test", "dotnet test Manage-KPI-or-OKR-System.sln --no-restore --nologo"],
        ["EF drift", "dotnet ef migrations has-pending-model-changes --no-build"],
        ["Migration script", "dotnet ef migrations script --idempotent"],
        ["Whitespace", "git diff --check"],
    ]
    add_table_with_caption(doc, ["Mục đích", "Lệnh"], commands, "Lệnh tái kiểm chứng phiên bản bàn giao", [4.0, 14.0])
    add_para(doc, f"Kết quả đóng gói ngày 11/08/2026: build 0 warning/0 error và {test_count} kiểm thử tự động đạt. Các kiểm thử relational sử dụng SQL Server cô lập và xóa database tạm sau khi hoàn tất.")

    add_heading(doc, "PHỤ LỤC D: THÔNG TIN BÀN GIAO", 1, page_break=True)
    handoff = [
        ["Kho mã nguồn", "https://github.com/nhubui2008/Manage-KPI-or-OKR-System"],
        ["Nhánh làm việc", "main"],
        ["Nền tảng", ".NET 10 / ASP.NET Core MVC / SQL Server"],
        ["Đường dẫn demo", "URL localhost được hiển thị bởi dotnet run"],
        ["Tài khoản demo", "Cấp riêng trong buổi trình bày; không lưu mật khẩu trong báo cáo"],
        ["Nguyên tắc AI", "AI đề xuất – con người quyết định"],
    ]
    add_table_with_caption(doc, ["Hạng mục", "Thông tin"], handoff, "Thông tin bàn giao dự án", [5.0, 13.0])


def build(source: Path, output: Path, status_path: Path, test_count_override: str | None = None):
    doc = Document(str(source))
    trim_after_cover(doc)
    configure_styles(doc)
    compact_cover(doc)
    set_update_fields(doc)
    test_count = test_count_override or infer_test_count(status_path)
    docs_dir = source.parent

    write_front_matter(doc, test_count)
    write_chapter_1(doc)
    write_chapter_2(doc)
    write_chapter_3(doc, docs_dir)
    write_chapter_4(doc, docs_dir)
    write_chapter_5(doc, test_count)
    write_chapter_6(doc)
    write_chapter_7(doc, test_count)
    write_appendices(doc, test_count)

    doc.core_properties.title = "Báo cáo Dự án Tốt nghiệp – Hệ thống KPI/OKR AI-native"
    doc.core_properties.subject = "Phiên bản bàn giao hoàn thiện ngày 11/08/2026"
    doc.core_properties.author = "Nhóm NEXTGEN"
    doc.core_properties.keywords = "KPI, OKR, AI-native, RAG, ASP.NET Core, SQL Server"
    doc.core_properties.comments = "Dựng từ file nguồn; chỉ mô tả phạm vi đã hoàn tất và được kiểm chứng."
    doc.core_properties.modified = datetime(2026, 8, 11, 6, 30, 0)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path(__file__).resolve().parents[1] / "BaoCao_DuAn_TotNghiep.docx",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).resolve().parents[1] / "BaoCao_DuAn_TotNghiep_HoanThien_20260811.docx",
    )
    parser.add_argument(
        "--status",
        type=Path,
        default=Path(__file__).resolve().parents[1] / "AI_NATIVE_PLAN_STATUS.md",
    )
    parser.add_argument(
        "--test-count",
        default="536",
        help="Verified full-suite count embedded in the submission report.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build(
        args.source.resolve(),
        args.output.resolve(),
        args.status.resolve(),
        args.test_count,
    )
    print(args.output.resolve())
