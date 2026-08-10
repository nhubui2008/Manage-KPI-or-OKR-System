"""
Script tạo phần 2.3: Use case & Activity Diagrams cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_spec_table(doc, data_dict):
    """Tạo bảng đặc tả Use Case chi tiết"""
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, (key, val) in enumerate(data_dict.items()):
        # Cột 1 (Tên thuộc tính)
        cell_key = table.rows[i].cells[0]
        set_cell(cell_key, key, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=RGBColor(255, 255, 255))
        set_cell_shading(cell_key, HEADER_BG)
        cell_key.width = Cm(3.5)

        # Cột 2 (Giá trị thuộc tính)
        cell_val = table.rows[i].cells[1]
        set_cell(cell_val, val, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_val.width = Cm(12.5)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 2.3 =====================

def write_section_2_3(doc):
    """2.3. Use case"""

    add_heading1(doc, "2.3. Use case")

    add_para(doc,
        "Để làm rõ cách thức tương tác giữa các tác nhân và hệ thống, mục này trình bày các sơ đồ "
        "Use Case tổng quan, danh sách Use Case chi tiết, đặc tả các Use Case cốt lõi và sơ đồ hoạt động "
        "(Activity Diagrams) mô tả luồng quy trình vận hành chính của hệ thống."
    )

    # ============================================================
    # 2.3.1. SƠ ĐỒ USE CASE VÀ DANH SÁCH USE CASE
    # ============================================================
    add_heading2(doc, "2.3.1. Sơ đồ Use Case tổng quan")
    
    add_para(doc,
        "Sơ đồ Use Case tổng quan của hệ thống mô tả mối quan hệ giữa 5 tác nhân (Admin, Director, HR, "
        "Manager, Employee) và các nhóm chức năng cốt lõi được trình bày chi tiết tại Hình 1 (Mục 2.1). "
        "Đồng thời, danh sách 28 Use Case chi tiết cùng mô tả tóm tắt được trình bày tại Bảng 9 (Mục 2.1) "
        "là cơ sở phân rã để thực hiện thiết kế chi tiết dưới đây."
    )

    # ============================================================
    # 2.3.2. ĐẶC TẢ CÁC USE CASE CỐT LÕI
    # ============================================================
    add_heading2(doc, "2.3.2. Đặc tả các Use Case cốt lõi")

    add_para(doc,
        "Hệ thống bao gồm nhiều chức năng, dưới đây là đặc tả chi tiết của 4 Use Case cốt lõi "
        "đóng vai trò xương sống cho hoạt động quản lý hiệu suất và tích hợp AI của dự án:"
    )

    # --- UC 1: Giao KPI cho nhân viên ---
    add_heading3(doc, "a) Đặc tả Use Case UC_MA_02 – Giao KPI nhân viên")
    
    spec_uc_ma_02 = {
        "Mã Use Case": "UC_MA_02",
        "Tên Use Case": "Giao KPI cho nhân viên",
        "Độ ưu tiên": "Cao (Trọng yếu)",
        "Tác nhân chính": "Manager (Trưởng phòng ban)",
        "Mô tả": "Cho phép Trưởng phòng ban tạo chỉ tiêu KPI cho nhân viên trực thuộc trong kỳ đánh giá. KPI bao gồm target cụ thể, trọng số đóng góp, kỳ hạn check-in và liên kết trực tiếp với mục tiêu OKR/Key Result tương ứng.",
        "Điều kiện tiên quyết": "Kỳ đánh giá đang hoạt động (Mở); Trưởng phòng và Nhân viên đã được HR gán vào cùng phòng ban; OKR phòng ban đã được thiết lập.",
        "Luồng sự kiện chính (Basic Flow)": 
            "1. Trưởng phòng truy cập vào module 'Quản lý KPI', chọn 'Tạo mới KPI'.\n"
            "2. Hệ thống hiển thị biểu mẫu yêu cầu nhập: Tên KPI, loại KPI (Định lượng/Định tính/Hành vi), thuộc tính (Tăng trưởng/Giảm thiểu).\n"
            "3. Trưởng phòng nhập các tham số: Giá trị Target, Pass/Fail Threshold, đơn vị đo, tần suất check-in và deadline.\n"
            "4. Trưởng phòng gán KPI cho nhân viên cụ thể, thiết lập Trọng số (Weight) và chọn liên kết với một Key Result cụ thể của OKR phòng ban.\n"
            "5. Trưởng phòng bấm 'Gửi duyệt'. Hệ thống kiểm tra tổng trọng số KPI của nhân viên (không vượt quá 100%), lưu trạng thái là 'Chờ duyệt' và gửi email thông báo cho Director.",
        "Luồng phụ / Ngoại lệ (Alternative Flow)":
            "Luồng 4a (Chọn gợi ý từ AI):\n"
            "   - Trưởng phòng click vào nút 'Gợi ý từ AI'.\n"
            "   - Server xác minh quyền, kỳ đang mở và OKR/Key Result rồi gửi snapshot tối thiểu không chứa tên, mã hoặc email nhân viên qua model gateway.\n"
            "   - Advisor trả 3-5 bản nháp strict JSON có citation; server kiểm tra đơn vị, chiều KPI và target/pass/fail. Trưởng phòng chọn một bản nháp để điền form rồi vẫn phải tự gửi qua validator chuẩn.\n"
            "Luồng 5a (Vượt quá tổng trọng số):\n"
            "   - Hệ thống hiển thị thông báo lỗi 'Tổng trọng số KPI của nhân viên trong kỳ vượt quá 100%'. Yêu cầu Trưởng phòng điều chỉnh lại.",
        "Điều kiện sau (Post-condition)": "KPI được lưu vào bảng KPIs và KPIDetails với trạng thái 'Chờ duyệt' (Pending Decision)."
    }
    create_spec_table(doc, spec_uc_ma_02)
    add_table_caption(doc, "Bảng 11: Đặc tả Use Case UC_MA_02 – Giao KPI nhân viên")

    # --- UC 2: Check-in tiến độ KPI ---
    add_heading3(doc, "b) Đặc tả Use Case UC_EM_02 – Check-in tiến độ KPI")
    
    spec_uc_em_02 = {
        "Mã Use Case": "UC_EM_02",
        "Tên Use Case": "Check-in tiến độ KPI",
        "Độ ưu tiên": "Cao (Trọng yếu)",
        "Tác nhân chính": "Employee (Nhân viên)",
        "Mô tả": "Cho phép nhân viên báo cáo kết quả thực hiện KPI định kỳ theo lịch check-in đã giao, gửi kèm ghi chú giải trình tiến độ để gửi lên cấp quản lý phê duyệt.",
        "Điều kiện tiên quyết": "KPI của nhân viên ở trạng thái 'Đang thực hiện' (Active); Đến lịch check-in định kỳ hoặc nhân viên chủ động cập nhật.",
        "Luồng sự kiện chính (Basic Flow)": 
            "1. Nhân viên truy cập trang cá nhân, chọn KPI cần check-in và click 'Báo cáo tiến độ'.\n"
            "2. Hệ thống hiển thị form check-in gồm: Giá trị đạt được thực tế (Achieved Value) và Ghi chú/Ý kiến giải trình.\n"
            "3. Nhân viên nhập giá trị thực tế mới nhất và viết nội dung giải trình tiến độ thực hiện.\n"
            "4. Nhân viên bấm 'Gửi báo cáo'.\n"
            "5. Hệ thống ghi nhận giá trị, tự động tính toán % hoàn thành thực tế và so sánh với tiến độ kỳ vọng tại thời điểm hiện tại (schedule progress).\n"
            "6. Hệ thống tạo bản ghi trong bảng KPICheckIns với trạng thái 'Chờ duyệt' (Pending Review), đẩy bản ghi vào Review Queue của Trưởng phòng và gửi thông báo nhắc nhở.",
        "Luồng phụ / Ngoại lệ (Alternative Flow)":
            "Luồng 3a (Nhân viên đính kèm link tài liệu minh chứng):\n"
            "   - Nhân viên nhập link tài liệu chứng minh kết quả check-in trong phần ghi chú.\n"
            "Luồng 5a (Check-in muộn quá hạn):\n"
            "   - Hệ thống tự động ghi nhận thuộc tính 'Quá hạn check-in' và gửi cảnh báo rủi ro (Smart Alerts) lên cho Trưởng phòng.",
        "Điều kiện sau (Post-condition)": "Tiến độ check-in được lưu tạm thời, chờ Trưởng phòng duyệt duyệt trong Review Queue."
    }
    create_spec_table(doc, spec_uc_em_02)
    add_table_caption(doc, "Bảng 12: Đặc tả Use Case UC_EM_02 – Check-in tiến độ KPI")

    # --- UC 3: Phê duyệt Check-in ---
    add_heading3(doc, "c) Đặc tả Use Case UC_MA_03 – Phê duyệt Check-in")
    
    spec_uc_ma_03 = {
        "Mã Use Case": "UC_MA_03",
        "Tên Use Case": "Phê duyệt Check-in",
        "Độ ưu tiên": "Cao",
        "Tác nhân chính": "Manager (Trưởng phòng ban)",
        "Mô tả": "Cho phép Trưởng phòng phê duyệt hoặc từ chối kết quả check-in tiến độ của nhân viên trực thuộc từ hàng đợi Review Queue, ghi nhận nhận xét và chấm điểm.",
        "Điều kiện tiên quyết": "Nhân viên đã thực hiện check-in và gửi báo cáo tiến độ (ở trạng thái Pending Review).",
        "Luồng sự kiện chính (Basic Flow)": 
            "1. Trưởng phòng truy cập vào 'Review Queue' của phòng ban quản lý.\n"
            "2. Hệ thống hiển thị danh sách các bản ghi check-in đang chờ duyệt của nhân viên.\n"
            "3. Trưởng phòng click xem chi tiết một bản ghi, xem xét giá trị báo cáo, so sánh biểu đồ tiến độ thực tế vs tiến độ kỳ vọng và đọc giải trình.\n"
            "4. Trưởng phòng viết nhận xét đánh giá và chọn nút 'Duyệt' (Approve).\n"
            "5. Hệ thống gọi OKRProgressService để tự động cập nhật tiến độ phần trăm (%) của KPI và đồng bộ hóa tiến độ lên các Key Results và OKR phòng ban/công ty liên quan.\n"
            "6. Trạng thái bản ghi check-in chuyển thành 'Đã duyệt' (Approved) và hệ thống gửi thông báo kết quả cho nhân viên.",
        "Luồng phụ / Ngoại lệ (Alternative Flow)":
            "Luồng 4a (Trưởng phòng chọn 'Từ chối' - Reject):\n"
            "   - Trưởng phòng bắt buộc phải nhập lý do từ chối (Fail Reason) và viết nhận xét yêu cầu nhân viên làm rõ.\n"
            "   - Trưởng phòng click 'Từ chối'. Hệ thống chuyển trạng thái check-in thành 'Từ chối' (Rejected) và gửi thông báo yêu cầu nhân viên cập nhật lại.",
        "Điều kiện sau (Post-condition)": "Tiến độ KPI được cập nhật chính thức vào cơ sở dữ liệu nếu được duyệt."
    }
    create_spec_table(doc, spec_uc_ma_03)
    add_table_caption(doc, "Bảng 13: Đặc tả Use Case UC_MA_03 – Phê duyệt Check-in")

    # --- UC 4: Trợ lý AI có nguồn ---
    add_heading3(doc, "d) Đặc tả Use Case UC_EM_06 – Trợ lý AI có nguồn")
    
    spec_uc_em_06 = {
        "Mã Use Case": "UC_EM_06",
        "Tên Use Case": "Tương tác với Trợ lý AI có nguồn",
        "Độ ưu tiên": "Trung bình (Tính năng giá trị gia tăng)",
        "Tác nhân chính": "Employee, Manager, Director (Người dùng hệ thống)",
        "Mô tả": "Cung cấp Chat Advisor nhận biết ngữ cảnh được cấp quyền, dùng nguồn SQL/RAG có citation và chủ động abstain khi bằng chứng không đủ.",
        "Điều kiện tiên quyết": "Tài khoản có membership/role tenant hiện hành, gói dịch vụ hỗ trợ AI và model gateway/RAG đã được cấu hình an toàn.",
        "Luồng sự kiện chính (Basic Flow)": 
            "1. Người dùng click mở 'Chatbot AI' từ thanh công cụ hoặc trang cá nhân.\n"
            "2. Người dùng nhập câu hỏi hoặc yêu cầu (Ví dụ: 'Phòng Công nghệ của tôi đang chậm KPI X, làm sao để khắc phục?').\n"
            "3. Server dựng lại principal từ membership/role hiện hành, nạp KPI/OKR và check-in đã duyệt trong đúng scope.\n"
            "4. Hệ thống truy xuất tối đa các nguồn RAG mà tenant/ACL cho phép và gửi context tạm thời qua IAIModelClient.\n"
            "5. Advisor kiểm tra strict JSON, source ID và trạng thái nguồn trước khi trả câu trả lời có citation.\n"
            "6. Widget hiển thị câu trả lời tư vấn và nguồn; server chỉ lưu AgentRun/citation metadata, không lưu nội dung hội thoại hoặc raw provider response.",
        "Luồng phụ / Ngoại lệ (Alternative Flow)":
            "Luồng 4a (Thiếu bằng chứng): Advisor abstain và nêu rõ không đủ nguồn để trả lời.\n"
            "Luồng 4b (Nguồn hoặc quyền thay đổi): Server trả conflict, không sử dụng kết quả stale và yêu cầu thử lại.\n"
            "Luồng 4c (Provider timeout/invalid JSON): Hệ thống retry hữu hạn rồi trả thông báo an toàn, không lộ lỗi nội bộ.",
        "Điều kiện sau (Post-condition)": "Không có dữ liệu nghiệp vụ chính thức nào bị thay đổi; chỉ metadata run/citation tối thiểu được lưu để kiểm toán."
    }
    create_spec_table(doc, spec_uc_em_06)
    add_table_caption(doc, "Bảng 14: Đặc tả Use Case UC_EM_06 – Trợ lý AI có nguồn")

    # ============================================================
    # 2.3.3. SƠ ĐỒ HOẠT ĐỘNG (ACTIVITY DIAGRAMS)
    # ============================================================
    add_heading2(doc, "2.3.3. Sơ đồ hoạt động cho các luồng quy trình chính")

    add_para(doc,
        "Sơ đồ hoạt động (Activity Diagram) thể hiện trình tự luồng xử lý nghiệp vụ đi qua các làn tác nhân "
        "(Swimlanes) khác nhau. Dưới đây là 3 sơ đồ hoạt động cho các quy trình cốt lõi của hệ thống:"
    )

    # --- Sơ đồ 1 ---
    add_heading3(doc, "a) Quy trình thiết lập & giao KPI/OKR đa cấp")
    add_para(doc,
        "Sơ đồ dưới đây mô tả chuỗi hoạt động từ lúc Ban Giám đốc thiết lập chiến lược công ty, phân bổ OKR "
        "xuống phòng ban, Trưởng phòng tạo và giao KPI cụ thể cho nhân viên, cho đến khi KPI được duyệt và "
        "kích hoạt chính thức:"
    )
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    p_act1 = doc.add_paragraph()
    p_act1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path1 = os.path.join(script_dir, "activity_kpi_okr.png")
    if os.path.exists(img_path1):
        p_act1.add_run().add_picture(img_path1, width=Cm(14.5))
    else:
        p_act1.add_run("[SƠ ĐỒ HOẠT ĐỘNG GIAO KPI/OKR - LỖI HÌNH ẢNH]")
    add_table_caption(doc, "Hình 2: Sơ đồ hoạt động quy trình thiết lập & giao KPI/OKR đa cấp")

    # --- Sơ đồ 2 ---
    add_heading3(doc, "b) Quy trình check-in và phê duyệt tiến độ")
    add_para(doc,
        "Sơ đồ dưới đây mô tả luồng báo cáo tiến độ (check-in) của nhân viên, hệ thống tự động tính toán %, "
        "đẩy vào Review Queue của Trưởng phòng để duyệt hoặc từ chối, và tự động đồng bộ hóa ngược lên OKR liên quan:"
    )
    
    p_act2 = doc.add_paragraph()
    p_act2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path2 = os.path.join(script_dir, "activity_checkin.png")
    if os.path.exists(img_path2):
        p_act2.add_run().add_picture(img_path2, width=Cm(14.5))
    else:
        p_act2.add_run("[SƠ ĐỒ HOẠT ĐỘNG CHECK-IN - LỖI HÌNH ẢNH]")
    add_table_caption(doc, "Hình 3: Sơ đồ hoạt động quy trình check-in và phê duyệt tiến độ")

    # --- Sơ đồ 3 ---
    add_heading3(doc, "c) Quy trình trợ lý AI có nguồn tư vấn & hỗ trợ quyết định")
    add_para(doc,
        "Sơ đồ dưới đây mô tả luồng Chat Advisor: xác thực tenant/scope, thu thập nguồn SQL/RAG được cấp quyền, "
        "gọi model gateway, kiểm tra strict schema/citation và chỉ trả kết quả khi nguồn còn hiện hành:"
    )
    
    p_act3 = doc.add_paragraph()
    p_act3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path3 = os.path.join(script_dir, "activity_ai.png")
    if os.path.exists(img_path3):
        p_act3.add_run().add_picture(img_path3, width=Cm(14.5))
    else:
        p_act3.add_run("[SƠ ĐỒ HOẠT ĐỘNG CHAT ADVISOR - LỖI HÌNH ẢNH]")
    add_table_caption(doc, "Hình 4: Sơ đồ hoạt động quy trình trợ lý AI có nguồn")

    # Kết luận mục
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Tóm lại, việc phân tích chi tiết sơ đồ Use Case và các sơ đồ hoạt động giúp nhóm nắm rõ luồng xử lý "
        "nghiệp vụ và cơ chế tương tác đa bên trong hệ thống. Đây là cơ sở cốt lõi để nhóm tiến hành "
        "thiết kế cấu trúc dữ liệu quan hệ thực thể (ERD) ở phần tiếp theo.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 2.3. Use case")
    print("      ✓ 2.3.1. Sơ đồ và danh sách Use Case")
    print("      ✓ 2.3.2. Đặc tả 4 Use Case cốt lõi (Bảng 11 - 14)")
    print("      ✓ 2.3.3. Sơ đồ hoạt động (Hình 2 - 4)")
    write_section_2_3(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 2.3 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
