"""
Script tạo phần 3.2: Thiết kế giao diện cho báo cáo tốt nghiệp
Chèn sơ đồ sitemap và 6 hình ảnh chụp màn hình thực tế từ hệ thống đang chạy
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 3.2 =====================

def write_section_3_2(doc):
    """3.2. Thiết kế giao diện"""

    add_heading1(doc, "3.2. Thiết kế giao diện")

    add_para(doc,
        "Thiết kế giao diện người dùng (UI/UX) quyết định mức độ tiện dụng và khả năng áp dụng thực tế "
        "của hệ thống trong doanh nghiệp. Giao diện được thiết kế theo tiêu chí tối giản, trực quan, hỗ trợ "
        "chế độ Dark Mode hiện đại và Responsive hoàn toàn tương thích với các thiết bị di động. "
        "Dưới đây là chi tiết sơ đồ tổ chức, danh sách các màn hình và đặc tả hình ảnh giao diện thực tế của hệ thống."
    )

    # ============================================================
    # 3.2.1. SƠ ĐỒ TỔ CHỨC GIAO DIỆN (SITE MAP)
    # ============================================================
    add_heading2(doc, "3.2.1. Sơ đồ tổ chức giao diện (Site Map)")
    
    add_para(doc,
        "Sơ đồ phân cấp cấu trúc trang hiển thị mối liên kết chặt chẽ từ cổng đăng nhập Auth, "
        "đến màn hình Dashboard trung tâm và 5 trục chức năng chính cùng với widget Trợ lý AI có nguồn:"
    )

    script_dir = os.path.dirname(os.path.abspath(__file__))
    sitemap_path = os.path.join(script_dir, "site_map.png")

    p_sitemap = doc.add_paragraph()
    p_sitemap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(sitemap_path):
        p_sitemap.add_run().add_picture(sitemap_path, width=Cm(15.0))
    else:
        p_sitemap.add_run("[SƠ ĐỒ SITEMAP GIAO DIỆN - LỖI HÌNH ẢNH]")
    add_figure_caption(doc, "Hình 6: Sơ đồ cấu trúc phân cấp giao diện hệ thống")

    # ============================================================
    # 3.2.2. DANH SÁCH CÁC MÀN HÌNH CHÍNH
    # ============================================================
    add_heading2(doc, "3.2.2. Danh sách các màn hình chính")

    add_para(doc, "Hệ thống bao gồm các màn hình chính được phân nhóm theo nhiệm vụ sử dụng dưới đây:")

    # Bảng 27: Danh sách màn hình
    headers_screens = ["STT", "Tên màn hình", "Đường dẫn (URL)", "Đối tượng sử dụng", "Mô tả chức năng chính"]
    rows_screens = [
        ["1", "Cổng Đăng Nhập", "/Auth/Login", "Tất cả người dùng", "Nhập tài khoản hoặc click đăng nhập bằng Google OAuth, lựa chọn Demo switcher."],
        ["2", "Dashboard Giám Đốc", "/Dashboard", "Director / Giám đốc", "Hiển thị KPI công ty, OKR chung, nút chạy AI Performance Analysis, Expected Bonus."],
        ["3", "Quản lý OKR Chiến lược", "/OKRs", "Director, Manager", "Xem biểu đồ OKR 3 cấp, tiến độ Key Results dưới dạng thanh tiến độ trực quan."],
        ["4", "Quản lý chỉ tiêu KPI", "/KPIs", "Manager, Employee", "Danh sách các KPI, gán trọng số, trạng thái (Duyệt/Thực hiện) và nút gọi AI gợi ý KPI."],
        ["5", "Review Queue check-in", "/KPICheckIns", "Manager / Trưởng phòng", "Hàng đợi duyệt tiến độ báo cáo của nhân viên trực thuộc phòng ban."],
        ["6", "Check-in tiến độ cá nhân", "/KPICheckIns/Create", "Employee / Nhân viên", "Biểu mẫu nhập kết quả thực tế, ghi chú và giải trình cho lần check-in."],
        ["7", "Dự án & Kanban", "/WorkProjects", "Manager, Employee", "Bảng Kanban board kéo thả (Backlog->Done) để quản lý các đầu việc chi tiết."],
        ["8", "Lên lịch họp 1-on-1", "/OneOnOneMeetings", "Manager, Employee", "Lên lịch họp, ghi nhận biên bản trao đổi và các kế hoạch hành động."],
        ["9", "Đặc tả lương thưởng HR", "/EvaluationResults", "HR / Nhân sự", "Cấu hình quy tắc tính thưởng theo bậc Rank và xuất báo cáo Excel."],
        ["10", "Widget Trợ lý AI", "Floating Widget", "Tất cả người dùng", "Widget trượt mở ở góc phải hỗ trợ chat tư vấn, nhận diện context và Smart Alerts."],
    ]

    create_table(doc, headers_screens, rows_screens, col_widths=[1.0, 3.5, 2.5, 3.0, 6.0])
    add_table_caption(doc, "Bảng 27: Danh sách các màn hình giao diện chính của hệ thống")

    # ============================================================
    # 3.2.3. THIẾT KẾ CHI TIẾT CÁC MÀN HÌNH CHÍNH (SCREENS)
    # ============================================================
    add_heading2(doc, "3.2.3. Thiết kế chi tiết các màn hình chính (Screenshots thực tế)")

    add_para(doc,
        "Dưới đây là hình ảnh giao diện thực tế của hệ thống được chụp từ ứng dụng đang vận hành tại cổng localhost:5208, "
        "thể hiện độ trực quan, nhất quán và thẩm mỹ cao:"
    )

    # 1. Cổng đăng nhập
    add_heading3(doc, "a) Giao diện Cổng đăng nhập (Auth/Login)")
    add_para(doc,
        "Màn hình đăng nhập được thiết kế theo phong cách tối giản và hiện đại. Hỗ trợ biểu mẫu đăng nhập "
        "an toàn, ghi nhớ mật khẩu, nút liên kết Google OAuth và phần 'Demo Accounts' giúp kiểm thử nhanh "
        "với các vai trò (admin, director, manager, hr, employee) chỉ bằng một click chuột:"
    )
    
    img_login = os.path.join(script_dir, "screenshot_login.png")
    p_img_login = doc.add_paragraph()
    p_img_login.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_login):
        p_img_login.add_run().add_picture(img_login, width=Cm(14.5))
    else:
        p_img_login.add_run("[HÌNH ẢNH LOGIN LỖI]")
    add_figure_caption(doc, "Hình 7: Giao diện cổng đăng nhập hệ thống")

    # 2. Dashboard Giám đốc
    add_heading3(doc, "b) Giao diện Dashboard của Giám đốc (Director Dashboard)")
    add_para(doc,
        "Màn hình Dashboard trung tâm dành cho Giám đốc hiển thị các chỉ số vĩ mô:achievement rate, "
        "tiến độ OKR công ty (79.7%), số nhân viên active. Đặc biệt, tích hợp hai nút tương tác AI: "
        "'AI Performance Analysis' (Phân tích hiệu suất bằng AI) và 'AI Customer Segments' (Phân khúc khách hàng bằng AI):"
    )
    
    img_dash = os.path.join(script_dir, "screenshot_dashboard.png")
    p_img_dash = doc.add_paragraph()
    p_img_dash.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_dash):
        p_img_dash.add_run().add_picture(img_dash, width=Cm(14.5))
    else:
        p_img_dash.add_run("[HÌNH ẢNH DASHBOARD LỖI]")
    add_figure_caption(doc, "Hình 8: Giao diện Dashboard tổng quan dành cho Giám đốc")

    # 3. Quản lý OKR
    add_heading3(doc, "c) Giao diện Quản lý OKR Chiến lược (OKRs)")
    add_para(doc,
        "Màn hình quản lý OKR hiển thị sơ đồ phân rã mục tiêu chiến lược Q2-2026. Mỗi mục tiêu OKR "
        "đều có thanh tiến độ (Progress Bar) được tính toán tự động từ các kết quả then chốt (Key Results) "
        "trực thuộc, hỗ trợ tìm kiếm nhanh và lọc theo kỳ đánh giá:"
    )
    
    img_okrs = os.path.join(script_dir, "screenshot_okrs.png")
    p_img_okrs = doc.add_paragraph()
    p_img_okrs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_okrs):
        p_img_okrs.add_run().add_picture(img_okrs, width=Cm(14.5))
    else:
        p_img_okrs.add_run("[HÌNH ẢNH OKRS LỖI]")
    add_figure_caption(doc, "Hình 9: Giao diện quản lý OKR và thanh tiến độ tự động")

    # 4. Quản lý KPI
    add_heading3(doc, "d) Giao diện Danh sách chỉ tiêu KPI (KPIs)")
    add_para(doc,
        "Giao diện quản lý danh sách KPI liệt kê tất cả các chỉ tiêu hiệu suất được giao trong kỳ. "
        "Màn hình hiển thị đầy đủ thông tin: tên KPI, loại chỉ tiêu, người được giao, người giao, trạng thái phê duyệt "
        "và đặc biệt là nút gọi tính năng 'AI KPI Generator' giúp tự động sinh KPI thông minh từ mục tiêu phòng ban:"
    )
    
    img_kpis = os.path.join(script_dir, "screenshot_kpis.png")
    p_img_kpis = doc.add_paragraph()
    p_img_kpis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_kpis):
        p_img_kpis.add_run().add_picture(img_kpis, width=Cm(14.5))
    else:
        p_img_kpis.add_run("[HÌNH ẢNH KPIS LỖI]")
    add_figure_caption(doc, "Hình 10: Giao diện quản lý và giao chỉ tiêu KPI")

    # 5. Check-in tiến độ
    add_heading3(doc, "e) Giao diện Nhật ký Check-in tiến độ nhân viên (KPICheckIns)")
    add_para(doc,
        "Màn hình hiển thị lịch sử check-in tiến độ của nhân viên. Hiển thị rõ giá trị thực tế đạt được "
        "(Achieved Value), các rào cản/khó khăn gặp phải (Barriers), nhận xét phản hồi từ quản lý và điểm đánh giá "
        "lần check-in tương ứng, đảm bảo tính minh bạch tối đa trong quá trình theo dõi:"
    )
    
    img_checkin = os.path.join(script_dir, "screenshot_checkin.png")
    p_img_checkin = doc.add_paragraph()
    p_img_checkin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_checkin):
        p_img_checkin.add_run().add_picture(img_checkin, width=Cm(14.5))
    else:
        p_img_checkin.add_run("[HÌNH ẢNH CHECKIN LỖI]")
    add_figure_caption(doc, "Hình 11: Giao diện nhật ký check-in tiến độ và nhận xét của quản lý")

    # 6. Trợ lý AI có nguồn
    add_heading3(doc, "f) Giao diện Trợ lý AI có nguồn (Bizen AI Assistant)")
    add_para(doc,
        "Bảng điều khiển slide-out của Chat Advisor (Bizen AI Widget) hiển thị ở góc phải màn hình, "
        "tự động liên kết ngữ cảnh thực tế của người dùng. Cung cấp các nút tắt nhanh: 'Phân tích tiến độ KPI', "
        "'Tìm KPI có rủi ro', 'Đề xuất giải pháp khắc phục', đồng thời cho phép nhập câu hỏi tự nhiên để AI hỗ trợ:"
    )
    
    img_ai = os.path.join(script_dir, "screenshot_ai.png")
    p_img_ai = doc.add_paragraph()
    p_img_ai.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_ai):
        p_img_ai.add_run().add_picture(img_ai, width=Cm(14.5))
    else:
        p_img_ai.add_run("[HÌNH ẢNH AI LỖI]")
    add_figure_caption(doc, "Hình 12: Giao diện Trợ lý AI có nguồn (Bizen AI Assistant) trượt mở")

    # Kết luận mục
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Tổng kết Chương 3: Việc thiết kế cơ sở dữ liệu chuẩn hóa 3NF và xây dựng hệ thống giao diện trực quan, "
        "nhất quán có sự hỗ trợ của các advisor AI có kiểm soát đã hoàn thiện toàn bộ phần thiết kế logic và thiết kế vật lý của dự án. "
        "Đây là cơ sở đầy đủ và chi tiết để nhóm NEXTGEN triển khai lập trình mã nguồn và cấu hình hệ thống ở chương tiếp theo.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 3.2. Thiết kế giao diện")
    print("      ✓ 3.2.1. Sơ đồ cấu trúc giao diện Site Map (Hình 6)")
    print("      ✓ 3.2.2. Danh sách 10 màn hình chính (Bảng 27)")
    print("      ✓ 3.2.3. Đặc tả 6 ảnh chụp màn hình thực tế (Hình 7 - 12)")
    write_section_3_2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 3.2 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
