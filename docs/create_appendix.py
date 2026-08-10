"""
Script tạo PHẦN PHỤ LỤC (Phụ lục A - E) cho báo cáo tốt nghiệp
Dựa trên cấu trúc chuẩn của báo cáo tốt nghiệp FPT Polytechnic
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_appendix_title(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def create_spec_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0, 2, 4]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG PHỤ LỤC =====================

def write_appendices(doc):
    """Viết Phụ lục A, B, C, D, E vào file Word"""

    # ============================================================
    # PHỤ LỤC A: ĐẶC TẢ USE CASE BỔ SUNG
    # ============================================================
    add_page_break(doc)
    add_appendix_title(doc, "PHỤ LỤC A: ĐẶC TẢ USE CASE BỔ SUNG")
    
    add_para(doc,
        "Để bổ sung cho mục 2.3 Chương 2, phụ lục này đặc tả chi tiết 2 Use Cases nghiệp vụ quan trọng khác "
        "là UC10: Quản lý hồ sơ nhân sự (CRUD) và UC25: Tra cứu lịch sử cuộc gọi AI."
    )

    # Bảng đặc tả UC10
    add_heading2(doc, "1. Đặc tả Use Case UC10: Quản lý hồ sơ nhân sự")
    headers_uc10 = ["Thuộc tính Use Case", "Nội dung đặc tả chi tiết"]
    rows_uc10 = [
        ["Tên Use Case", "UC10: Quản lý hồ sơ nhân sự (CRUD)"],
        ["Tác nhân (Actor)", "HR (Nhân sự), Admin (Quản trị viên)"],
        ["Mục đích", "Cho phép HR thêm mới, cập nhật, khóa hồ sơ nhân sự và import hàng loạt nhân sự từ tệp Excel."],
        ["Điều kiện tiên quyết", "HR đã đăng nhập thành công và có quyền 'EMPLOYEE_EDIT'."],
        ["Kịch bản chính (Luồng cơ bản)",
         "1. HR vào mục 'Nhân sự' -> Chọn 'Tạo hồ sơ nhân viên'.\n"
         "2. Nhập thông tin: Họ tên, Email công ty, Phone, Phòng ban, Chức vụ và Ngày vào làm.\n"
         "3. Click 'Lưu hồ sơ'.\n"
         "4. Hệ thống kiểm tra trùng lặp Email và tự động sinh mã nhân viên EMPxxx.\n"
         "5. Hệ thống lưu database và tạo tài khoản SystemUser tương ứng, gửi email password tạm thời."],
        ["Luồng rẽ nhánh (Alternative Flows)",
         "2a. HR chọn 'Import Excel':\n"
         "  1. Tải tệp Excel mẫu.\n"
         "  2. Điền thông tin hàng loạt nhân viên và upload tệp.\n"
         "  3. Hệ thống validate định dạng và import toàn bộ bản ghi hợp lệ."],
        ["Kịch bản lỗi (Exception Flows)",
         "4a. Email bị trùng lặp: Hệ thống báo lỗi 'Email đã tồn tại' và giữ nguyên form để sửa."],
        ["Kết quả đầu ra", "Hồ sơ nhân viên được lưu vào bảng Employees; tài khoản đăng nhập được kích hoạt."],
    ]
    create_table(doc, headers_uc10, rows_uc10, col_widths=[4.5, 11.5])
    add_table_caption(doc, "Bảng A.1: Đặc tả chi tiết Use Case UC10: Quản lý hồ sơ nhân sự")

    # Bảng đặc tả UC25
    add_heading2(doc, "2. Đặc tả Use Case UC25: Theo dõi vận hành AI")
    headers_uc25 = ["Thuộc tính Use Case", "Nội dung đặc tả chi tiết"]
    rows_uc25 = [
        ["Tên Use Case", "UC25: Theo dõi vận hành AI an toàn"],
        ["Tác nhân (Actor)", "Admin, Director"],
        ["Mục đích", "Giúp quản trị viên theo dõi trạng thái run, độ trễ, retry/dead-letter, citation và abstain mà không đọc prompt hoặc raw provider response."],
        ["Điều kiện tiên quyết", "Admin đã đăng nhập thành công và có quyền 'AUDIT_LOGS_VIEW'."],
        ["Kịch bản chính (Luồng cơ bản)",
         "1. Quản trị tenant mở trang vận hành AI/RAG được phân quyền.\n"
         "2. Hệ thống hiển thị trạng thái run/outbox/ingestion, độ trễ, số lần retry, dead-letter, citation và abstain.\n"
         "3. Quản trị viên có thể retry đúng bản ghi DeadLetter sau khi server kiểm tra row-version, source và quyền hiện hành.\n"
         "4. Giao diện không hiển thị prompt, nội dung tài liệu, hội thoại hoặc raw provider response."],
        ["Kết quả đầu ra", "Metadata vận hành được truy vấn theo tenant; mọi retry được audit và không tự thay đổi dữ liệu nghiệp vụ chính thức."],
    ]
    create_table(doc, headers_uc25, rows_uc25, col_widths=[4.5, 11.5])
    add_table_caption(doc, "Bảng A.2: Đặc tả chi tiết Use Case UC25: Theo dõi vận hành AI")


    # ============================================================
    # PHỤ LỤC B: ĐẶC TẢ BẢNG CƠ SỞ DỮ LIỆU BỔ SUNG
    # ============================================================
    add_page_break(doc)
    add_appendix_title(doc, "PHỤ LỤC B: ĐẶC TẢ BẢNG CƠ SỞ DỮ LIỆU BỔ SUNG")
    
    add_para(doc,
        "Phụ lục này đặc tả chi tiết 4 bảng dữ liệu phụ trợ quan trọng trong hệ thống bao gồm: "
        "bảng dự án Kanban (WorkProjects), bảng thẻ công việc (WorkItems), bảng quy tắc thưởng (BonusRules) "
        "và bảng lịch sử AI (AIGenerationHistories)."
    )

    headers_spec = ["STT", "Tên trường", "Kiểu dữ liệu", "Ý nghĩa / Mô tả", "Ràng buộc"]
    col_w = [1.0, 3.5, 3.0, 6.0, 2.5]

    # --- 1. WorkProjects ---
    add_heading2(doc, "1. Đặc tả bảng WorkProjects (Dự án phòng ban)")
    rows_proj = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh dự án", "PK, IDENTITY(1,1)"],
        ["2", "ProjectName", "NVARCHAR(255)", "Tên dự án Kanban", "NOT NULL"],
        ["3", "Description", "NVARCHAR(1000)", "Mô tả mục tiêu dự án", "NULL"],
        ["4", "DepartmentId", "INT", "Mã phòng ban quản lý dự án", "FK, NOT NULL"],
        ["5", "SourceOKRId", "INT", "OKR nguồn mà dự án đóng góp tiến độ", "FK, NULL"],
        ["6", "IsActive", "BIT", "Trạng thái hoạt động", "DEFAULT 1"],
    ]
    create_spec_table(doc, headers_spec, rows_proj, col_widths=col_w)
    add_table_caption(doc, "Bảng B.1: Đặc tả chi tiết bảng WorkProjects")

    # --- 2. WorkItems ---
    add_heading2(doc, "2. Đặc tả bảng WorkItems (Thẻ công việc Kanban)")
    rows_items = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh thẻ việc", "PK, IDENTITY(1,1)"],
        ["2", "WorkProjectId", "INT", "Mã dự án (Liên kết WorkProjects)", "FK, NOT NULL"],
        ["3", "TaskName", "NVARCHAR(255)", "Tên đầu việc chi tiết", "NOT NULL"],
        ["4", "AssigneeId", "INT", "Nhân viên thực hiện (Liên kết Employees)", "FK, NULL"],
        ["5", "KPIId", "INT", "KPI liên kết để ghi nhận tiến độ", "FK, NULL"],
        ["6", "Status", "VARCHAR(30)", "Trạng thái (Backlog/Todo/InProgress/Done)", "DEFAULT 'Todo'"],
        ["7", "CompletedAt", "DATETIME", "Thời điểm hoàn thành thực tế", "NULL"],
    ]
    create_spec_table(doc, headers_spec, rows_items, col_widths=col_w)
    add_table_caption(doc, "Bảng B.2: Đặc tả chi tiết bảng WorkItems")

    # --- 3. BonusRules ---
    add_heading2(doc, "3. Đặc tả bảng BonusRules (Quy tắc lương thưởng HR)")
    rows_bonus = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh quy tắc", "PK, IDENTITY(1,1)"],
        ["2", "GradingRankId", "INT", "Xếp hạng liên quan (GradingRanks)", "FK, NOT NULL"],
        ["3", "RankName", "VARCHAR(10)", "Ký hiệu rank (S, A, B, C, D)", "NOT NULL"],
        ["4", "BonusPercentageOfSalary", "DECIMAL(5,2)", "% thưởng theo lương cứng", "NOT NULL, DEFAULT 0"],
        ["5", "FixedBonusAmount", "DECIMAL(18,2)", "Số tiền thưởng cứng", "NOT NULL, DEFAULT 0"],
    ]
    create_spec_table(doc, headers_spec, rows_bonus, col_widths=col_w)
    add_table_caption(doc, "Bảng B.3: Đặc tả chi tiết bảng BonusRules")

    # --- 4. AIGenerationHistories ---
    add_heading2(doc, "4. Đặc tả bảng AIGenerationHistories (Dữ liệu legacy chờ retention)")
    rows_ai = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh log AI", "PK, IDENTITY(1,1)"],
        ["2", "SystemUserId", "INT", "ID tài khoản gọi AI", "FK, NOT NULL"],
        ["3", "PromptType", "VARCHAR(50)", "Loại gọi AI (Chat/KPI_Generator/Analysis)", "NOT NULL"],
        ["4", "PromptText", "NVARCHAR(MAX)", "Nội dung câu hỏi gửi đi", "NOT NULL"],
        ["5", "ResponseText", "NVARCHAR(MAX)", "Phản hồi provider lịch sử; runtime mới không ghi trường này", "NOT NULL"],
        ["6", "TokenCount", "INT", "Kích thước token tiêu dùng", "DEFAULT 0"],
        ["7", "CreatedAt", "DATETIME", "Thời điểm thực hiện cuộc gọi", "DEFAULT GETDATE()"],
    ]
    create_spec_table(doc, headers_spec, rows_ai, col_widths=col_w)
    add_table_caption(doc, "Bảng B.4: Đặc tả bảng legacy AIGenerationHistories (không có reader/writer runtime mới)")


    # ============================================================
    # PHỤ LỤC C: KỊCH BẢN KIỂM THỬ NGOẠI LỆ / EDGE CASES
    # ============================================================
    add_page_break(doc)
    add_appendix_title(doc, "PHỤ LỤC C: BẢNG KỊCH BẢN KIỂM THỬ NGOẠI LỆ (EDGE CASES)")
    
    add_para(doc,
        "Để đảm bảo hệ thống vận hành trơn tru ngay cả khi gặp lỗi người dùng nhập liệu "
        "hoặc các tình huống biên phức tạp, phụ lục này mô tả các kịch bản kiểm thử ngoại lệ đã thực thi:"
    )

    # Bảng Test case ngoại lệ
    headers_tce = ["Mã TC", "Tình huống Edge Case", "Các bước thực hiện", "Kết quả kỳ vọng", "Kết quả thực tế"]
    rows_tce = [
        ["TC_EDGE_01", "Check-in vượt quá target cam kết",
         "1. KPI có TargetValue là 10 (tài liệu).\n2. Nhân viên nhập check-in Achieved Value là 15.",
         "Hệ thống cho phép lưu, ghi nhận tỷ lệ hoàn thành là 150% và tô xanh lá báo vượt chỉ tiêu.", "Đúng kỳ vọng (Pass)"],
        
        ["TC_EDGE_02", "Check-in nhập giá trị âm",
         "1. Mở form check-in.\n2. Nhập Achieved Value là -5. Click Gửi.",
         "Hệ thống validate phía client và báo lỗi 'Giá trị đạt được không được phép nhỏ hơn 0'. Chặn lưu.", "Đúng kỳ vọng (Pass)"],
        
        ["TC_EDGE_03", "Hết hạn dùng thử gói SaaS",
         "1. Tài khoản tenant có TrialEndTime nhỏ hơn thời gian hiện tại.\n2. User cố truy cập Dashboard.",
         "Hệ thống chặn truy cập, hiển thị trang thông báo hết hạn và hướng dẫn thanh toán gia hạn.", "Đúng kỳ vọng (Pass)"],
        
        ["TC_EDGE_04", "Gọi AI khi thiếu secret provider",
         "1. Không cấp DeepSeek__ApiKey trong secret store của môi trường.\n2. User mở AI chat widget.",
         "Hệ thống bắt ngoại lệ, hiển thị thông báo thân thiện 'Tính năng AI đang bảo trì, vui lòng quay lại sau' thay vì báo lỗi crash hệ thống.", "Đúng kỳ vọng (Pass)"],
    ]
    create_table(doc, headers_tce, rows_tce, col_widths=[2.0, 3.5, 4.5, 4.5, 2.0])
    add_table_caption(doc, "Bảng C.1: Danh sách các kịch bản kiểm thử ngoại lệ và kết quả thực tế")


    # ============================================================
    # PHỤ LỤC D: LIÊN KẾT MÃ NGUỒN VÀ TÀI KHOẢN DUYỆT DEMO
    # ============================================================
    add_page_break(doc)
    add_appendix_title(doc, "PHỤ LỤC D: LIÊN KẾT MÃ NGUỒN VÀ TÀI KHOẢN DUYỆT DEMO")
    
    add_para(doc,
        "Hội đồng chấm tốt nghiệp và Giảng viên hướng dẫn có thể duyệt mã nguồn dự án "
        "và đăng nhập trải nghiệm hệ thống theo thông tin liên kết dưới đây:"
    )

    add_bullet(doc, "https://github.com/nhubui2008/Manage-KPI-or-OKR-System", bold_prefix="Đường dẫn kho lưu trữ mã nguồn (Github)")
    add_bullet(doc, "http://localhost:5208 (Môi trường máy chủ thử nghiệm nội bộ)", bold_prefix="Đường dẫn Demo vận hành")

    add_para(doc, "Bảng danh sách tài khoản khảo sát demo (Mật khẩu đăng nhập mặc định cho tất cả tài khoản: 123):")

    # Bảng tài khoản test
    headers_acc = ["STT", "Username đăng nhập", "Vai trò (Role)", "Phòng ban tương ứng", "Quyền hạn kiểm thử nổi bật"]
    rows_acc = [
        ["1", "admin", "Admin (Quản trị viên)", "Phòng Kỹ thuật hệ thống", "Cấu hình phân quyền Roles, xem Audit Logs, thiết lập API Key."],
        ["2", "director", "Director (Giám đốc)", "Ban Giám đốc công ty", "Xem Dashboard quỹ thưởng thực tế, duyệt đánh giá, gọi AI Performance Analysis."],
        ["3", "manager", "Manager (Trưởng phòng)", "Phòng Kỹ thuật (IT Dept)", "Giao KPI cho nhân viên, gọi AI KPI Generator, duyệt check-in tiến độ."],
        ["4", "hr", "HR Specialist (Nhân sự)", "Phòng Hành chính nhân sự", "Cấu hình BonusRules tính thưởng, import nhân viên bằng Excel."],
        ["5", "employee", "Employee (Nhân viên)", "Phòng Kỹ thuật (IT Dept)", "Xem KPI cá nhân, gửi báo cáo check-in kèm Barriers, chat tư vấn AI."],
    ]
    create_table(doc, headers_acc, rows_acc, col_widths=[1.0, 3.5, 3.5, 3.5, 5.0])
    add_table_caption(doc, "Bảng D.1: Danh sách tài khoản demo tích hợp sẵn phục vụ duyệt sản phẩm")


    # ============================================================
    # PHỤ LỤC E: TÀI LIỆU CẤU HÌNH API ENDPOINTS
    # ============================================================
    add_page_break(doc)
    add_appendix_title(doc, "PHỤ LỤC E: TÀI LIỆU CẤU HÌNH API ENDPOINTS")
    
    add_para(doc,
        "Hệ thống thiết lập một số API nội bộ để phục vụ gọi Ajax từ client lên server "
        "và liên kết dữ liệu với các dịch vụ AI qua model gateway. Dưới đây là đặc tả kỹ thuật:"
    )

    # API 1
    add_heading2(doc, "1. API lấy gợi ý KPI thông minh")
    add_bullet(doc, "/AI/SuggestKPI", bold_prefix="Endpoint URL")
    add_bullet(doc, "POST", bold_prefix="Phương thức (Method)")
    add_bullet(doc, "application/json", bold_prefix="Định dạng Header")
    add_bullet(doc, "{ \"periodId\": 1, \"departmentId\": 2, \"employeeId\": 5, \"okrId\": 10, \"okrKeyResultId\": 20 }", bold_prefix="Tham số truyền lên (Body)")
    add_bullet(doc, "{ \"success\": true, \"advisoryOnly\": true, \"agentRunId\": \"...\", \"suggestions\": [{ \"name\": \"...\", \"targetValue\": 100, \"unit\": \"%\", \"passThreshold\": 90, \"failThreshold\": 70, \"isInverse\": false, \"rationale\": \"...\", \"sourceIds\": [\"authorized-kpi-planning-snapshot:...\"] }], \"citations\": [] }", bold_prefix="Kết quả trả về (Response)")
    add_bullet(doc, "Yêu cầu KPIS_CREATE + anti-forgery; chỉ nhận kỳ/phạm vi được phép, strict schema và source ID do server cấp. Kết quả chỉ điền form, không tự tạo KPI và không lưu prompt/raw response.", bold_prefix="Ràng buộc an toàn")

    # API 2
    add_heading2(doc, "2. API phân tích hiệu suất bằng AI")
    add_bullet(doc, "/AI/AnalyzePerformance", bold_prefix="Endpoint URL")
    add_bullet(doc, "POST", bold_prefix="Phương thức (Method)")
    add_bullet(doc, "{ \"periodId\": 1, \"employeeId\": null, \"departmentId\": null } (tenant và actor luôn lấy từ session)", bold_prefix="Tham số truyền lên")
    add_bullet(doc, "{ \"success\": true, \"advisoryOnly\": true, \"overview\": { ... }, \"strengths\": [], \"risks\": [], \"recommendedActions\": [], \"citations\": [] }", bold_prefix="Kết quả trả về")
    add_bullet(doc, "Chỉ dùng check-in đã duyệt, yêu cầu DASHBOARD_VIEW, abstain khi thiếu tiến độ đo lường và không lưu prompt/raw response.", bold_prefix="Ràng buộc an toàn")

    # Kết luận phụ lục
    add_para(doc, "", space_before=12, space_after=0, indent=False)
    add_para(doc,
        "Các phụ lục trên đã khép lại toàn bộ hồ sơ thuyết minh dự án tốt nghiệp. "
        "Tài liệu cung cấp đầy đủ thông tin kỹ thuật từ sơ đồ, đặc tả dữ liệu, kịch bản kiểm thử "
        "đến hướng dẫn vận hành thực tế, hỗ trợ tối đa cho quá trình chuyển giao và nghiệm thu sản phẩm.",
        italic=True, space_before=12, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("  ✓ PHẦN PHỤ LỤC TÀI LIỆU")
    print("    ✓ Phụ lục A: Đặc tả Use Case bổ sung (UC10, UC25 - Bảng A.1, A.2)")
    print("    ✓ Phụ lục B: Đặc tả bảng dữ liệu phụ trợ (Bảng B.1 - B.4)")
    print("    ✓ Phụ lục C: Kiểm thử ngoại lệ Edge Cases (Bảng C.1)")
    print("    ✓ Phụ lục D: Tài khoản Demo (Bảng D.1)")
    print("    ✓ Phụ lục E: Tài liệu API Endpoints")
    write_appendices(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm các Phụ lục vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
