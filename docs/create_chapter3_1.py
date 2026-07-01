"""
Script tạo CHƯƠNG 3: THIẾT KẾ - Phần 3.1 Thiết kế cơ sở dữ liệu
Bao gồm: Danh sách bảng và Đặc tả chi tiết 10 bảng dữ liệu cốt lõi
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    """Tiêu đề chương (CHƯƠNG X: ...) - căn giữa, in đậm, viết hoa (Heading 1)"""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Căn giữa cho cột 0 (STT), các cột khác căn trái
            if c_idx in [0]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def create_spec_table(doc, headers, rows, col_widths=None):
    """Tạo bảng đặc tả chi tiết các cột của bảng"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Căn giữa cột 0 (STT), cột 2 (Kiểu dữ liệu), cột 4 (Ràng buộc). Cột 1 (Tên trường), cột 3 (Mô tả) căn trái.
            if c_idx in [0, 2, 4]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG CHƯƠNG 3 =====================

def write_chapter3(doc):
    """CHƯƠNG 3: THIẾT KẾ - 3.1 Thiết kế cơ sở dữ liệu"""

    add_chapter_title(doc, "CHƯƠNG 3: THIẾT KẾ SẢN PHẨM")

    # ============================================================
    # 3.1. Thiết kế cơ sở dữ liệu
    # ============================================================
    add_heading1(doc, "3.1. Thiết kế cơ sở dữ liệu")

    # 3.1.1. Danh sách bảng dữ liệu
    add_heading2(doc, "3.1.1. Danh sách bảng dữ liệu hệ thống")
    
    add_para(doc,
        "Để quản lý toàn diện các thực thể vận hành bao gồm cơ cấu tổ chức, mục tiêu chiến lược OKR, "
        "chỉ tiêu KPI, check-in, đánh giá hiệu suất, dự án Kanban và tích hợp AI, cơ sở dữ liệu hệ thống "
        "được cấu trúc thành 22 bảng cốt lõi phục vụ lưu trữ nghiệp vụ chính của doanh nghiệp:"
    )

    # Bảng 16: Danh sách bảng
    headers_list = ["STT", "Tên bảng trong CSDL", "Ý nghĩa nghiệp vụ lưu trữ", "Thuộc phân hệ"]
    rows_list = [
        ["1", "SystemUsers", "Tài khoản người dùng đăng nhập hệ thống", "Foundation"],
        ["2", "Employees", "Hồ sơ thông tin chi tiết nhân viên", "Organization"],
        ["3", "Departments", "Cơ cấu cây phòng ban phân cấp đa tầng", "Organization"],
        ["4", "Positions", "Danh mục các chức danh và bậc xếp hạng (Rank)", "Organization"],
        ["5", "MissionVisions", "Thiết lập Sứ mệnh, Tầm nhìn, Mục tiêu năm công ty", "Strategy & OKR"],
        ["6", "OKRs", "Quản lý mục tiêu OKR của công ty, phòng ban, cá nhân", "Strategy & OKR"],
        ["7", "OKRKeyResults", "Kết quả then chốt đo lường tiến độ OKR", "Strategy & OKR"],
        ["8", "EvaluationPeriods", "Các kỳ đánh giá hiệu suất (Quý/Năm)", "KPI setup"],
        ["9", "KPIs", "Danh sách các chỉ tiêu KPI được thiết lập và giao", "KPI setup"],
        ["10", "KPIDetails", "Tham số chi tiết của KPI (Target, Threshold, Unit)", "KPI setup"],
        ["11", "KPICheckIns", "Báo cáo check-in tiến độ định kỳ của nhân viên", "Check-in"],
        ["12", "CheckInDetails", "Báo cáo chi tiết các đầu việc hoàn thành khi check-in", "Check-in"],
        ["13", "EvaluationResults", "Điểm số, xếp hạng (S->D) và thưởng cuối kỳ", "Evaluation"],
        ["14", "BonusRules", "Chính sách quy định tiền thưởng theo xếp hạng", "Evaluation"],
        ["15", "RealtimeExpectedBonuses", "Dự toán quỹ thưởng của công ty theo thời gian thực", "Evaluation"],
        ["16", "WorkProjects", "Dự án công việc phòng ban (Kanban)", "Kanban"],
        ["17", "WorkItems", "Các công việc chi tiết được giao trên Kanban board", "Kanban"],
        ["18", "SystemParameters", "Cấu hình tham số hệ thống động", "Foundation"],
        ["19", "SystemAlerts", "Cảnh báo tiến độ và vận hành hệ thống", "System"],
        ["20", "AuditLogs", "Nhật ký truy vết thao tác người dùng", "System"],
        ["21", "AIGenerationHistories", "Lịch sử cuộc gọi AI và hội thoại Gemini", "System & AI"],
        ["22", "SaaSPackages", "Gói cước dịch vụ SaaS đa doanh nghiệp", "SaaS Admin"],
    ]

    create_table(doc, headers_list, rows_list, col_widths=[1.0, 4.0, 7.5, 3.5])
    add_table_caption(doc, "Bảng 16: Danh sách các bảng dữ liệu cốt lõi trong hệ thống")

    # 3.1.2. Đặc tả chi tiết các bảng cốt lõi
    add_heading2(doc, "3.1.2. Đặc tả chi tiết các bảng cốt lõi")
    
    add_para(doc,
        "Dưới đây là bảng đặc tả chi tiết cấu trúc thuộc tính, kiểu dữ liệu, các ràng buộc và ý nghĩa "
        "của từng trường dữ liệu cho 10 bảng cốt lõi nhất của hệ thống:"
    )

    headers_spec = ["STT", "Tên trường", "Kiểu dữ liệu", "Ý nghĩa / Mô tả", "Ràng buộc"]
    col_w = [1.0, 3.5, 3.0, 6.0, 2.5]

    # --- 1. SystemUsers ---
    add_heading3(doc, "1) Đặc tả bảng SystemUsers (Tài khoản người dùng)")
    rows_user = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh tài khoản", "PK, IDENTITY(1,1)"],
        ["2", "Username", "VARCHAR(255)", "Tên đăng nhập hệ thống", "NOT NULL, UNIQUE"],
        ["3", "Email", "VARCHAR(255)", "Địa chỉ email của tài khoản", "NOT NULL, UNIQUE"],
        ["4", "PasswordHash", "VARCHAR(255)", "Mật khẩu đã mã hóa SHA-256", "NOT NULL"],
        ["5", "LastPasswordChange", "DATETIME", "Thời điểm đổi mật khẩu gần nhất", "DEFAULT GETDATE()"],
        ["6", "RoleId", "INT", "Mã vai trò tài khoản (Liên kết bảng Roles)", "FK, NULL"],
        ["7", "IsActive", "BIT", "Trạng thái tài khoản (1: Hoạt động, 0: Khóa)", "DEFAULT 1"],
        ["8", "CreatedAt", "DATETIME", "Thời điểm tạo tài khoản", "DEFAULT GETDATE()"],
        ["9", "CreatedById", "INT", "ID người tạo tài khoản", "NULL"],
        ["10", "TrialEndTime", "DATETIME", "Hạn kết thúc dùng thử dịch vụ SaaS", "NULL"],
    ]
    create_spec_table(doc, headers_spec, rows_user, col_widths=col_w)
    add_table_caption(doc, "Bảng 17: Đặc tả chi tiết bảng SystemUsers")

    # --- 2. Employees ---
    add_heading3(doc, "2) Đặc tả bảng Employees (Hồ sơ nhân viên)")
    rows_emp = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh nhân viên", "PK, IDENTITY(1,1)"],
        ["2", "EmployeeCode", "VARCHAR(20)", "Mã nhân viên tự sinh (EMP001, EMP002...)", "NOT NULL, UNIQUE"],
        ["3", "FullName", "NVARCHAR(100)", "Họ tên đầy đủ của nhân viên", "NOT NULL"],
        ["4", "DateOfBirth", "DATETIME", "Ngày tháng năm sinh của nhân viên", "NULL"],
        ["5", "Phone", "VARCHAR(15)", "Số điện thoại liên lạc", "NOT NULL"],
        ["6", "Email", "VARCHAR(255)", "Địa chỉ email nhân viên", "NOT NULL, UNIQUE"],
        ["7", "TaxCode", "VARCHAR(50)", "Mã số thuế cá nhân", "NULL"],
        ["8", "JoinDate", "DATETIME", "Ngày bắt đầu làm việc", "NULL"],
        ["9", "SystemUserId", "INT", "Liên kết với bảng SystemUsers", "FK, NULL"],
        ["10", "IsActive", "BIT", "Trạng thái nhân sự (1: Đang làm, 0: Đã nghỉ)", "DEFAULT 1"],
        ["11", "StrategicGoalId", "INT", "Mục tiêu chiến lược nhân sự phụ trách", "FK, NULL"],
        ["12", "CreatedAt", "DATETIME", "Thời điểm tạo hồ sơ", "DEFAULT GETDATE()"],
    ]
    create_spec_table(doc, headers_spec, rows_emp, col_widths=col_w)
    add_table_caption(doc, "Bảng 18: Đặc tả chi tiết bảng Employees")

    # --- 3. Departments ---
    add_heading3(doc, "3) Đặc tả bảng Departments (Phòng ban)")
    rows_dept = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh phòng ban", "PK, IDENTITY(1,1)"],
        ["2", "DepartmentCode", "VARCHAR(20)", "Mã phòng ban tự sinh (DEPT001...)", "NOT NULL, UNIQUE"],
        ["3", "DepartmentName", "NVARCHAR(100)", "Tên gọi phòng ban", "NOT NULL"],
        ["4", "ManagerId", "INT", "ID Trưởng phòng (Liên kết bảng Employees)", "FK, NULL"],
        ["5", "ParentId", "INT", "ID phòng ban cấp cha (Liên kết đệ quy)", "FK, NULL"],
    ]
    create_spec_table(doc, headers_spec, rows_dept, col_widths=col_w)
    add_table_caption(doc, "Bảng 19: Đặc tả chi tiết bảng Departments")

    # --- 4. Positions ---
    add_heading3(doc, "4) Đặc tả bảng Positions (Chức vụ)")
    rows_pos = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh chức vụ", "PK, IDENTITY(1,1)"],
        ["2", "PositionCode", "VARCHAR(20)", "Mã chức vụ tự sinh (POS001...)", "NOT NULL, UNIQUE"],
        ["3", "PositionName", "NVARCHAR(100)", "Tên gọi chức danh chức vụ", "NOT NULL"],
        ["4", "RankLevel", "INT", "Bậc chức vụ phục vụ tính toán thưởng", "NOT NULL, DEFAULT 1"],
    ]
    create_spec_table(doc, headers_spec, rows_pos, col_widths=col_w)
    add_table_caption(doc, "Bảng 20: Đặc tả chi tiết bảng Positions")

    # --- 5. OKRs ---
    add_heading3(doc, "5) Đặc tả bảng OKRs (Mục tiêu OKR)")
    rows_okr = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh OKR", "PK, IDENTITY(1,1)"],
        ["2", "OKRName", "NVARCHAR(255)", "Tên gọi mục tiêu OKR", "NOT NULL"],
        ["3", "Description", "NVARCHAR(1000)", "Mô tả chi tiết mục tiêu OKR", "NULL"],
        ["4", "PeriodId", "INT", "Kỳ đánh giá (Liên kết EvaluationPeriods)", "FK, NOT NULL"],
        ["5", "OKRTypeId", "INT", "Phân loại OKR (1: Công ty, 2: Phòng ban, 3: Cá nhân)", "FK, NOT NULL"],
        ["6", "Progress", "DECIMAL(5,2)", "Tiến độ hoàn thành OKR tính từ các KR", "DEFAULT 0.00"],
        ["7", "IsActive", "BIT", "Trạng thái hoạt động", "DEFAULT 1"],
        ["8", "CreatedAt", "DATETIME", "Thời điểm thiết lập OKR", "DEFAULT GETDATE()"],
    ]
    create_spec_table(doc, headers_spec, rows_okr, col_widths=col_w)
    add_table_caption(doc, "Bảng 21: Đặc tả chi tiết bảng OKRs")

    # --- 6. OKRKeyResults ---
    add_heading3(doc, "6) Đặc tả bảng OKRKeyResults (Kết quả then chốt)")
    rows_kr = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh KR", "PK, IDENTITY(1,1)"],
        ["2", "OKRId", "INT", "Liên kết OKR (Liên kết bảng OKRs)", "FK, NOT NULL"],
        ["3", "KRName", "NVARCHAR(255)", "Tên kết quả then chốt KR", "NOT NULL"],
        ["4", "TargetValue", "DECIMAL(18,2)", "Giá trị mục tiêu cần đạt được", "NOT NULL"],
        ["5", "CurrentValue", "DECIMAL(18,2)", "Giá trị thực tế đạt được hiện tại", "DEFAULT 0.00"],
        ["6", "Unit", "VARCHAR(50)", "Đơn vị đo lường (%, VND, cái...)", "NOT NULL"],
        ["7", "IsInverse", "BIT", "Chỉ số nghịch (1: Càng giảm càng tốt, 0: Càng tăng càng tốt)", "DEFAULT 0"],
    ]
    create_spec_table(doc, headers_spec, rows_kr, col_widths=col_w)
    add_table_caption(doc, "Bảng 22: Đặc tả chi tiết bảng OKRKeyResults")

    # --- 7. KPIs ---
    add_heading3(doc, "7) Đặc tả bảng KPIs (Chỉ tiêu KPI)")
    rows_kpi = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh KPI", "PK, IDENTITY(1,1)"],
        ["2", "PeriodId", "INT", "ID Kỳ đánh giá liên quan", "FK, NULL"],
        ["3", "KPIName", "NVARCHAR(255)", "Tên chỉ tiêu hiệu suất KPI", "NOT NULL"],
        ["4", "Description", "NVARCHAR(1000)", "Mô tả chi tiết yêu cầu KPI", "NULL"],
        ["5", "PropertyId", "INT", "Thuộc tính KPI (Tăng trưởng/Ổn định/Giảm thiểu)", "FK, NULL"],
        ["6", "KPITypeId", "INT", "Phân loại KPI (Định lượng/Định tính/Hành vi)", "FK, NULL"],
        ["7", "OKRId", "INT", "OKR liên quan để đồng bộ tiến độ", "FK, NULL"],
        ["8", "OKRKeyResultId", "INT", "Key Result liên kết trực tiếp để tính %", "FK, NULL"],
        ["9", "AssignerId", "INT", "ID Trưởng phòng giao KPI", "FK, NULL"],
        ["10", "StatusId", "INT", "Trạng thái KPI (Bản nháp/Duyệt/Thực hiện...)", "FK, NULL"],
        ["11", "IsActive", "BIT", "Trạng thái hoạt động", "DEFAULT 1"],
    ]
    create_spec_table(doc, headers_spec, rows_kpi, col_widths=col_w)
    add_table_caption(doc, "Bảng 23: Đặc tả chi tiết bảng KPIs")

    # --- 8. KPIDetails ---
    add_heading3(doc, "8) Đặc tả bảng KPIDetails (Tham số cấu hình KPI)")
    rows_kpidet = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh tham số KPI", "PK, IDENTITY(1,1)"],
        ["2", "KPIId", "INT", "Liên kết chỉ tiêu KPI (Bảng KPIs)", "FK, NULL"],
        ["3", "TargetValue", "DECIMAL(18,2)", "Giá trị mục tiêu cam kết", "NULL"],
        ["4", "PassThreshold", "DECIMAL(18,2)", "Ngưỡng đạt chỉ tiêu tối thiểu", "NULL"],
        ["5", "FailThreshold", "DECIMAL(18,2)", "Ngưỡng không đạt (cảnh báo rủi ro)", "NULL"],
        ["6", "MeasurementUnit", "VARCHAR(50)", "Đơn vị tính chỉ tiêu KPI", "NULL"],
        ["7", "IsInverse", "BIT", "Chỉ số nghịch (càng thấp càng tốt)", "DEFAULT 0"],
        ["8", "DeadlineDate", "DATETIME", "Thời hạn kết thúc thực hiện KPI", "NULL"],
        ["9", "CheckInFrequencyDays", "INT", "Tần suất check-in định kỳ (số ngày)", "DEFAULT 1"],
        ["10", "CheckInDeadlineTime", "TIME", "Giờ giới hạn check-in trong ngày", "DEFAULT '10:00:00'"],
        ["11", "ReminderBeforeHours", "INT", "Thời gian thông báo nhắc nhở trước (giờ)", "DEFAULT 24"],
    ]
    create_spec_table(doc, headers_spec, rows_kpidet, col_widths=col_w)
    add_table_caption(doc, "Bảng 24: Đặc tả chi tiết bảng KPIDetails")

    # --- 9. KPICheckIns ---
    add_heading3(doc, "9) Đặc tả bảng KPICheckIns (Báo cáo check-in)")
    rows_ci = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh check-in", "PK, IDENTITY(1,1)"],
        ["2", "EmployeeId", "INT", "ID nhân viên check-in báo cáo", "FK, NULL"],
        ["3", "KPIId", "INT", "ID chỉ tiêu KPI liên quan", "FK, NULL"],
        ["4", "SubmittedById", "INT", "ID tài khoản gửi báo cáo", "FK, NULL"],
        ["5", "CheckInDate", "DATETIME", "Thời điểm báo cáo thực tế", "DEFAULT GETDATE()"],
        ["6", "DeadlineAt", "DATETIME", "Thời hạn check-in của kỳ báo cáo đó", "NULL"],
        ["7", "IsLate", "BIT", "Cờ đánh dấu báo cáo quá hạn", "NULL"],
        ["8", "StatusId", "INT", "Mã trạng thái check-in", "FK, NULL"],
        ["9", "FailReasonId", "INT", "Mã nguyên nhân chậm tiến độ (nếu có)", "FK, NULL"],
        ["10", "ReviewStatus", "VARCHAR(30)", "Trạng thái duyệt (Pending/Approved/Rejected)", "DEFAULT 'Pending'"],
        ["11", "ReviewedById", "INT", "ID Quản lý duyệt báo cáo", "FK, NULL"],
        ["12", "ReviewedAt", "DATETIME", "Thời điểm phê duyệt thực tế", "NULL"],
        ["13", "ReviewComment", "NVARCHAR(2000)", "Ý kiến nhận xét của quản lý", "NULL"],
        ["14", "ReviewScore", "DECIMAL(5,2)", "Điểm đánh giá tiến độ của lần check-in", "NULL"],
    ]
    create_spec_table(doc, headers_spec, rows_ci, col_widths=col_w)
    add_table_caption(doc, "Bảng 25: Đặc tả chi tiết bảng KPICheckIns")

    # --- 10. EvaluationResults ---
    add_heading3(doc, "10) Đặc tả bảng EvaluationResults (Kết quả đánh giá)")
    rows_eval = [
        ["1", "Id", "INT", "Khóa chính, tự tăng định danh kết quả đánh giá", "PK, IDENTITY(1,1)"],
        ["2", "EmployeeId", "INT", "ID nhân viên được đánh giá hiệu suất", "FK, NULL"],
        ["3", "PeriodId", "INT", "ID Kỳ đánh giá liên quan", "FK, NULL"],
        ["4", "TotalScore", "DECIMAL(5,2)", "Tổng điểm trung bình KPI đạt được", "NULL"],
        ["5", "RankId", "INT", "ID Xếp hạng hiệu suất (Liên kết bảng GradingRanks)", "FK, NULL"],
        ["6", "Classification", "VARCHAR(50)", "Chuỗi phân loại rank (S, A+, A, B+, B, C, D)", "NULL"],
        ["7", "ReviewComment", "NVARCHAR(2000)", "Nhận xét tổng hợp của Trưởng phòng", "NULL"],
        ["8", "SubmissionStatus", "VARCHAR(30)", "Trạng thái đánh giá (Draft/Submitted/Reviewed)", "DEFAULT 'Draft'"],
        ["9", "SubmittedById", "INT", "ID Trưởng phòng gửi bảng đánh giá", "FK, NULL"],
        ["10", "SubmittedAt", "DATETIME", "Thời điểm nộp bảng đánh giá", "NULL"],
        ["11", "DirectorReviewedById", "INT", "ID Giám đốc duyệt chốt đánh giá", "FK, NULL"],
        ["12", "DirectorReviewedAt", "DATETIME", "Thời điểm Giám đốc phê duyệt", "NULL"],
        ["13", "DirectorReviewComment", "NVARCHAR(2000)", "Nhận xét bổ sung của Giám đốc", "NULL"],
    ]
    create_spec_table(doc, headers_spec, rows_eval, col_widths=col_w)
    add_table_caption(doc, "Bảng 26: Đặc tả chi tiết bảng EvaluationResults")

    # Kết luận mục
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Việc xây dựng đặc tả chi tiết 10 bảng dữ liệu cốt lõi trên với đầy đủ các thuộc tính khóa, "
        "kiểu dữ liệu và ràng buộc chặt chẽ đảm bảo tính đồng bộ hoàn toàn giữa cơ sở dữ liệu vật lý trên SQL Server "
        "và các lớp model Entity Framework Core của mã nguồn ASP.NET 10 MVC.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    add_page_break(doc)

    print("  ✓ CHƯƠNG 3: THIẾT KẾ SẢN PHẨM")
    print("    ✓ 3.1. Thiết kế cơ sở dữ liệu")
    print("      ✓ 3.1.1. Danh sách 22 bảng chính (Bảng 16)")
    print("      ✓ 3.1.2. Đặc tả chi tiết 10 bảng cốt lõi (Bảng 17 - 26)")
    write_chapter3(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm Chương 3 (3.1) vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
