"""
Script tạo phần 5.2: Kịch bản kiểm thử (Test Cases) cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Căn lề: cột 0 (STT), cột 1 (Mã TC), cột 5 (Trạng thái) căn giữa. Các cột khác căn trái.
            if c_idx in [0, 1, 5]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 5.2 =====================

def write_section_5_2(doc):
    """5.2. Test case"""

    add_heading1(doc, "5.2. Kịch bản kiểm thử (Test Cases)")

    add_para(doc,
        "Dựa trên kế hoạch kiểm thử đã lập, Tester Đoàn Quốc Khánh đã thiết lập kịch bản kiểm thử chi tiết "
        "cho các chức năng trọng yếu của hệ thống. Dưới đây là bảng tổng hợp danh sách các Test Cases thực tế:"
    )

    # Bảng 30: Danh sách Test Cases
    headers_tc = ["STT", "Mã TC", "Chức năng / Tên Test Case", "Các bước thực hiện", "Kết quả kỳ vọng", "Trạng thái"]
    rows_tc = [
        ["1", "TC_AUTH_01", "Đăng nhập tài khoản hợp lệ",
         "1. Truy cập /Auth/Login.\n2. Nhập Username 'director', Password '123'.\n3. Click 'Đăng nhập'.",
         "Xác thực thành công. Hệ thống thiết lập Cookie và tự động chuyển hướng sang trang Dashboard chính.", "Đạt (Pass)"],
        
        ["2", "TC_AUTH_02", "Đăng nhập sai mật khẩu",
         "1. Truy cập /Auth/Login.\n2. Nhập Username 'director', Password 'wrong'.\n3. Click 'Đăng nhập'.",
         "Hệ thống từ chối xác thực. Hiển thị thông báo lỗi 'Tên đăng nhập hoặc mật khẩu không chính xác'.", "Đạt (Pass)"],
        
        ["3", "TC_AUTH_03", "Đăng nhập bằng Google OAuth",
         "1. Truy cập /Auth/Login.\n2. Click nút 'Đăng nhập bằng Google'.\n3. Chọn email Google đã đăng ký.",
         "Google trả về token hợp lệ. Hệ thống thiết lập Claims và chuyển hướng sang Dashboard thành công.", "Đạt (Pass)"],
        
        ["4", "TC_OKR_01", "Thiết lập OKR cấp công ty",
         "1. Đăng nhập với vai trò Director.\n2. Vào mục OKRs -> 'Tạo mới OKR'.\n3. Nhập tên OKR, kỳ đánh giá, click 'Lưu'.",
         "OKR mới được lưu vào CSDL bảng OKRs. Hiển thị chính xác trên sơ đồ phân cấp OKR 3 cấp.", "Đạt (Pass)"],
        
        ["5", "TC_KPI_01", "Giao KPI cho nhân viên trực thuộc",
         "1. Đăng nhập với vai trò Manager.\n2. Vào KPIs -> 'Giao KPI'.\n3. Chọn nhân viên, nhập Target, Trọng số (20%), click 'Gửi'.",
         "KPI mới được tạo trong bảng KPIs ở trạng thái 'Chờ duyệt' (Pending Decision) và gửi mail báo cho Director.", "Đạt (Pass)"],
        
        ["6", "TC_KPI_02", "Giao KPI vượt quá 100% trọng số",
         "1. Manager chọn nhân viên có tổng trọng số KPI hiện tại là 90%.\n2. Tiến hành giao thêm KPI mới với Trọng số 20%.",
         "Hệ thống phát hiện tổng trọng số vượt quá 100% (90+20=110%). Báo lỗi và từ chối lưu vào database.", "Đạt (Pass)"],
        
        ["7", "TC_CHECKIN_01", "Check-in báo cáo tiến độ KPI",
         "1. Đăng nhập với vai trò Employee.\n2. Vào KPIs của tôi -> chọn KPI và click 'Check-in'.\n3. Nhập kết quả đạt được, giải trình.",
         "Hệ thống tự động tính %, ghi nhận trạng thái 'Chờ duyệt' (Pending Review) và đẩy vào Review Queue của Manager.", "Đạt (Pass)"],
        
        ["8", "TC_CHECKIN_02", "Manager phê duyệt tiến độ check-in",
         "1. Đăng nhập với vai trò Manager.\n2. Vào Review Queue -> Click 'Xem' bản check-in.\n3. Nhận xét và click 'Duyệt'.",
         "Trạng thái chuyển sang Approved. OKRProgressService tự động đồng bộ tăng tiến độ KR và OKR liên quan.", "Đạt (Pass)"],
        
        ["9", "TC_AI_01", "Gọi trợ lý AI gợi ý KPI",
         "1. Mở form giao KPI.\n2. Click nút 'Gợi ý từ AI'.\n3. AI đề xuất KPI dựa trên context phòng ban.",
         "Model gateway trả 3-5 bản nháp strict JSON có citation; server chặn schema/nguồn/ngưỡng sai, không tạo KPI tự động và người dùng có thể áp dụng bản nháp vào form.", "Đạt (Pass)"],
        
        ["10", "TC_AI_02", "Kiểm tra Rate Limit gọi AI",
         "1. Thực hiện gửi câu hỏi liên tục >15 lần/phút trong widget AI chat.",
         "Hệ thống kích hoạt Rate Limit. Chặn các cuộc gọi tiếp theo và hiển thị thông báo 'Trợ lý AI đang quá tải'.", "Đạt (Pass)"],
        
        ["11", "TC_EVAL_01", "Tính toán thưởng tự động HR",
         "1. Đăng nhập vai trò HR.\n2. Xem bảng EvaluationResults cuối kỳ sau khi Director chốt.",
         "Hệ thống so sánh Target vs Achieved, tự động xếp rank S->D và tính tiền thưởng chính xác theo BonusRules.", "Đạt (Pass)"],
        
        ["12", "TC_SEC_01", "Bảo mật dữ liệu chéo (Access Scope)",
         "1. Đăng nhập vai trò Employee thuộc IT.\n2. Cố tình truy cập URL `/KPICheckIns` của phòng ban Marketing.",
         "Hệ thống phát hiện vi phạm phân cấp Scope. Chặn truy cập và chuyển hướng về trang chủ hoặc báo lỗi 403.", "Đạt (Pass)"],
    ]

    create_table(doc, headers_tc, rows_tc, col_widths=[1.0, 2.0, 3.5, 4.5, 4.5, 1.5])
    add_table_caption(doc, "Bảng 30: Danh sách kịch bản kiểm thử chi tiết các chức năng")

    # Đánh giá chung kết quả kiểm thử
    add_heading2(doc, "5.2.2. Đánh giá chung kết quả kiểm thử")
    
    add_para(doc,
        "Thông qua quá trình thực thi 12 kịch bản kiểm thử trọng yếu trên các trình duyệt khác nhau, "
        "kết quả cho thấy toàn bộ các kịch bản đều đạt trạng thái ĐẠT (Pass). Hệ thống hoạt động chính xác theo đúng "
        "thiết kế logic nghiệp vụ, cơ chế phân quyền Claims và Access Scope bảo mật hoạt động hiệu quả, "
        "tốc độ xử lý của backend nhanh chóng và các luồng AI có kiểm soát hoạt động ổn định."
    )

    # Kết luận chương
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Tổng kết Chương 5: Việc lập kế hoạch kiểm thử khoa học và thực thi chi tiết các kịch bản kiểm thử chức năng, "
        "hiệu năng và bảo mật đã xác minh tính chính xác và độ tin cậy của phần mềm. Kết quả này đảm bảo "
        "sản phẩm hoàn toàn đủ điều kiện nghiệm thu và chuyển sang giai đoạn vận hành, hướng dẫn sử dụng ở chương cuối.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 5.2. Test cases")
    print("      ✓ 5.2.1. Bảng 12 kịch bản kiểm thử chi tiết (Bảng 30)")
    print("      ✓ 5.2.2. Đánh giá chung kết quả kiểm thử")
    write_section_5_2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 5.2 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
