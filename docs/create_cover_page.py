"""
Script tạo trang bìa Báo cáo Dự án Tốt nghiệp - FPT Polytechnic
Theo bố cục mẫu SD-17/SD-29/SD-38
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import os

# ===================== CẤU HÌNH THÔNG TIN =====================
TRUONG = "TRƯỜNG CAO ĐẲNG FPT POLYTECHNIC"
BAO_CAO = "BÁO CÁO DỰ ÁN TỐT NGHIỆP"
DE_TAI = "ĐỀ TÀI: HỆ THỐNG HỖ TRỢ VẬN HÀNH THÔNG MINH CHO DOANH NGHIỆP VỪA VÀ NHỎ HỖ TRỢ QUẢN LÝ ĐA CẤP VÀ ĐƯA RA QUYẾT ĐỊNH BẰNG AI"
GIANG_VIEN = "Phan Hoàng Khải"
CHUYEN_NGANH = "Phát triển phần mềm"
NHOM = "NEXTGEN"
DIA_DIEM_NAM = "Hà Nội - 2026"

# Danh sách sinh viên: (Tên, Mã sinh viên)
SINH_VIEN = [
    ("Phạm Trần Anh Quân", "TB01758"),
    ("Phạm Trần An An", "TB01817"),
    ("Bùi Nguyễn Anh Như", "TB01785"),
    ("Trần Thanh Phong", "TB01649"),
    ("Nguyễn Thế Bảo", "TB01573"),
    ("Đoàn Quốc Khánh", "TB01544"),
    ("Vũ Hoàng Huy Nhật", "TB01605"),
]

# ===================== ĐƯỜNG DẪN =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(SCRIPT_DIR, "img_page0_2.png")  # Logo FPT Polytechnic
BORDER_PATH = os.path.join(SCRIPT_DIR, "img_page0_1.png")  # Khung viền trang bìa
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")


def set_page_margins(section):
    """Thiết lập lề trang A4"""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def add_page_border(section):
    """Thêm border trang kiểu hoa văn (art border) giống mẫu"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'thinThickSmallGap')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '1F4E79')  # Xanh đậm giống mẫu
        pgBorders.append(border)
    
    sectPr.append(pgBorders)


def add_empty_lines(doc, count=1, font_size=14, font_name='Times New Roman'):
    """Thêm dòng trống"""
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(font_size + 4)


def add_centered_text(doc, text, font_size=14, bold=False, font_name='Times New Roman',
                      color=RGBColor(0, 0, 0), space_before=0, space_after=0):
    """Thêm đoạn text căn giữa"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    run.font.color.rgb = color
    return p


def set_tab_stops(paragraph, positions):
    """Thiết lập tab stops cho paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    for pos in positions:
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:pos'), str(pos))
        tabs.append(tab)
    pPr.append(tabs)


def add_info_row(doc, label, value, font_size=14, font_name='Times New Roman'):
    """Thêm dòng thông tin: Label:    Value"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.left_indent = Cm(3.0)
    
    # Label (bold)
    run_label = p.add_run(label)
    run_label.font.size = Pt(font_size)
    run_label.font.name = font_name
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run_label.bold = True
    
    # Tab
    run_sep = p.add_run('\t')
    run_sep.font.size = Pt(font_size)
    run_sep.font.name = font_name
    
    # Value (normal)
    run_value = p.add_run(value)
    run_value.font.size = Pt(font_size)
    run_value.font.name = font_name
    run_value._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    set_tab_stops(p, [5670])
    return p


def add_student_row(doc, name, student_id, font_size=14, font_name='Times New Roman',
                    is_first=False):
    """Thêm dòng sinh viên"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.left_indent = Cm(3.0)
    
    if is_first:
        run_label = p.add_run('Sinh viên thực hiện:')
        run_label.font.size = Pt(font_size)
        run_label.font.name = font_name
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run_label.bold = True
        
        run_sep = p.add_run('\t')
        run_sep.font.size = Pt(font_size)
        run_sep.font.name = font_name
    else:
        run_sep = p.add_run('\t')
        run_sep.font.size = Pt(font_size)
        run_sep.font.name = font_name
    
    # Tên
    run_name = p.add_run(name)
    run_name.font.size = Pt(font_size)
    run_name.font.name = font_name
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    # Tab 2
    run_tab = p.add_run('\t')
    run_tab.font.size = Pt(font_size)
    run_tab.font.name = font_name
    
    # Mã SV
    run_id = p.add_run(student_id)
    run_id.font.size = Pt(font_size)
    run_id.font.name = font_name
    run_id._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    set_tab_stops(p, [5670, 8505])
    return p


def create_cover_page():
    """Tạo trang bìa theo mẫu FPT Polytechnic"""
    doc = Document()
    
    # === THIẾT LẬP TRANG ===
    section = doc.sections[0]
    set_page_margins(section)
    
    # === Font mặc định ===
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # === BORDER TRANG ===
    add_page_border(section)
    print("  ✓ Đã thêm khung viền trang")
    
    # === 1. KHOẢNG TRỐNG ĐẦU ===
    add_empty_lines(doc, 1, font_size=6)
    
    # === 2. LOGO FPT POLYTECHNIC ===
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(6)
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO_PATH, width=Cm(5.5))
        print(f"  ✓ Đã thêm logo: {LOGO_PATH}")
    else:
        print(f"  ⚠ Không tìm thấy logo: {LOGO_PATH}")
    
    # === 3. TÊN TRƯỜNG ===
    add_centered_text(doc, TRUONG, font_size=14, bold=True, space_before=6, space_after=0)
    
    # === 4. KHOẢNG TRỐNG GIỮA ===
    add_empty_lines(doc, 6, font_size=14)
    
    # === 5. BÁO CÁO DỰ ÁN TỐT NGHIỆP ===
    add_centered_text(doc, BAO_CAO, font_size=20, bold=True, space_before=0, space_after=6)
    
    # === 6. ĐỀ TÀI ===
    add_centered_text(doc, DE_TAI, font_size=14, bold=True, space_before=0, space_after=0)
    
    # === 7. KHOẢNG TRỐNG ===
    add_empty_lines(doc, 6, font_size=14)
    
    # === 8. THÔNG TIN GIẢNG VIÊN, CHUYÊN NGÀNH, NHÓM ===
    add_info_row(doc, 'Giảng viên hướng dẫn:', GIANG_VIEN)
    add_info_row(doc, 'Chuyên ngành:', CHUYEN_NGANH)
    add_info_row(doc, 'Nhóm thực hiện:', NHOM)
    
    # === 9. DANH SÁCH SINH VIÊN ===
    for i, (name, sid) in enumerate(SINH_VIEN):
        add_student_row(doc, name, sid, is_first=(i == 0))
    
    # === 10. KHOẢNG TRỐNG CUỐI ===
    add_empty_lines(doc, 4, font_size=14)
    
    # === 11. ĐỊA ĐIỂM - NĂM ===
    add_centered_text(doc, DIA_DIEM_NAM, font_size=14, bold=True, space_before=0, space_after=0)
    
    # === LƯU FILE ===
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã tạo file Word thành công!")
    print(f"   📄 {OUTPUT_PATH}")
    print(f"\n📝 Lưu ý: Hãy cập nhật danh sách sinh viên trong script nếu cần.")


if __name__ == '__main__':
    create_cover_page()
