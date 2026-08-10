"""
Script tạo phần 2.2: Tác nhân hệ thống cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            # Căn lề: cột 0 (STT), cột 1 (Tên tác nhân), cột 2 (Loại tác nhân) căn giữa. Cột 3 (Phạm vi), cột 4 (Nhiệm vụ) căn trái.
            if c_idx in [0, 1, 2]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 2.2 =====================

def write_section_2_2(doc):
    """2.2. Tác nhân hệ thống"""

    add_heading1(doc, "2.2. Tác nhân hệ thống")

    add_para(doc,
        "Việc làm rõ các tác nhân (Actors) giúp nhóm xây dựng đúng cơ chế phân quyền "
        "(Role-Based Access Control - RBAC) và thiết kế phạm vi truy cập dữ liệu (Access Scope) "
        "chính xác. Trong hệ thống hỗ trợ vận hành thông minh này, 5 tác nhân chính được định nghĩa "
        "chi tiết như sau:"
    )

    # Bảng 10: Tác nhân hệ thống
    headers_actor = ["STT", "Tên tác nhân", "Phân loại", "Phạm vi dữ liệu (Scope)", "Mô tả vai trò và nhiệm vụ chính"]
    rows_actor = [
        ["1", "Admin\n(Quản trị SaaS)", "Tác nhân trong\n(Cấp hệ thống)", "Toàn bộ hệ thống SaaS\n(Đa doanh nghiệp)",
         "Quản lý cơ sở hạ tầng SaaS, CRUD các gói dịch vụ (SaaSPackage), duyệt đăng ký mua gói của các công ty mới, cấu hình tham số hệ thống chung và giám sát nhật ký Audit Logs toàn cục."],
        ["2", "Director\n(Ban Giám Đốc)", "Người dùng\n(Nội bộ)", "Toàn bộ doanh nghiệp sở tại\n(Company Scope)",
         "Thiết lập sứ mệnh, tầm nhìn, mục tiêu năm (MissionVisions). Tạo OKR cấp công ty. Phê duyệt kết quả đánh giá cuối kỳ và dự toán ngân sách thưởng. Theo dõi dashboard tổng quan và nhận cảnh báo rủi ro AI."],
        ["3", "HR\n(Nhân sự)", "Người dùng\n(Nội bộ)", "Toàn bộ nhân sự & phòng ban\n(Company HR Scope)",
         "Quản lý cây phòng ban, danh mục chức vụ và hồ sơ nhân viên (CRUD, import Excel). Thiết lập kỳ đánh giá hiệu suất, cấu hình quy tắc thưởng theo Rank (BonusRules) và xuất báo cáo tổng hợp."],
        ["4", "Manager\n(Trưởng phòng)", "Người dùng\n(Nội bộ)", "Trong phòng ban trực thuộc\n(Department Scope)",
         "Nhận OKR công ty để phân rã OKR phòng ban. Giao KPI cho nhân viên. Phê duyệt tiến độ check-in (Review Queue). Quản lý dự án, phân công việc trên Kanban. Chấm điểm nhân viên và họp 1-on-1."],
        ["5", "Employee\n(Nhân viên)", "Người dùng\n(Nội bộ)", "Chỉ dữ liệu cá nhân được giao\n(Personal Scope)",
         "Xem OKR/KPI cá nhân. Check-in tiến độ KPI định kỳ kèm giải trình. Nhận việc và cập nhật thẻ Kanban công việc. Xem điểm xếp hạng và dự toán thưởng cá nhân. Tương tác với Chat Advisor có nguồn để nhận tư vấn."],
    ]

    create_table(doc, headers_actor, rows_actor, col_widths=[1.0, 2.5, 2.5, 3.8, 6.5])
    add_table_caption(doc, "Bảng 10: Đặc tả chi tiết các tác nhân tương tác với hệ thống")

    # Phân cấp và Mối quan hệ giữa các tác nhân
    add_heading2(doc, "2.2.2. Phân cấp dữ liệu và cơ chế bảo mật giữa các tác nhân")
    
    add_para(doc,
        "Để đảm bảo an toàn thông tin và phù hợp với mô hình quản lý của doanh nghiệp vừa và nhỏ, "
        "hệ thống áp dụng cơ chế phân quyền đa cấp chặt chẽ được xử lý thông qua AccessScopeHelper và PermissionClaimsTransformation:"
    )

    add_bullet(doc, "Admin hệ thống đứng ở tầng cao nhất, kiểm soát hạ tầng và tài khoản doanh nghiệp. Admin không can thiệp vào dữ liệu vận hành cụ thể (KPI, OKR, Thưởng) của các doanh nghiệp thành viên.", bold_prefix="Cấp quản trị SaaS")
    
    add_bullet(doc, "Director có quyền hạn cao nhất trong doanh nghiệp sở tại, được quyền xem tất cả dữ liệu chiến lược, hiệu suất và thưởng của mọi phòng ban, nhân sự. Director là người duy nhất chốt kỳ đánh giá.", bold_prefix="Cấp quyết định vĩ mô")
    
    add_bullet(doc, "Manager có quyền quản lý toàn quyền trong phạm vi phòng ban mình trực thuộc. Manager không thể xem dữ liệu KPI/OKR hay bảng lương thưởng của các phòng ban khác (bảo mật ngang hàng).", bold_prefix="Cấp quản lý trung gian")
    
    add_bullet(doc, "Employee chỉ có quyền thao tác trên các KPI được giao và các thẻ công việc (WorkItems) cá nhân. Employee không thể xem kết quả đánh giá của đồng nghiệp khác.", bold_prefix="Cấp thực thi cá nhân")

    # Kết luận
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Việc định nghĩa rõ các tác nhân và cơ chế phân cấp dữ liệu giúp hệ thống đảm bảo tính "
        "bảo mật, giảm thiểu tối đa rủi ro lộ lọt thông tin nhạy cảm của doanh nghiệp, đồng thời "
        "tối ưu hóa luồng xử lý của các sơ đồ hoạt động (Activity Diagrams) tiếp theo.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 2.2. Tác nhân hệ thống")
    print("      ✓ 2.2.1. Đặc tả chi tiết các tác nhân (Bảng 10)")
    print("      ✓ 2.2.2. Cơ chế phân cấp dữ liệu và bảo mật (AccessScope)")
    write_section_2_2(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 2.2 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
