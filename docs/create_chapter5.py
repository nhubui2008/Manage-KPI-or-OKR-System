"""
Script tạo CHƯƠNG 5: KIỂM THỬ HỆ THỐNG - Phần 5.1 Kế hoạch kiểm thử
Dựa trên cấu trúc chuẩn của báo cáo tốt nghiệp FPT Polytechnic
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    """Tiêu đề chương (CHƯƠNG X: ...) - căn giữa, in đậm, viết hoa (Heading 1)"""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if c_idx in [0, 2]:
                align = WD_ALIGN_PARAGRAPH.CENTER
            else:
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=11, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG CHƯƠNG 5 =====================

def write_chapter5(doc):
    """CHƯƠNG 5: KIỂM THỬ HỆ THỐNG - 5.1 Kế hoạch kiểm thử"""

    add_chapter_title(doc, "CHƯƠNG 5: KIỂM THỬ HỆ THỐNG")

    # ============================================================
    # 5.1. Kế hoạch kiểm thử
    # ============================================================
    add_heading1(doc, "5.1. Kế hoạch kiểm thử")

    # 5.1.1. Mục tiêu kiểm thử
    add_heading2(doc, "5.1.1. Mục tiêu kiểm thử")
    
    add_para(doc,
        "Hoạt động kiểm thử nhằm phát hiện tối đa các lỗi phần mềm trước khi nghiệm thu, "
        "đảm bảo hệ thống đáp ứng đúng và đủ các yêu cầu nghiệp vụ thực tế của doanh nghiệp mục tiêu. "
        "Các mục tiêu cụ thể bao gồm:"
    )
    add_bullet(doc, "Đảm bảo tất cả 28 Use Cases và các chức năng đã xây dựng chạy đúng nghiệp vụ như mô tả đặc tả.", bold_prefix="Độ chính xác chức năng")
    add_bullet(doc, "Đảm bảo tính chính xác tuyệt đối trong việc phân rã OKR đa cấp, tính % tiến độ và quy tắc tính thưởng HR.", bold_prefix="Tính toán dữ liệu")
    add_bullet(doc, "Xác minh cơ chế Claims-based Auth kiểm soát đúng 60 permissions, phân định đúng phạm vi Access Scope.", bold_prefix="Độ an toàn bảo mật")
    add_bullet(doc, "Kiểm tra khả năng phản hồi mượt mà dưới 500ms đối với các tác vụ thông thường, kiểm tra rate limit AI.", bold_prefix="Hiệu năng & Ổn định")
    add_bullet(doc, "Đảm bảo giao diện responsive hiển thị chuẩn trên cả màn hình desktop và điện thoại di động.", bold_prefix="Trải nghiệm người dùng")

    # 5.1.2. Phạm vi kiểm thử
    add_heading2(doc, "5.1.2. Phạm vi kiểm thử")
    
    add_para(doc, "Phạm vi kiểm thử bao phủ toàn bộ các phân hệ chức năng trên môi trường Web ứng dụng:")
    add_bullet(doc, "Đăng nhập, đổi mật khẩu, Google OAuth2, quên mật khẩu nhận OTP qua Email.", bold_prefix="Module 1: Xác thực & Bảo mật")
    add_bullet(doc, "Thiết lập cây phòng ban, chức vụ (Rank) và hồ sơ nhân sự (CRUD, import/export Excel).", bold_prefix="Module 2: Nhân sự & Tổ chức")
    add_bullet(doc, "Tạo OKR công ty/phòng ban, phân bổ KR, tự động cập nhật tiến độ liên cấp.", bold_prefix="Module 3: Chiến lược & OKR")
    add_bullet(doc, "Giao chỉ tiêu KPI, cấu hình chi tiết (Target, Threshold, Unit, deadline, trọng số), duyệt workflow.", bold_prefix="Module 4: Thiết lập KPI")
    add_bullet(doc, "Gửi báo cáo check-in, đính kèm giải trình, duyệt check-in trong Review Queue, goal comments.", bold_prefix="Module 5: Check-in tiến độ")
    add_bullet(doc, "Tạo dự án phòng ban, phân công việc, kéo thả Kanban board, liên kết WorkItem với KPI.", bold_prefix="Module 6: Dự án & Kanban")
    add_bullet(doc, "Mở/đóng kỳ đánh giá, xếp hạng rank tự động, tính toán dự toán thưởng, export Excel tổng hợp.", bold_prefix="Module 7: Đánh giá & Thưởng")
    add_bullet(doc, "Tư vấn chatbot, gọi AI gợi ý KPI, phân tích hiệu suất và Smart Alerts cảnh báo rủi ro.", bold_prefix="Module 8: Trợ lý AI Gemini")

    # 5.1.3. Phương pháp kiểm thử
    add_heading2(doc, "5.1.3. Phương pháp kiểm thử")
    
    add_para(doc, "Để đạt hiệu quả tối đa, dự án áp dụng kết hợp các phương pháp kiểm thử sau:")
    add_bullet(doc, "Tập trung kiểm thử chức năng (Functional Testing) thông qua phương pháp kiểm thử hộp đen (Black-box Testing) bằng cách xây dựng kịch bản kiểm thử (Test Cases) đầu vào và đối chiếu kết quả đầu ra thực tế.", bold_prefix="Kiểm thử hộp đen")
    add_bullet(doc, "Kiểm tra tích hợp (Integration Testing) để xác minh các API giao tiếp đúng đắn giữa các layer và liên kết đồng bộ giữa các module (ví dụ: duyệt check-in phải đồng bộ OKR).", bold_prefix="Kiểm thử tích hợp")
    add_bullet(doc, "Thử nghiệm đăng nhập trái phép, cố ý truy cập chéo ID phòng ban để kiểm tra tính hiệu quả của AccessScopeHelper.", bold_prefix="Kiểm thử bảo mật (Security)")
    add_bullet(doc, "Thử nghiệm import tệp Excel lớn chứa 240 nhân viên để kiểm tra tính ổn định và tốc độ xử lý.", bold_prefix="Kiểm thử tải & Hiệu năng (Performance)")

    # 5.1.4. Tiêu chí Đạt/Không đạt (Pass/Fail Criteria)
    add_heading2(doc, "5.1.4. Tiêu chí Đạt/Không đạt (Pass/Fail Criteria)")
    
    add_para(doc, "Một chức năng hoặc đợt kiểm thử được coi là ĐẠT (Pass) khi đáp ứng toàn bộ các tiêu chuẩn sau:")
    add_bullet(doc, "Tất cả các bước trong kịch bản kiểm thử (Test Steps) đều thực hiện thành công và kết quả thực tế (Actual Result) khớp hoàn toàn với kết quả kỳ vọng (Expected Result).", bold_prefix="Kết quả khớp kỳ vọng")
    add_bullet(doc, "Hệ thống không còn bất kỳ lỗi nghiêm trọng nào thuộc Severity 1 (Lỗi hệ thống crash, mất dữ liệu) hoặc Severity 2 (Lỗi chức năng chính không hoạt động).", bold_prefix="Không còn lỗi nghiêm trọng")
    
    add_para(doc, "Một chức năng bị coi là KHÔNG ĐẠT (Fail) và cần ghi nhận lỗi (Bug) khi:")
    add_bullet(doc, "Xảy ra lỗi crash, lỗi trắng trang (HTTP 500), hoặc dữ liệu lưu vào database sai lệch so với giá trị nhập.", bold_prefix="Lỗi vật lý & logic")
    add_bullet(doc, "Giao diện bị vỡ, không bấm được các nút tương tác trên thiết bị di động.", bold_prefix="Lỗi tương thích UI")

    # 5.1.5. Môi trường kiểm thử
    add_heading2(doc, "5.1.5. Môi trường kiểm thử")
    
    add_para(doc, "Đợt kiểm thử hệ thống được thực hiện trên môi trường tiêu chuẩn sau:")
    
    # Bảng 29: Môi trường kiểm thử
    headers_env = ["Thành phần", "Cấu hình / Thông số kiểm thử", "Mô tả vai trò"]
    rows_env = [
        ["Thiết bị kiểm thử", "Laptop Dell Vostro 15 (CPU Intel Core i7, 16GB RAM, SSD 512GB)", "Thiết bị vật lý của Tester"],
        ["Hệ điều hành", "Windows 11 Home 64-bit", "Nền tảng chạy máy chủ cục bộ"],
        ["Cơ sở dữ liệu", "Microsoft SQL Server 2019 LocalDB", "RDBMS lưu trữ dữ liệu test"],
        ["Trình duyệt test", "Google Chrome (v124+), Microsoft Edge (v124+)", "Kiểm thử tương thích web"],
        ["Thiết bị di động", "iPhone 13 (Safari iOS 17), Samsung Galaxy S22 (Chrome Android 14)", "Kiểm thử Responsive UI"],
        ["Công cụ test", "Postman (v10+)", "Kiểm thử API endpoints"],
    ]
    create_table(doc, headers_env, rows_env, col_widths=[4.0, 7.5, 4.5])
    add_table_caption(doc, "Bảng 29: Đặc tả môi trường kiểm thử hệ thống")

    # 5.1.6. Người thực hiện kiểm thử
    add_heading2(doc, "5.1.6. Người thực hiện kiểm thử")
    
    add_para(doc,
        "Để đảm bảo tính khách quan và phát hiện lỗi tối đa, quy trình kiểm thử được phân công cụ thể "
        "giữa vai trò Tester chuyên trách và đội ngũ lập trình viên:"
    )
    add_bullet(doc, "Chịu trách nhiệm chính thiết lập kịch bản test (Test Cases), thực hiện Manual Test toàn bộ giao diện, ghi nhận log bug lên hệ thống theo dõi.", bold_prefix="Đoàn Quốc Khánh (Tester)")
    add_bullet(doc, "Thực hiện Unit Test trên mã nguồn lớp Backend Services và hỗ trợ kiểm thử tích hợp (Integration Test) sau khi ghép nối các module.", bold_prefix="Đội ngũ Developer (Quân, An, Bảo, Phong, Nhật, Như)")

    # Kết luận mục
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Kế hoạch kiểm thử chi tiết này là cơ sở định hướng cho quá trình thực thi, "
        "đảm bảo các kịch bản kiểm thử tiếp theo được tiến hành một cách khoa học, "
        "giúp nâng cao chất lượng phần mềm trước khi bàn giao.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 5.1. Kế hoạch kiểm thử")
    print("      ✓ 5.1.1. Mục tiêu kiểm thử")
    print("      ✓ 5.1.2. Phạm vi kiểm thử")
    print("      ✓ 5.1.3. Phương pháp kiểm thử")
    print("      ✓ 5.1.4. Tiêu chí Đạt/Không đạt")
    print("      ✓ 5.1.5. Môi trường (Bảng 29)")
    print("      ✓ 5.1.6. Phân công người thực hiện")
    write_chapter5(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 5.1 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
