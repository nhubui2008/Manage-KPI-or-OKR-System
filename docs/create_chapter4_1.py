"""
Script tạo CHƯƠNG 4: THỰC THI MÃ NGUỒN - Phần 4.1 Tổ chức mã nguồn
Bao gồm: Sơ đồ thư mục dự án, Kiến trúc hệ thống (MVC + Service Layer) và Phân chia layer chi tiết
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    """Tiêu đề chương (CHƯƠNG X: ...) - căn giữa, in đậm, viết hoa (Heading 1)"""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_code_block(doc, code_lines):
    """Vẽ khối mã nguồn (thư mục) thụt lề, font chữ Courier New"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(16)
    pf.left_indent = Cm(1.5)
    
    # Vẽ khung xám nhẹ
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)

    run = p.add_run(code_lines)
    run.font.name = 'Courier New'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG CHƯƠNG 4 =====================

def write_chapter4_1(doc):
    """CHƯƠNG 4: THUỐC THI MÃ NGUỒN - 4.1. Tổ chức mã nguồn"""

    add_chapter_title(doc, "CHƯƠNG 4: THỰC THI MÃ NGUỒN")

    # ============================================================
    # 4.1. Tổ chức mã nguồn
    # ============================================================
    add_heading1(doc, "4.1. Tổ chức mã nguồn")

    # 4.1.1. Sơ đồ thư mục dự án
    add_heading2(doc, "4.1.1. Sơ đồ cấu trúc thư mục dự án")
    
    add_para(doc,
        "Dự án được phát triển dựa trên nền tảng .NET 10.0 MVC. Thư mục mã nguồn được tổ chức "
        "khoa học, phân tách rõ ràng giữa các phần giao diện (Views), điều hướng (Controllers), thực thể cơ sở "
        "dữ liệu (Models), dịch vụ xử lý logic (Services) và các thư viện tiện ích (Helpers). "
        "Dưới đây là sơ đồ cấu trúc thư mục chi tiết của dự án:"
    )

    folder_tree = (
        "Manage-KPI-or-OKR-System/                      # Thư mục gốc giải pháp\n"
        "├── Controllers/                                # Lớp điều hướng (22 controllers)\n"
        "│   ├── AIController.cs                         # API hỗ trợ chat, gợi ý và phân tích bằng AI\n"
        "│   ├── AuthController.cs                       # Điều phối Đăng nhập, OAuth2, OTP\n"
        "│   ├── DashboardController.cs                  # Điều phối Dashboard & Vẽ biểu đồ\n"
        "│   ├── KPIsController.cs                       # Điều phối Nghiệp vụ giao & duyệt KPI\n"
        "│   └── OKRsController.cs                       # Điều phối Thiết lập OKR đa cấp\n"
        "├── Models/                                     # Lớp thực thể CSDL (45 Entity Models)\n"
        "│   ├── ViewModels/                             # ViewModels truyền dữ liệu cho Views\n"
        "│   ├── KPI.cs                                  # Entity lưu trữ thông tin KPI chính\n"
        "│   ├── OKR.cs                                  # Entity lưu trữ mục tiêu OKR\n"
        "│   └── Employee.cs                             # Entity lưu trữ hồ sơ nhân viên\n"
        "├── Views/                                      # Giao diện Razor HTML (.cshtml)\n"
        "│   ├── Shared/                                 # Giao diện dùng chung cho hệ thống\n"
        "│   │   ├── _Layout.cshtml                      # Giao diện khung (Sidebar, Navbar)\n"
        "│   │   └── _AIChatWidget.cshtml                # Widget Trợ lý AI Gemini\n"
        "│   ├── Dashboard/                              # Thư mục Views trang chủ\n"
        "│   └── KPIs/                                   # Thư mục Views CRUD & Giao KPI\n"
        "├── Services/                                   # Lớp xử lý nghiệp vụ chính (Business Logic)\n"
        "│   ├── GeminiService.cs                        # Gọi và cấu hình tham số Gemini API\n"
        "│   ├── AIDataService.cs                        # Xử lý nạp dữ liệu ngữ cảnh (6 partial classes)\n"
        "│   ├── AIAlertService.cs                       # Quét và sinh cảnh báo rủi ro tự động\n"
        "│   ├── OKRProgressService.cs                   # Tính toán tiến độ OKR liên cấp\n"
        "│   └── NotificationService.cs                  # Gửi mail thông báo và deadline\n"
        "├── Helpers/                                    # Thư viện tiện ích (Utility classes)\n"
        "│   ├── AccessScopeHelper.cs                    # Ràng buộc phạm vi dữ liệu theo vai trò\n"
        "│   ├── WorkflowStatusHelper.cs                 # Quản lý trạng thái vòng đời KPI/OKR\n"
        "│   └── PermissionAuthorizationHelper.cs        # Hỗ trợ phân quyền phân cấp\n"
        "├── Filters/                                    # Bộ lọc phân quyền Claims\n"
        "│   └── HasPermissionAttribute.cs               # Authorize Filter kiểm soát 60 permissions\n"
        "├── Data/                                       # Cấu hình kết nối CSDL (DbContext)\n"
        "│   └── MiniERPDbContext.cs                     # EF Core DbContext mapping 45 thực thể\n"
        "├── wwwroot/                                    # Tĩnh nguyên (Static assets)\n"
        "│   ├── css/site.css                            # CSS tùy chỉnh của hệ thống (89KB)\n"
        "│   ├── js/site.js                              # JS xử lý Ajax và gọi AI widget\n"
        "│   └── lib/                                    # Thư viện ngoài (Bootstrap, ApexCharts)\n"
        "├── Program.cs                                  # Khởi chạy hệ thống & Cấu hình DI\n"
        "└── appsettings.json                            # Tệp cấu hình tham số kết nối CSDL & AI"
    )
    add_code_block(doc, folder_tree)

    # 4.1.2. Kiến trúc dự án và cách chia layer
    add_heading2(doc, "4.1.2. Kiến trúc hệ thống và phân chia Layer")
    
    add_para(doc,
        "Dự án được xây dựng dựa trên sự kết hợp giữa mô hình kiến trúc MVC (Model-View-Controller) "
        "truyền thống của ASP.NET Core và mô hình Layered Architecture (Kiến trúc phân tầng) hiện đại. "
        "Sự kết hợp này giúp phân tách rõ ràng trách nhiệm của từng thành phần, dễ bảo trì, dễ viết test case "
        "và thuận tiện cho việc tích hợp các API bên thứ ba (như Google Gemini API). Dự án chia làm 5 lớp chính:"
    )

    layers = [
        ("Presentation Layer (Lớp hiển thị & điều hướng)",
         "Bao gồm các Controllers và Views. Controllers đóng vai trò tiếp nhận HTTP requests từ client, "
         "gọi xuống lớp Services tương ứng để lấy dữ liệu, đổ vào ViewModels và trả về Razor Views (.cshtml) "
         "để kết xuất HTML ra trình duyệt. Phía client sử dụng JavaScript (Ajax) để gọi không đồng bộ lên "
         "AIController nhằm đem lại trải nghiệm mượt mà cho Widget Chat."),
        
        ("Service Layer (Lớp xử lý nghiệp vụ - Business Logic)",
         "Đây là 'trái tim' của hệ thống, chứa toàn bộ các business rules của doanh nghiệp. "
         "Services (như OKRProgressService, AIAlertService, GeminiService) trực tiếp xử lý các phép toán, "
         "cấu hình tham số Prompt và kiểm soát tiến độ. Đặc biệt, AIDataService được chia thành 6 tệp partial classes "
         "để tách biệt ngữ cảnh lấy dữ liệu (OKR context, Customer context, Alerts context...) giúp mã nguồn không bị phình to."),
        
        ("Data Access Layer (Lớp truy cập dữ liệu)",
         "Sử dụng Entity Framework Core 10 (ORM) để kết nối vật lý với SQL Server thông qua lớp MiniERPDbContext. "
         "Toàn bộ truy vấn SQL được viết dưới dạng LINQ giúp tối ưu hóa bảo mật (chống SQL Injection), "
         "và sử dụng cơ chế AsNoTracking cho các câu lệnh đọc dữ liệu (Read-only query) để đạt hiệu năng tối đa."),
        
        ("Security & Authorization Layer (Lớp bảo mật & kiểm soát quyền)",
         "Bảo mật hệ thống dựa trên cơ chế Claims-based Authentication. Lớp này sử dụng Filter HasPermissionAttribute "
         "để đánh dấu trực tiếp trên các Action Method của Controller. Khi một request được gửi lên, bộ lọc sẽ "
         "quét Claims của tài khoản xem có đủ permission trong số 60 permissions hệ thống không trước khi cho phép thực thi. "
         "Đồng thời, AccessScopeHelper đóng vai trò lọc dữ liệu trả về theo đúng cấp bậc của tác nhân."),
        
        ("Helper & Utility Layer (Lớp tiện ích phụ trợ)",
         "Chứa các hàm dùng chung trong hệ thống như CodeGeneratorHelper (tự động tạo mã EMP/DEPT theo quy tắc), "
         "KpiCheckInScheduleHelper (tính toán hạn chót check-in), PaginatedList (hỗ trợ phân trang danh sách lớn) "
         "và SeoHelper (tối ưu hóa SEO cho trang web)."),
    ]

    for title, desc in layers:
        add_bullet(doc, desc, bold_prefix=title)

    # 4.1.3. Cơ chế giao tiếp giữa các Layer
    add_heading2(doc, "4.1.3. Cơ chế giao tiếp và Dependency Injection")
    
    add_para(doc,
        "Để đảm bảo nguyên lý thiết kế Loose Coupling (liên kết lỏng lẻo) và Dependency Inversion (đảo ngược phụ thuộc), "
        "tất cả các thành phần trong Service Layer đều giao tiếp thông qua các Interfaces (như IAIDataService, IEmailService...). "
        "Tại file khởi chạy Program.cs, hệ thống thực hiện đăng ký Dependency Injection (DI) dưới dạng Scoped "
        "(ví dụ: builder.Services.AddScoped<IAIDataService, AIDataService>()) hoặc Singleton đối với các dịch vụ background."
    )

    add_para(doc,
        "Mô hình giao tiếp tuần tự của một request (ví dụ: yêu cầu tư vấn AI trên Dashboard) sẽ đi qua các bước:\n"
        "1. Trình duyệt gửi Ajax request tới Action trong AIController.\n"
        "2. HasPermissionAttribute xác thực quyền truy cập của người dùng.\n"
        "3. AIController gọi xuống IAIDataService để lấy context dữ liệu hiện tại từ Database.\n"
        "4. AIController gửi context đó sang IGeminiService để đóng gói prompt và gọi API Google Gemini.\n"
        "5. Gemini trả kết quả về Service, Service lưu lịch sử và AIController trả kết quả JSON về cho client hiển thị.",
        italic=True, space_before=4, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 4.1. Tổ chức mã nguồn")
    print("      ✓ 4.1.1. Sơ đồ thư mục dự án")
    print("      ✓ 4.1.2. Phân chia 5 Layer chính")
    print("      ✓ 4.1.3. Cơ chế giao tiếp & Dependency Injection")
    write_chapter4_1(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 4.1 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
