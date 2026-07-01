"""
Script tạo phần 6.4: Hướng dẫn sử dụng cho vai trò Nhân viên (Employee) cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=FONT_SIZE, bold=True)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG 6.4 =====================

def write_section_6_4(doc):
    """6.4. Hướng dẫn sử dụng cho vai trò Nhân viên (Employee)"""

    add_heading1(doc, "6.4. Hướng dẫn sử dụng cho vai trò Nhân viên (Employee)")

    add_para(doc,
        "Vai trò Nhân viên (Employee) tập trung hoàn toàn vào việc thực thi các chỉ tiêu hiệu suất KPI được giao, "
        "thực hiện báo cáo check-in tiến độ đúng thời hạn, cập nhật các đầu việc trên bảng Kanban "
        "và tương tác với Trợ lý AI để tối ưu hóa hiệu quả công việc."
    )

    # 6.4.1. Tổng quan các nhiệm vụ chính
    add_heading2(doc, "6.4.1. Tổng quan nhiệm vụ của Nhân viên")
    add_para(doc, "Nhân viên thực hiện các quy trình thao tác nghiệp vụ hàng ngày bao gồm:")
    add_bullet(doc, "Thực hiện nhập giá trị đạt được thực tế, giải trình khó khăn của KPI định kỳ.", bold_prefix="Báo cáo check-in tiến độ")
    add_bullet(doc, "Nhận các thẻ công việc được giao, cập nhật tiến độ công việc trên bảng Kanban board.", bold_prefix="Cập nhật trạng thái công việc")
    add_bullet(doc, "Sử dụng widget AI để chat hỏi đáp chuyên môn, xin giải pháp khắc phục chậm trễ KPI.", bold_prefix="Hỏi đáp & Nhận gợi ý từ AI Gemini")
    add_bullet(doc, "Tra cứu biên bản họp 1-on-1 đã thống nhất, theo dõi các Action Items được giao.", bold_prefix="Theo dõi kế hoạch họp 1-on-1")
    add_bullet(doc, "Xem kết quả điểm số đánh giá, xếp hạng rank hiệu suất và thưởng cuối kỳ sau khi Giám đốc chốt.", bold_prefix="Tra cứu kết quả hiệu suất & Thưởng")

    # 6.4.2. Hướng dẫn thao tác chi tiết
    add_heading2(doc, "6.4.2. Hướng dẫn các thao tác thực thi chính")

    # Thao tác 1
    add_heading3(doc, "a) Thực hiện check-in tiến độ KPI (URL: /KPICheckIns/Create)")
    add_para(doc,
        "Nhân viên báo cáo kết quả thực thi các chỉ tiêu KPI định kỳ theo tần suất quy định:\n"
        "1. Đăng nhập vai trò Nhân viên. Vào mục 'KPI của tôi' trên thanh Menu.\n"
        "2. Click chọn chỉ tiêu KPI cần cập nhật, nhấn nút 'Báo cáo check-in' (URL: `/KPICheckIns/Create`).\n"
        "3. Nhập giá trị thực tế đạt được tại ô 'Achieved Value' (Ví dụ: đã hoàn thành 5 tài liệu, nhập số '5'). "
        "Hệ thống sẽ tự động tính toán tỷ lệ hoàn thành phần trăm.\n"
        "4. Nếu chỉ tiêu bị chậm hoặc gặp khó khăn, nhập chi tiết tại ô 'Barriers' (Rào cản/Khó khăn) để "
        "người quản lý hiểu rõ ngữ cảnh và có giải pháp hỗ trợ kịp thời.\n"
        "5. Click 'Gửi báo cáo'. Hệ thống ghi nhận trạng thái 'Chờ duyệt' (Pending Review) và tự động "
        "gửi email thông báo cho Trưởng phòng trực tiếp duyệt.",
        italic=False
    )

    # Thao tác 2
    add_heading3(doc, "b) Kéo thả cập nhật công việc Kanban (URL: /WorkProjects)")
    add_para(doc,
        "Nhân viên quản lý và báo cáo trạng thái các đầu việc chi tiết trên bảng Kanban:\n"
        "1. Truy cập mục 'Dự án & Kanban' (URL: `/WorkProjects`). Click chọn dự án phòng ban.\n"
        "2. Hệ thống hiển thị bảng Kanban board chứa các thẻ công việc được giao cho nhân viên (có ảnh avatar định danh).\n"
        "3. Khi bắt đầu thực hiện một công việc: Click chuột giữ thẻ công việc tại cột To Do và kéo thả sang cột In Progress.\n"
        "4. Khi hoàn thành công việc: Kéo thả thẻ từ cột In Progress sang cột Done. Hệ thống sẽ tự động ghi nhận thời gian "
        "hoàn thành thực tế và đồng bộ tiến độ của KPI liên kết tương ứng.",
        italic=False
    )

    # Thao tác 3
    add_heading3(doc, "c) Hội thoại và nhận trợ giúp từ Trợ lý AI Gemini (Widget)")
    add_para(doc,
        "Nhân viên tương tác với Trợ lý AI Gemini để tìm kiếm giải pháp tháo gỡ khó khăn trong công việc:\n"
        "1. Tại bất kỳ trang nào của hệ thống, click biểu tượng Chatbot (Bizen AI Widget) floating ở góc dưới bên phải màn hình.\n"
        "2. Khung chat trượt mở. Hệ thống tự động đính kèm context công việc hiện tại của nhân viên gửi đi. "
        "Nhân viên có thể click nút tắt nhanh 'Tư vấn cải thiện KPI chậm hạn' hoặc nhập câu hỏi tự nhiên "
        "(Ví dụ: 'Tôi đang bị chậm tiến độ KPI viết code do thiếu nhân sự, AI hãy đề xuất giải pháp').\n"
        "3. AI Gemini xử lý ngữ cảnh và trả về câu trả lời chi tiết: gợi ý cách tối ưu quy trình làm việc, "
        "hoặc đề xuất kịch bản trao đổi 1-on-1 với quản lý để xin hỗ trợ. "
        "Mọi lịch sử hội thoại sẽ được mã hóa lưu trữ phục vụ việc tiếp tục ngữ cảnh ở các lần mở sau.",
        italic=False
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 6.4. Vai trò Nhân viên (Employee)")
    print("      ✓ 6.4.1. Tổng quan nhiệm vụ")
    print("      ✓ 6.4.2. Thao tác chính (Check-in, Kanban, Trợ lý AI Gemini)")
    write_section_6_4(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 6.4 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
