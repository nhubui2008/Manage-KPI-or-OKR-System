"""
Script tạo phần 1.4: Khảo sát cho báo cáo tốt nghiệp
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell(cell, text, bold=False, size=FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell(cell, h, bold=True, color=RGBColor(255, 255, 255))
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        for c_idx, text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cell, text, size=12, align=align)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    return table


def add_table_caption(doc, caption):
    """Thêm caption cho bảng (Heading style 'Caption' + SEQ field cho Word) """
    import re
    m = re.match(r"^(Bảng|Hình)\s+(\d+)\s*:\s*(.*)$", caption, re.IGNORECASE)
    
    p = doc.add_paragraph(style='Caption')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(12)
    pf.keep_with_next = True
    
    if m:
        label = m.group(1) # Bảng hoặc Hình
        num = m.group(2)
        desc = m.group(3)
        
        # Thêm nhãn
        r1 = p.add_run(label + " ")
        set_font(r1, size=12, italic=True)
        r1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Thêm SEQ field
        run_seq = p.add_run()
        set_font(run_seq, size=12, italic=True)
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_seq._r.append(fldChar1)
        
        run_instr = p.add_run()
        set_font(run_instr, size=12, italic=True)
        run_instr.font.color.rgb = RGBColor(0, 0, 0)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f'SEQ {label} \\* ARABIC'
        run_instr._r.append(instrText)
        
        run_sep = p.add_run()
        set_font(run_sep, size=12, italic=True)
        run_sep.font.color.rgb = RGBColor(0, 0, 0)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._r.append(fldChar2)
        
        run_num = p.add_run(num)
        set_font(run_num, size=12, italic=True, bold=True)
        run_num.font.color.rgb = RGBColor(0, 0, 0)
        
        run_end = p.add_run()
        set_font(run_end, size=12, italic=True)
        run_end.font.color.rgb = RGBColor(0, 0, 0)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fldChar3)
        
        # Thêm mô tả
        r2 = p.add_run(": " + desc)
        set_font(r2, size=12, italic=True)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = p.add_run(caption)
        set_font(run, size=12, italic=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_figure_caption(doc, caption):
    """Thêm caption cho hình (Heading style 'Caption' + SEQ field cho Word) """
    return add_table_caption(doc, caption)


# ===================== NỘI DUNG 1.4 =====================

def write_section_1_4(doc):
    """1.4. Khảo sát"""

    add_heading1(doc, "1.4. Khảo sát")

    # ============================================================
    # 1.4.1. KHẢO SÁT CÁC HỆ THỐNG TƯƠNG TỰ
    # ============================================================
    add_heading2(doc, "1.4.1. Khảo sát các hệ thống tương tự")

    add_para(doc,
        "Để xây dựng một hệ thống tối ưu và học hỏi các mô hình quản trị hiệu quả, nhóm đã tiến hành "
        "nghiên cứu và khảo sát 3 nền tảng quản lý dự án, công việc và cơ sở dữ liệu phổ biến nhất hiện nay. "
        "Dưới đây là chi tiết phân tích ưu và nhược điểm của từng hệ thống:"
    )

    # --- Lark Base ---
    add_heading3(doc, "a) Lark Base (Lark Suite)")
    add_para(doc,
        "Lark Base là một cơ sở dữ liệu đa chiều nằm trong hệ sinh thái Lark Suite, cho phép người dùng tự "
        "thiết lập các bảng dữ liệu, luồng công việc tự động và liên kết linh hoạt giữa các phòng ban. "
        "Lark Base rất được các doanh nghiệp vừa và nhỏ ưa chuộng nhờ tính năng All-in-one."
    )
    add_bullet(doc, "Tích hợp sâu với Chatbot, Messenger, Calendar và Drive trong cùng một ứng dụng. Hỗ trợ tự động hóa luồng công việc (Workflows) rất tốt thông qua giao diện kéo thả.", bold_prefix="Ưu điểm")
    add_bullet(doc, "Hệ thống chỉ cung cấp cơ sở dữ liệu chung chung, không có nghiệp vụ chuyên biệt về đánh giá hiệu suất (xếp hạng tự động S-D) hay liên kết OKR 3 cấp chuyên sâu. Phân quyền đa cấp khá phức tạp đối với người dùng không chuyên.", bold_prefix="Nhược điểm")

    # --- Trello ---
    add_heading3(doc, "b) Trello")
    add_para(doc,
        "Trello là công cụ quản lý công việc cực kỳ phổ biến dựa trên phương pháp Kanban. Giao diện trực quan "
        "với các thẻ công việc (Cards) nằm trên các cột trạng thái (Lists) giúp người dùng dễ dàng theo dõi "
        "tiến độ dự án."
    )
    add_bullet(doc, "Giao diện tối giản, cực kỳ trực quan, dễ sử dụng cho mọi đối tượng. Thao tác kéo thả thẻ công việc mượt mà, hỗ trợ làm việc nhóm nhanh.", bold_prefix="Ưu điểm")
    add_bullet(doc, "Rất yếu trong việc liên kết công việc với mục tiêu chiến lược vĩ mô (OKR/KPI). Không hỗ trợ tính toán chỉ số hiệu suất tự động và phân quyền đa cấp chặt chẽ.", bold_prefix="Nhược điểm")

    # --- Jira Software ---
    add_heading3(doc, "c) Jira Software")
    add_para(doc,
        "Jira là hệ thống quản lý dự án mạnh mẽ nhất dành cho các đội ngũ phát triển phần mềm theo mô hình Agile. "
        "Hệ thống cung cấp các bộ lọc, biểu đồ báo cáo và quy trình làm việc (Workflow) cực kỳ chặt chẽ."
    )
    add_bullet(doc, "Tính năng quản lý công việc và phân cấp dự án rất sâu. Hỗ trợ tốt quy trình Agile Scrum/Kanban, báo cáo chi tiết về tốc độ làm việc (Velocity).", bold_prefix="Ưu điểm")
    add_bullet(doc, "Giao diện cực kỳ phức tạp và khó tiếp cận đối với các phòng ban phi kỹ thuật (HR, Marketing, Sales). Chi phí bản quyền rất cao và đòi hỏi người dùng phải qua đào tạo mới sử dụng thành thạo.", bold_prefix="Nhược điểm")

    # ============================================================
    # 1.4.2. KHẢO SÁT DOANH NGHIỆP VÀ NGƯỜI DÙNG MỤC TIÊU
    # ============================================================
    add_heading2(doc, "1.4.2. Khảo sát doanh nghiệp và người dùng mục tiêu")

    add_para(doc,
        "Để đảm bảo hệ thống giải quyết đúng các bài toán thực tế, nhóm đã tiến hành khảo sát trực tuyến "
        "qua Google Forms kết hợp phỏng vấn sâu 50 nhà quản lý (Director, Manager, HR) và 150 nhân viên "
        "đang làm việc tại các doanh nghiệp vừa và nhỏ (SMEs) tại Hà Nội. Dưới đây là kết quả khảo sát:"
    )

    # Bảng 8: Kết quả khảo sát
    headers_survey = ["STT", "Nội dung câu hỏi khảo sát", "Kết quả ghi nhận", "Nhận xét / Đánh giá của nhóm"]
    rows_survey = [
        ["1", "Công cụ hiện tại doanh nghiệp đang dùng để theo dõi KPI/OKR là gì?",
         "82% chọn Excel/Google Sheets.\n12% chọn Phần mềm CRM/ERP.\n6% chọn Giấy tờ/Email.",
         "Excel vẫn là công cụ thống trị do tính linh hoạt, nhưng gây phân tán dữ liệu và mất nhiều thời gian thu thập."],
        ["2", "Nhân viên có hiểu rõ công việc hàng ngày ảnh hưởng thế nào đến mục tiêu công ty?",
         "74% trả lời Không rõ ràng.\n26% trả lời Có hiểu sơ bộ.",
         "Sự đứt gãy thông tin giữa mục tiêu chiến lược và thực thi là nỗi đau phổ biến nhất ở các SME."],
        ["3", "Quản lý có nhận được cảnh báo khi nhân viên chậm tiến độ KPI không?",
         "91% trả lời Không có cảnh báo tự động.\n9% trả lời Tự kiểm tra thủ công.",
         "Thiếu cơ chế cảnh báo sớm dẫn đến việc phát hiện lỗi quá muộn khi kỳ đánh giá kết thúc."],
        ["4", "Mất bao lâu để HR tính toán xong điểm đánh giá và thưởng cuối kỳ?",
         "68% trả lời Mất từ 5 - 7 ngày.\n22% trả lời Trên 7 ngày.\n10% trả lời Dưới 3 ngày.",
         "Quy trình tính thưởng thủ công làm chậm trễ tiến độ chi trả thưởng và dễ xảy ra sai sót khi đối soát."],
        ["5", "Mức độ quan tâm đến việc tích hợp AI hỗ trợ gợi ý KPI và phân tích hiệu suất?",
         "85% Rất quan tâm.\n12% Quan tâm trung bình.\n3% Không quan tâm.",
         "Doanh nghiệp có nhu cầu cao về trợ lý AI giúp giảm tải việc soạn thảo KPI và viết nhận xét đánh giá."],
    ]

    create_table(doc, headers_survey, rows_survey, col_widths=[1.0, 5.0, 5.0, 5.0])
    add_table_caption(doc, "Bảng 8: Kết quả khảo sát thực trạng quản trị tại doanh nghiệp mục tiêu")

    # ============================================================
    # 1.4.3. BÀI HỌC RÚT RA CHO HỆ THỐNG
    # ============================================================
    add_heading2(doc, "1.4.3. Bài học rút ra cho hệ thống")

    add_para(doc,
        "Từ kết quả khảo sát thực tế và phân tích ưu nhược điểm của các hệ thống đi trước, "
        "nhóm NEXTGEN đã đúc rút được những bài học quan trọng để định hình và phát triển hệ thống của mình:"
    )

    lessons = [
        ("Tối giản hóa giao diện người dùng",
         "Giao diện cần trực quan và dễ tiếp cận như Trello. Nhân viên có thể cập nhật trạng thái "
         "công việc thông qua bảng Kanban kéo thả, giúp phòng ban phi kỹ thuật dễ dàng sử dụng."),
        ("Liên kết chặt chẽ ba cấp độ mục tiêu",
         "Hệ thống phải thiết lập mối liên kết bắt buộc: WorkItem (Công việc) -> KPI -> OKR Key Result. "
         "Khi nhân viên hoàn thành một công việc trên Kanban, tiến độ KPI và OKR tương ứng sẽ tự động "
         "được cập nhật theo thời gian thực."),
        ("Cơ chế cảnh báo rủi ro tự động (Smart Alerts)",
         "Hệ thống cần tự động gửi email/thông báo nhắc nhở khi deadline check-in sắp đến. Đồng thời, "
         "tích hợp AI để phát hiện và đưa ra cảnh báo sớm khi tốc độ hoàn thành KPI thực tế thấp hơn "
         "ngưỡng kỳ vọng."),
        ("Tối ưu hóa quy trình tính điểm và thưởng",
         "Module đánh giá phải xếp hạng tự động theo 7 bậc (S-D) dựa trên điểm số thực tế. "
         "Tích hợp module BonusRules để HR cấu hình quy tắc thưởng linh hoạt, xuất báo cáo "
         "Excel tự động để rút ngắn thời gian xử lý cuối kỳ từ 7 ngày xuống còn vài phút."),
        ("Tích hợp AI có nguồn như một trợ lý chuyên môn",
         "AI không chỉ là chatbot hỏi đáp thông thường mà phải có context-aware (đọc hiểu dữ liệu thực tế). "
         "AI hỗ trợ đề xuất KPI phù hợp cho Manager, phân tích hiệu suất tự động cuối kỳ và hỗ trợ viết "
         "nhận xét đánh giá (AI Review) nhằm tiết kiệm 80% thời gian soạn thảo."),
    ]

    for i, (title, desc) in enumerate(lessons, 1):
        add_bullet(doc, desc, bold_prefix=f"Bài học {i} – {title}")

    # Kết luận
    add_para(doc, "", space_before=6, space_after=0, indent=False)
    add_para(doc,
        "Tóm lại, phần khảo sát đã giúp nhóm NEXTGEN hiểu rõ \"nỗi đau\" thực tế của doanh nghiệp "
        "vừa và nhỏ, từ đó thiết kế một hệ thống không chỉ giải quyết bài toán số hóa dữ liệu "
        "mà còn mang lại giá trị vượt trội nhờ sự kết hợp giữa quản lý đa cấp chặt chẽ và trí tuệ nhân tạo.",
        italic=True, space_before=6, space_after=12
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("    ✓ 1.4. Khảo sát")
    print("      ✓ 1.4.1. Khảo sát các hệ thống tương tự (Lark Base, Trello, Jira)")
    print("      ✓ 1.4.2. Khảo sát doanh nghiệp/người dùng mục tiêu (Bảng khảo sát)")
    print("      ✓ 1.4.3. Bài học rút ra cho hệ thống")
    write_section_1_4(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 1.4 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
