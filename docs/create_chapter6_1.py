"""
Script tạo CHƯƠNG 6: HƯỚNG DẪN SỬ DỤNG - Phần 6.1 Vai trò Admin
Dựa trên cấu trúc chuẩn của báo cáo tốt nghiệp FPT Polytechnic
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ===================== CẤU HÌNH =====================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "BaoCao_DuAn_TotNghiep.docx")

FONT_NAME = 'Times New Roman'
FONT_SIZE = 14
HEADER_BG = "1F4E79"


def set_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_chapter_title(doc, text):
    """Tiêu đề chương (CHƯƠNG X: ...) - căn giữa, in đậm, viết hoa (Heading 1)"""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading1(doc, text):
    """Tiêu đề cấp 1 (1.1. ...) - in đậm (Heading 2)"""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading2(doc, text):
    """Tiêu đề cấp 2 (1.1.1. ...) - in đậm (Heading 3)"""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_heading3(doc, text):
    """Tiêu đề cấp 3 (1.1.1.1. ...) - in đậm (Heading 4)"""
    p = doc.add_paragraph(style='Heading 4')
    p.alignment = 0 # WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_para(doc, text, bold=False, italic=False, indent=True, space_before=3, space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(22)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(22)
    pf.left_indent = Cm(1.27 + level * 0.63)
    pf.first_line_indent = Cm(-0.63)

    if bold_prefix:
        run_b = p.add_run("● " + bold_prefix + ": ")
        set_font(run_b, bold=True)
        run_t = p.add_run(text)
        set_font(run_t)
    else:
        run = p.add_run("● " + text)
        set_font(run)
    return p


# ===================== NỘI DUNG CHƯƠNG 6 =====================

def write_chapter6_1(doc):
    """CHƯƠNG 6: HƯỚNG DẪN SỬ DỤNG - 6.1. Vai trò Admin"""

    add_page_break(doc)
    add_chapter_title(doc, "CHƯƠNG 6: HƯỚNG DẪN SỬ DỤNG HỆ THỐNG")

    # ============================================================
    # 6.1. Vai trò Admin (Quản trị viên hệ thống)
    # ============================================================
    add_heading1(doc, "6.1. Hướng dẫn sử dụng cho vai trò Quản trị viên (Admin)")

    add_para(doc,
        "Vai trò Quản trị viên (Admin) nắm giữ quyền hạn cao nhất về mặt cấu hình hệ thống phần cứng, "
        "quản trị tài khoản đăng nhập, cấu hình bảo mật phân quyền và giám sát nhật ký an toàn thông tin. "
        "Dưới đây là các tài liệu hướng dẫn vận hành chi tiết các chức năng thuộc quyền hạn của Admin."
    )

    # 6.1.1. Tổng quan các nhiệm vụ chính
    add_heading2(doc, "6.1.1. Tổng quan các nhiệm vụ chính của Admin")
    add_para(doc, "Admin chịu trách nhiệm vận hành 5 mảng chức năng kỹ thuật bao gồm:")
    add_bullet(doc, "Quản trị danh sách người dùng đăng nhập hệ thống, cấp phát tài khoản mới, đặt lại mật khẩu và khóa/mở khóa tài khoản.", bold_prefix="Quản lý tài khoản (SystemUsers)")
    add_bullet(doc, "Cấu hình quyền hạn (60 permissions) cho 5 vai trò mặc định (Admin, Director, Manager, HR, Employee) hoặc tạo thêm vai trò mới.", bold_prefix="Phân quyền vai trò (Roles & Permissions)")
    add_bullet(doc, "Thiết lập cấu hình Gmail SMTP và thời hạn dùng thử hệ thống (SaaS Trial); secret AI phải nằm trong secret store của môi trường.", bold_prefix="Cấu hình hệ thống (SystemParameters)")
    add_bullet(doc, "Tra cứu lịch sử thao tác của tất cả tài khoản trong hệ thống để phục vụ công tác rà soát lỗi và bảo mật an toàn thông tin.", bold_prefix="Giám sát nhật ký (AuditLogs)")
    add_bullet(doc, "Thống kê tình hình đăng ký của các doanh nghiệp khách hàng (Tenants), gia hạn gói cước dịch vụ.", bold_prefix="Quản trị SaaS & Thanh toán (SaaSAdmin)")

    # 6.1.2. Hướng dẫn các thao tác chính
    add_heading2(doc, "6.1.2. Hướng dẫn các thao tác kỹ thuật chính")

    # Thao tác 1
    add_heading3(doc, "a) Quản lý tài khoản người dùng (URL: /SystemUsers)")
    add_para(doc,
        "Chức năng này giúp Admin kiểm soát danh sách người dùng đăng nhập. "
        "Quy trình thực hiện bao gồm các bước sau:\n"
        "1. Từ Sidebar chính, Admin click chọn mục 'Quản lý tài khoản' (hoặc truy cập trực tiếp đường dẫn `/SystemUsers`).\n"
        "2. Hệ thống hiển thị danh sách tài khoản hiện tại kèm theo thông tin Username, Email, vai trò (Role), "
        "và trạng thái hoạt động (Active/Blocked).\n"
        "3. Để thêm tài khoản mới: Click nút 'Tạo tài khoản mới', nhập Email, Username, lựa chọn vai trò mặc định và click 'Lưu'. "
        "Hệ thống sẽ tự động gửi email chứa thông tin mật khẩu tạm thời cho người dùng.\n"
        "4. Để khóa tài khoản vi phạm: Click chọn tài khoản trong danh sách, tắt cờ 'IsActive' và click 'Lưu'. "
        "Tài khoản bị khóa sẽ lập tức bị logout khỏi tất cả thiết bị và không thể đăng nhập lại.",
        italic=False
    )

    # Thao tác 2
    add_heading3(doc, "b) Phân quyền và ma trận vai trò (URL: /Roles)")
    add_para(doc,
        "Chức năng phân quyền động cho phép Admin cấu hình chi tiết quyền hạn của từng nhóm vai trò "
        "nhằm đáp ứng chính xác yêu cầu bảo mật của doanh nghiệp:\n"
        "1. Truy cập mục 'Quản lý vai trò' (URL: `/Roles`).\n"
        "2. Click chọn nút 'Phân quyền' tại vai trò cần thay đổi (ví dụ: Trưởng phòng - Manager).\n"
        "3. Hệ thống hiển thị bảng ma trận chứa 60 permissions phân chia theo 8 phân hệ chức năng. "
        "Admin tiến hành tích chọn hoặc bỏ tích các quyền tương ứng (ví dụ: thêm quyền 'AI_KPI_GENERATOR_USE' cho Manager).\n"
        "4. Click 'Lưu thay đổi'. Hệ thống sẽ tự động cập nhật cache phân quyền. Người dùng thuộc vai trò đó "
        "sẽ được áp dụng quyền hạn mới ngay lập tức mà không cần đăng nhập lại.",
        italic=False
    )

    # Thao tác 3
    add_heading3(doc, "c) Quản lý tham số hệ thống (URL: /SystemParameters)")
    add_para(doc,
        "Admin thiết lập các tham số cấu hình toàn hệ thống giúp ứng dụng vận hành tự động:\n"
        "1. Truy cập mục 'Tham số hệ thống' (URL: `/SystemParameters`).\n"
        "2. Cấu hình các tham số SMTP: Nhập SMTP_Server (smtp.gmail.com), SMTP_Port (587), SMTP_Username "
        "và mật khẩu ứng dụng Gmail (App Password) để kích hoạt tính năng gửi mail tự động.\n"
        "3. Kiểm tra trạng thái tích hợp AI: secret provider và endpoint RAG được cấu hình ngoài database qua secret store của môi trường.\n"
        "4. Click 'Cập nhật tham số'. Hệ thống sẽ tự động tải lại (Reload) các tham số này vào bộ nhớ cache RAM của Web server.",
        italic=False
    )

    # Thao tác 4
    add_heading3(doc, "d) Giám sát nhật ký hệ thống (URL: /AuditLogs)")
    add_para(doc,
        "Tính năng Audit Logs ghi nhận toàn bộ hoạt động trong hệ thống giúp truy vết lỗi và phát hiện xâm nhập:\n"
        "1. Truy cập mục 'Nhật ký hệ thống' (URL: `/AuditLogs`).\n"
        "2. Hệ thống hiển thị danh sách nhật ký gồm: Thời gian thực hiện, ID người dùng, Tên tài khoản, "
        "Hành động thực hiện (Ví dụ: 'Đăng nhập thành công', 'Xóa KPI ID 105', 'Gọi AI gợi ý KPI'), "
        "địa chỉ IP Client và trình duyệt sử dụng.\n"
        "3. Admin có thể sử dụng bộ lọc tìm kiếm nhanh theo Username hoặc lọc theo khoảng thời gian "
        "để nhanh chóng khoanh vùng các hoạt động bất thường.",
        italic=False
    )

    # Thao tác 5
    add_heading3(doc, "e) Quản lý gói cước dịch vụ SaaS (URL: /SaaSAdmin)")
    add_para(doc,
        "Khi triển khai hệ thống dưới mô hình SaaS đa doanh nghiệp, Admin sử dụng chức năng này để quản lý thuê bao:\n"
        "1. Truy cập trang điều hướng `/SaaSAdmin` để xem tổng quan biểu đồ tăng trưởng doanh nghiệp đăng ký "
        "và biểu đồ doanh thu thanh toán theo tháng.\n"
        "2. Tại thẻ 'Danh sách Doanh nghiệp', Admin giám sát thông tin các công ty đăng ký dùng thử (Trial), "
        "thời điểm kết thúc dùng thử (TrialEndTime).\n"
        "3. Khi doanh nghiệp chuyển khoản thanh toán gia hạn, Admin click chọn doanh nghiệp, chọn gói dịch vụ "
        "(Standard, Premium, Enterprise), nhập số tháng gia hạn và click 'Gia hạn gói cước'. Hệ thống sẽ tự động "
        "mở khóa đầy đủ các tính năng tương ứng cho tenant đó.",
        italic=False
    )


# ===================== MAIN =====================

def main():
    print("📄 Đang mở file Word hiện có...")
    doc = Document(INPUT_PATH)

    print("  ✓ CHƯƠNG 6: HƯỚNG DẪN SỬ DỤNG HỆ THỐNG")
    print("    ✓ 6.1. Vai trò Admin (Quản trị viên)")
    print("      ✓ 6.1.1. Tổng quan nhiệm vụ")
    print("      ✓ 6.1.2. Thao tác chính (Tài khoản, Phân quyền, Tham số, Audit, SaaS)")
    write_chapter6_1(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅ Đã thêm phần 6.1 vào file Word!")
    print(f"   📄 {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
